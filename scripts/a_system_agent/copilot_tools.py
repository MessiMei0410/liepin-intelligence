"""Phase 2: ASA Copilot 工具系统。提供 Agent 可调用的数据库查询和业务操作能力。

每个工具由两部分组成：
1. OpenAI function calling schema（name/description/parameters）
2. 执行函数 execute_<name>(self, db_path, **kwargs) -> dict
"""

from __future__ import annotations

from typing import Any

# ---------------------------------------------------------------------------
# 工具 Schema 定义（OpenAI function calling 格式）
# ---------------------------------------------------------------------------

COPILOT_TOOLS: list[dict[str, Any]] = [
    {
        "type": "function",
        "function": {
            "name": "query_candidate",
            "description": "查询单个候选人的详细信息，包括当前阶段、匹配评分、风险点等。用于了解某个具体候选人的状态。",
            "parameters": {
                "type": "object",
                "properties": {
                    "candidate_id": {
                        "type": "integer",
                        "description": "候选人关系ID（job_candidate_id）",
                    },
                },
                "required": ["candidate_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "filter_candidates",
            "description": "按岗位专用硬性证据模型分级过滤候选池，当前支持机械、软件和电源岗位；未知岗位失败关闭，不套用其他岗位规则。输出 A-核心/A-强/B/C 名单并排除禁挖公司。",
            "parameters": {
                "type": "object",
                "properties": {
                    "job_id": {
                        "type": "integer",
                        "description": "岗位ID",
                    },
                    "limit": {
                        "type": "integer",
                        "description": "最多返回人数，默认 200",
                    },
                    "max_salary_k": {
                        "type": "integer",
                        "description": "期望月薪上限(K)。候选人期望薪资上限超过该值将归入 D-期望超限，默认不过滤",
                    },
                    "cities": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "允许的期望城市关键词列表（命中任一即保留），候选人期望城市不匹配将归入 D-城市不符，默认不过滤",
                    },
                },
                "required": ["job_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_job",
            "description": "查询岗位详情，包括候选池大小、当前状态、岗位要求等。",
            "parameters": {
                "type": "object",
                "properties": {
                    "job_id": {
                        "type": "integer",
                        "description": "岗位ID",
                    },
                },
                "required": ["job_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_candidates",
            "description": "搜索候选人池。可按客户、岗位、阶段、姓名关键词搜索。",
            "parameters": {
                "type": "object",
                "properties": {
                    "client": {
                        "type": "string",
                        "description": "客户名称（模糊匹配），如 '士兰微'",
                    },
                    "job_title": {
                        "type": "string",
                        "description": "岗位名称关键词（模糊匹配），如 'AE'",
                    },
                    "name": {
                        "type": "string",
                        "description": "候选人姓名关键词（模糊匹配）",
                    },
                    "stage": {
                        "type": "string",
                        "description": "阶段过滤，如 'S1 待复核'、'S2 待触达'",
                    },
                    "limit": {
                        "type": "integer",
                        "description": "最多返回数量，默认10，最大20",
                    },
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_dashboard",
            "description": "获取驾驶舱概览：活跃岗位数、候选人总数、待处理数量和异常列表。",
            "parameters": {
                "type": "object",
                "properties": {},
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_candidate_pipeline",
            "description": "获取某个岗位的候选人管道：列出该岗位下所有候选人及其阶段分布。",
            "parameters": {
                "type": "object",
                "properties": {
                    "job_id": {
                        "type": "integer",
                        "description": "岗位ID",
                    },
                },
                "required": ["job_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_knowledge",
            "description": "搜索ASA已保存的客户知识、岗位画像和经验记忆。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "搜索关键词",
                    },
                    "knowledge_type": {
                        "type": "string",
                        "description": "知识类型：client（客户知识）、job_profile（岗位画像）、experience（经验）、all（全部）",
                        "enum": ["client", "job_profile", "experience", "all"],
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "update_candidate_stage",
            "description": "更新候选人的流程阶段（clean_stage），并写入审计日志。",
            "parameters": {
                "type": "object",
                "properties": {
                    "candidate_id": {
                        "type": "integer",
                        "description": "候选人关系ID（job_candidate_id）",
                    },
                    "stage": {
                        "type": "string",
                        "description": "新阶段，如 'S3 已触达'、'S4 面试中'、'S5 offer'",
                    },
                    "note": {
                        "type": "string",
                        "description": "变更备注（可选）",
                    },
                },
                "required": ["candidate_id", "stage"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "record_communication",
            "description": "记录一次与候选人的沟通（电话、微信、邮件等）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "candidate_id": {
                        "type": "integer",
                        "description": "候选人关系ID（job_candidate_id）",
                    },
                    "channel": {
                        "type": "string",
                        "description": "沟通渠道，如 '电话'、'微信'、'邮件'、'猎聘'",
                    },
                    "summary": {
                        "type": "string",
                        "description": "沟通内容摘要",
                    },
                },
                "required": ["candidate_id", "channel", "summary"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_report",
            "description": "为某个岗位生成推荐报告（markdown），写入桌面客户项目目录。",
            "parameters": {
                "type": "object",
                "properties": {
                    "job_id": {
                        "type": "integer",
                        "description": "岗位ID",
                    },
                    "kind": {
                        "type": "string",
                        "description": "报告类型，默认 'recommendation'",
                    },
                },
                "required": ["job_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "memorize",
            "description": "把一条事实、经验或偏好存入长期记忆（幂等：相同内容只更新不重复插入）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "content": {
                        "type": "string",
                        "description": "要记住的内容",
                    },
                    "kind": {
                        "type": "string",
                        "description": "记忆类型，如 'fact'、'preference'、'experience'，默认 'fact'",
                    },
                    "importance": {
                        "type": "integer",
                        "description": "重要度 1-5，默认 3",
                    },
                },
                "required": ["content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "recall",
            "description": "从长期记忆中按关键词召回相关条目，按重要度和时间排序。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "搜索关键词",
                    },
                    "limit": {
                        "type": "integer",
                        "description": "最多返回数量，默认 8",
                    },
                },
                "required": ["query"],
            },
        },
    },
]


# ---------------------------------------------------------------------------
# 工具执行函数
# ---------------------------------------------------------------------------

def _safe_query(conn, sql: str, params: tuple = ()) -> list[dict[str, Any]]:
    """安全查询：返回 dict 列表。"""
    cur = conn.execute(sql, params)
    cols = [desc[0] for desc in cur.description] if cur.description else []
    return [dict(zip(cols, row)) for row in cur.fetchall()]


def execute_filter_candidates(
    db_path: str,
    job_id: int,
    limit: int = 200,
    max_salary_k: int | None = None,
    cities: list[str] | None = None,
) -> dict[str, Any]:
    """按岗位硬性证据分级过滤候选池（调用 candidate_pool_filter，含禁挖排除）。

    max_salary_k: 期望月薪上限(K)，超出者归入 D-期望超限。
    cities: 允许的期望城市关键词，不匹配者归入 D-城市不符。
    """
    from .candidate_pool_filter import filter_job_candidates
    import sqlite3
    conn = sqlite3.connect(f"file:{db_path}?mode=ro", uri=True)
    conn.row_factory = sqlite3.Row
    try:
        jrow = conn.execute(
            "SELECT c.name AS client FROM jobs j JOIN clients c ON c.id=j.client_id WHERE j.id=?",
            (job_id,),
        ).fetchone()
    finally:
        conn.close()
    client = str(jrow["client"]) if jrow is not None else ""
    try:
        result = filter_job_candidates(
            db_path, job_id, client=client, max_candidates=int(limit or 200),
            max_salary_k=max_salary_k, cities=cities,
        )
    except Exception as exc:
        return {"success": False, "error": f"分级过滤失败: {exc}"}
    # 岗位级口径记忆：工具路径做过严格筛选同样登记，之后任意会话问名单默认给分级结果。
    from .copilot_sessions import set_job_list_filter
    set_job_list_filter(db_path, job_id, "grade_filter")
    candidates = result.get("candidates") or []
    grades: dict[str, list[dict[str, Any]]] = {}
    for c in candidates:
        grades.setdefault(c.get("grade") or "未知", []).append(c)
    return {
        "success": True,
        "data": {
            "job_id": job_id,
            "client": client,
            "total": result.get("total") or 0,
            "summary": {g: len(items) for g, items in grades.items()},
            "groups": [
                {"grade": g, "label": g, "candidates": items[:50]}
                for g, items in sorted(grades.items(), key=lambda kv: -len(kv[1]))
            ],
        },
    }


def execute_query_candidate(db_path: str, candidate_id: int) -> dict[str, Any]:
    """查询候选人详情。"""
    import sqlite3
    conn = sqlite3.connect(str(db_path), timeout=5)
    conn.row_factory = sqlite3.Row
    try:
        row = conn.execute(
            """SELECT jc.id, jc.clean_stage, jc.stop_reason,
                      p.display_name as candidate_name,
                      j.title as job_title, cl.name as client_name
               FROM job_candidates jc
               LEFT JOIN people p ON p.id = jc.person_id
               LEFT JOIN jobs j ON j.id = jc.job_id
               LEFT JOIN clients cl ON cl.id = j.client_id
               WHERE jc.id = ?""",
            (candidate_id,),
        ).fetchone()
        if not row:
            return {"success": False, "error": f"未找到候选人 #{candidate_id}"}
        return {
            "success": True,
            "data": {
                "id": row["id"],
                "name": row["candidate_name"] or "未知",
                "stage": row["clean_stage"] or "未知",
                "job": row["job_title"] or "未分配",
                "client": row["client_name"] or "未知",
                "stopped": bool(row["stop_reason"]),
                "stop_reason": row["stop_reason"] or "",
            },
        }
    finally:
        conn.close()


def execute_query_job(db_path: str, job_id: int) -> dict[str, Any]:
    """查询岗位详情。"""
    import sqlite3
    conn = sqlite3.connect(str(db_path), timeout=5)
    conn.row_factory = sqlite3.Row
    try:
        row = conn.execute(
            """SELECT j.id, j.title, j.status, j.summary, j.lifecycle_stage,
                      cl.name as client_name,
                      COUNT(jc.id) as candidate_count,
                      SUM(CASE WHEN jc.clean_stage LIKE '%触达%' OR jc.clean_stage LIKE '%面试%' OR jc.clean_stage LIKE '%offer%' THEN 1 ELSE 0 END) as active_count
               FROM jobs j
               LEFT JOIN clients cl ON cl.id = j.client_id
               LEFT JOIN job_candidates jc ON jc.job_id = j.id
               WHERE j.id = ?
               GROUP BY j.id""",
            (job_id,),
        ).fetchone()
        if not row:
            return {"success": False, "error": f"未找到岗位 #{job_id}"}
        return {
            "success": True,
            "data": {
                "id": row["id"],
                "title": row["title"] or "未知",
                "client": row["client_name"] or "未知",
                "status": row["status"] or "未知",
                "lifecycle_stage": row["lifecycle_stage"] or "未知",
                "summary": (row["summary"] or "")[:500],
                "candidate_count": row["candidate_count"],
                "active_count": row["active_count"],
                "need_sourcing": row["active_count"] < 5,
            },
        }
    finally:
        conn.close()


def execute_search_candidates(
    db_path: str,
    client: str = "",
    job_title: str = "",
    name: str = "",
    stage: str = "",
    limit: int = 10,
) -> dict[str, Any]:
    """搜索候选人池。"""
    import sqlite3
    conn = sqlite3.connect(str(db_path), timeout=5)
    conn.row_factory = sqlite3.Row
    try:
        conditions = []
        params: list[Any] = []
        if client:
            conditions.append("cl.name LIKE ?")
            params.append(f"%{client}%")
        if job_title:
            conditions.append("j.title LIKE ?")
            params.append(f"%{job_title}%")
        if name:
            conditions.append("p.display_name LIKE ?")
            params.append(f"%{name}%")
        if stage:
            conditions.append("jc.clean_stage = ?")
            params.append(stage)
        where = " AND ".join(conditions) if conditions else "1=1"
        limit = min(max(1, limit), 20)
        rows = conn.execute(
            f"""SELECT jc.id, p.display_name as candidate_name, jc.clean_stage,
                       j.title as job_title, cl.name as client_name
                FROM job_candidates jc
                LEFT JOIN people p ON p.id = jc.person_id
                LEFT JOIN jobs j ON j.id = jc.job_id
                LEFT JOIN clients cl ON cl.id = j.client_id
                WHERE {where}
                ORDER BY jc.id DESC
                LIMIT ?""",
            params + [limit],
        ).fetchall()
        results = [
            {
                "id": row["id"],
                "name": row["candidate_name"] or "未知",
                "stage": row["clean_stage"] or "未知",
                "job": row["job_title"] or "未分配",
                "client": row["client_name"] or "未知",
            }
            for row in rows
        ]
        return {
            "success": True,
            "data": {
                "total": len(results),
                "candidates": results,
            },
        }
    finally:
        conn.close()


def execute_get_dashboard(db_path: str) -> dict[str, Any]:
    """获取驾驶舱概览。"""
    import sqlite3
    conn = sqlite3.connect(str(db_path), timeout=5)
    try:
        active_jobs = conn.execute(
            "SELECT COUNT(*) FROM jobs WHERE status NOT IN ('closed','archived')"
        ).fetchone()[0]
        total_candidates = conn.execute(
            "SELECT COUNT(*) FROM job_candidates"
        ).fetchone()[0]
        pending = conn.execute(
            "SELECT COUNT(*) FROM job_candidates WHERE clean_stage IN ('S1 待复核','S2 待触达')"
        ).fetchone()[0]
        interviews = conn.execute(
            "SELECT COUNT(*) FROM job_candidates WHERE clean_stage LIKE '%面试%'"
        ).fetchone()[0]
        # 候选池不足的岗位（活跃候选 < 5）
        low_pool = conn.execute(
            """SELECT COUNT(*) FROM (
                SELECT j.id FROM jobs j
                LEFT JOIN job_candidates jc ON jc.job_id = j.id
                WHERE j.status NOT IN ('closed','archived')
                GROUP BY j.id HAVING COUNT(jc.id) < 5
            )"""
        ).fetchone()[0]
        return {
            "success": True,
            "data": {
                "active_jobs": active_jobs,
                "total_candidates": total_candidates,
                "pending_review_or_outreach": pending,
                "in_interview": interviews,
                "jobs_with_low_pool": low_pool,
            },
        }
    finally:
        conn.close()


def execute_get_candidate_pipeline(db_path: str, job_id: int) -> dict[str, Any]:
    """获取岗位候选人管道分布。"""
    import sqlite3
    conn = sqlite3.connect(str(db_path), timeout=5)
    conn.row_factory = sqlite3.Row
    try:
        rows = conn.execute(
            """SELECT jc.id, p.display_name as candidate_name, jc.clean_stage,
                      jc.updated_at
               FROM job_candidates jc
               LEFT JOIN people p ON p.id = jc.person_id
               WHERE jc.job_id = ?
               ORDER BY jc.clean_stage, jc.updated_at DESC""",
            (job_id,),
        ).fetchall()
        # 按阶段分组
        stages: dict[str, list[dict[str, Any]]] = {}
        for row in rows:
            stage = row["clean_stage"] or "未分类"
            stages.setdefault(stage, []).append({
                "id": row["id"],
                "name": row["candidate_name"] or "未知",
            })
        return {
            "success": True,
            "data": {
                "job_id": job_id,
                "total": len(rows),
                "stages": {k: len(v) for k, v in stages.items()},
                "pipeline": stages,
            },
        }
    finally:
        conn.close()


def execute_search_knowledge(
    db_path: str, query: str, knowledge_type: str = "all"
) -> dict[str, Any]:
    """搜索知识库。"""
    import sqlite3
    conn = sqlite3.connect(str(db_path), timeout=5)
    conn.row_factory = sqlite3.Row
    try:
        results: list[dict[str, Any]] = []
        # 从 agent_memories 表搜索
        try:
            type_filter = ""
            type_params: list[Any] = []
            if knowledge_type != "all":
                type_filter = "AND memory_type = ?"
                type_params = [knowledge_type]
            rows = conn.execute(
                f"""SELECT id, memory_type, content, client, job, created_at
                    FROM agent_memories
                    WHERE content LIKE ? {type_filter}
                    ORDER BY created_at DESC LIMIT 5""",
                [f"%{query}%"] + type_params,
            ).fetchall()
            for row in rows:
                results.append({
                    "id": row["id"],
                    "type": row["memory_type"] or "unknown",
                    "content": (row["content"] or "")[:300],
                    "client": row["client"] or "",
                    "job": row["job"] or "",
                })
        except sqlite3.OperationalError:
            pass  # 表可能还不存在
        return {
            "success": True,
            "data": {
                "query": query,
                "total": len(results),
                "results": results,
            },
        }
    finally:
        conn.close()


# 工具执行路由表
TOOL_EXECUTORS = {
    "query_candidate": execute_query_candidate,
    "filter_candidates": execute_filter_candidates,
    "query_job": execute_query_job,
    "search_candidates": execute_search_candidates,
    "get_dashboard": execute_get_dashboard,
    "get_candidate_pipeline": execute_get_candidate_pipeline,
    "search_knowledge": execute_search_knowledge,
}


# ---- Phase 3: 主动建议引擎 ----

def generate_proactive_suggestions(db_path: str) -> list[dict[str, Any]]:
    """扫描业务状态，生成主动建议列表。返回去重后的建议卡片。"""
    import sqlite3
    conn = sqlite3.connect(str(db_path), timeout=5)
    conn.row_factory = sqlite3.Row
    suggestions: list[dict[str, Any]] = []
    try:
        # 1. 候选池不足的岗位（活跃候选 < 5）
        low_pool_jobs = conn.execute(
            """SELECT j.id, j.title, cl.name as client, COUNT(jc.id) as cnt
               FROM jobs j
               JOIN clients cl ON cl.id = j.client_id
               LEFT JOIN job_candidates jc ON jc.job_id = j.id
               WHERE j.status NOT IN ('closed','archived')
               GROUP BY j.id HAVING cnt < 5
               ORDER BY cnt ASC LIMIT 3"""
        ).fetchall()
        for job in low_pool_jobs:
            suggestions.append({
                "id": f"source_{job['id']}",
                "type": "source_candidates",
                "priority": "high",
                "message": f"{job['client']} {job['title']} 候选池仅 {job['cnt']} 人，需要寻访补充",
                "action": {"type": "start_sourcing", "job_id": job["id"], "label": "开始寻访"},
            })

        # 2. 长时间待触达的候选人（> 7 天）
        stale = conn.execute(
            """SELECT jc.id, p.display_name as name, jc.clean_stage, jc.updated_at,
                      j.title as job_title, cl.name as client
               FROM job_candidates jc
               LEFT JOIN people p ON p.id = jc.person_id
               LEFT JOIN jobs j ON j.id = jc.job_id
               LEFT JOIN clients cl ON cl.id = j.client_id
               WHERE jc.clean_stage = 'S2 待触达'
                 AND jc.updated_at < datetime('now', '-7 days', 'localtime')
               ORDER BY jc.updated_at ASC LIMIT 3"""
        ).fetchall()
        for row in stale:
            suggestions.append({
                "id": f"outreach_{row['id']}",
                "type": "outreach_reminder",
                "priority": "medium",
                "message": f"{row['name']}（{row['client']} {row['job_title']}）已待触达超7天",
                "action": {"type": "open_candidate", "id": row["id"], "label": "查看人选"},
            })

        # 3. S1 待复核积压
        pending_review = conn.execute(
            "SELECT COUNT(*) as cnt FROM job_candidates WHERE clean_stage = 'S1 待复核'"
        ).fetchone()
        if pending_review and pending_review["cnt"] >= 5:
            suggestions.append({
                "id": "review_backlog",
                "type": "review_backlog",
                "priority": "high",
                "message": f"待复核候选人有 {pending_review['cnt']} 人，建议集中处理",
                "action": {"type": "open_queue", "label": "打开待复核队列"},
            })

        # 4. 最近7天无活动的活跃岗位
        inactive = conn.execute(
            """SELECT j.id, j.title, cl.name as client
               FROM jobs j
               JOIN clients cl ON cl.id = j.client_id
               WHERE j.status NOT IN ('closed','archived')
                 AND j.id NOT IN (
                   SELECT DISTINCT jc.job_id FROM job_candidates jc
                   WHERE jc.updated_at > datetime('now', '-7 days', 'localtime')
                 )
               LIMIT 2"""
        ).fetchall()
        for job in inactive:
            suggestions.append({
                "id": f"inactive_{job['id']}",
                "type": "job_inactive",
                "priority": "low",
                "message": f"{job['client']} {job['title']} 近7天无活动",
                "action": {"type": "open_job", "id": job["id"], "label": "打开岗位"},
            })

    finally:
        conn.close()

    # 去重：同类型只保留最高优先级
    seen_types: set[str] = set()
    deduped: list[dict[str, Any]] = []
    for s in sorted(suggestions, key=lambda x: {"high": 0, "medium": 1, "low": 2}.get(x["priority"], 3)):
        if s["type"] not in seen_types:
            seen_types.add(s["type"])
            deduped.append(s)
    return deduped[:5]  # 最多5条


# ---------------------------------------------------------------------------
# Phase 4: 写操作与记忆工具
# ---------------------------------------------------------------------------

def execute_update_candidate_stage(
    db_path: str, candidate_id: int, stage: str, note: str = ""
) -> dict[str, Any]:
    """更新候选人阶段，并写入审计日志。"""
    import json
    import sqlite3
    import uuid
    conn = sqlite3.connect(str(db_path), timeout=5)
    conn.row_factory = sqlite3.Row
    try:
        row = conn.execute(
            "SELECT id, clean_stage FROM job_candidates WHERE id = ?",
            (candidate_id,),
        ).fetchone()
        if not row:
            return {"success": False, "error": f"未找到候选人 #{candidate_id}"}
        before_stage = row["clean_stage"]
        cur = conn.execute(
            "UPDATE job_candidates SET clean_stage = ?, updated_at = datetime('now','localtime') WHERE id = ?",
            (stage, candidate_id),
        )
        updated_rows = cur.rowcount
        conn.execute(
            """INSERT INTO audit_events
               (event_id, actor, surface, request_id, operation, target_type, target_id,
                before_json, after_json, result)
               VALUES (?, 'copilot', 'copilot_tool', ?, 'update_candidate_stage',
                       'job_candidate', ?, ?, ?, 'success')""",
            (
                uuid.uuid4().hex,
                uuid.uuid4().hex,
                str(candidate_id),
                json.dumps({"stage": before_stage}, ensure_ascii=False),
                json.dumps({"stage": stage, "note": note}, ensure_ascii=False),
            ),
        )
        conn.commit()
        return {"success": True, "updated_rows": updated_rows}
    finally:
        conn.close()


def execute_record_communication(
    db_path: str, candidate_id: int, channel: str, summary: str
) -> dict[str, Any]:
    """记录一次候选人沟通。表不存在则自动创建。"""
    import sqlite3
    conn = sqlite3.connect(str(db_path), timeout=5)
    try:
        conn.execute(
            """CREATE TABLE IF NOT EXISTS candidate_communications (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                candidate_id INTEGER NOT NULL,
                channel TEXT NOT NULL,
                summary TEXT,
                created_at TEXT DEFAULT (datetime('now','localtime'))
            )"""
        )
        cur = conn.execute(
            "INSERT INTO candidate_communications (candidate_id, channel, summary) VALUES (?, ?, ?)",
            (candidate_id, channel, summary),
        )
        conn.commit()
        return {"success": True, "communication_id": cur.lastrowid}
    finally:
        conn.close()


def execute_generate_report(
    db_path: str, job_id: int, kind: str = "recommendation"
) -> dict[str, Any]:
    """生成岗位推荐报告，写入 ~/Desktop/客户项目/{客户}/{岗位}/。
    kind='recommendation' 输出 .md；kind='docx' 输出排版 .docx（python-docx）。
    """
    import datetime
    import os
    import sqlite3
    conn = sqlite3.connect(str(db_path), timeout=5)
    conn.row_factory = sqlite3.Row
    try:
        job = conn.execute(
            """SELECT j.id, j.title, j.location, j.status, j.summary,
                      cl.name as client_name
               FROM jobs j
               LEFT JOIN clients cl ON cl.id = j.client_id
               WHERE j.id = ?""",
            (job_id,),
        ).fetchone()
        if not job:
            return {"success": False, "error": f"未找到岗位 #{job_id}"}
        candidates = conn.execute(
            """SELECT jc.id, jc.clean_stage, jc.updated_at,
                      p.display_name, p.current_company, p.current_title
               FROM job_candidates jc
               LEFT JOIN people p ON p.id = jc.person_id
               WHERE jc.job_id = ?
               ORDER BY jc.clean_stage, jc.updated_at DESC""",
            (job_id,),
        ).fetchall()
        client_name = job["client_name"] or "未知客户"
        job_title = job["title"] or f"岗位{job_id}"
        now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        lines = [
            f"# 推荐报告：{client_name} - {job_title}",
            "",
            f"- 报告类型：{kind}",
            f"- 岗位ID：{job_id}",
            f"- 地点：{job['location'] or '未知'}",
            f"- 状态：{job['status'] or '未知'}",
            f"- 生成时间：{now}",
            "",
            "## 岗位摘要",
            "",
            (job["summary"] or "（无）")[:1000],
            "",
            f"## 候选人列表（{len(candidates)} 人）",
            "",
            "| ID | 姓名 | 当前公司 | 当前职位 | 阶段 |",
            "| --- | --- | --- | --- | --- |",
        ]
        for c in candidates:
            lines.append(
                f"| {c['id']} | {c['display_name'] or '未知'} | {c['current_company'] or '-'} "
                f"| {c['current_title'] or '-'} | {c['clean_stage'] or '未知'} |"
            )
        lines.append("")
        report_dir = os.path.expanduser(f"~/Desktop/客户项目/{client_name}/{job_title}")
        os.makedirs(report_dir, exist_ok=True)
        if kind == "docx":
            path = os.path.join(report_dir, f"推荐报告_{job_id}.docx")
            _write_report_docx(path, client_name, job_title, job, candidates, kind, now)
        else:
            path = os.path.join(report_dir, f"推荐报告_{job_id}.md")
            with open(path, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
        return {"success": True, "path": path, "candidate_count": len(candidates)}
    finally:
        conn.close()


def _write_report_docx(
    path: str, client_name: str, job_title: str, job: Any,
    candidates: list[Any], kind: str, now: str,
) -> None:
    """用 python-docx 写推荐报告：标题 + 岗位摘要 + 候选人表格。"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading(f"推荐报告：{client_name} - {job_title}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta = doc.add_paragraph()
    for label, value in (
        ("报告类型", kind), ("岗位ID", job["id"]), ("地点", job["location"] or "未知"),
        ("状态", job["status"] or "未知"), ("生成时间", now),
    ):
        run = meta.add_run(f"{label}：{value}\n")
        run.font.size = Pt(9)
    doc.add_heading("岗位摘要", level=1)
    doc.add_paragraph((job["summary"] or "（无）")[:1000])
    doc.add_heading(f"候选人列表（{len(candidates)} 人）", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, header in enumerate(["ID", "姓名", "当前公司", "当前职位", "阶段"]):
        table.rows[0].cells[i].text = header
    for c in candidates:
        row = table.add_row().cells
        row[0].text = str(c["id"])
        row[1].text = c["display_name"] or "未知"
        row[2].text = c["current_company"] or "-"
        row[3].text = c["current_title"] or "-"
        row[4].text = c["clean_stage"] or "未知"
    doc.save(path)


def execute_memorize(
    db_path: str, content: str, kind: str = "fact", importance: int = 3
) -> dict[str, Any]:
    """写入长期记忆；content_hash 冲突时幂等更新 updated_at 并返回原 id。"""
    import hashlib
    import sqlite3
    content_hash = hashlib.sha256(content.encode("utf-8")).hexdigest()
    conn = sqlite3.connect(str(db_path), timeout=5)
    try:
        try:
            cur = conn.execute(
                """INSERT INTO agent_memories
                   (scope_type, memory_type, content, source_type, confidence, content_hash)
                   VALUES ('global', ?, ?, 'copilot', ?, ?)""",
                (kind, content, importance, content_hash),
            )
            conn.commit()
            return {"success": True, "memory_id": cur.lastrowid}
        except sqlite3.IntegrityError:
            row = conn.execute(
                "SELECT id FROM agent_memories WHERE content_hash = ?",
                (content_hash,),
            ).fetchone()
            conn.execute(
                "UPDATE agent_memories SET updated_at = datetime('now','localtime') WHERE id = ?",
                (row[0],),
            )
            conn.commit()
            return {"success": True, "memory_id": row[0]}
    finally:
        conn.close()


def execute_recall(db_path: str, query: str, limit: int = 8) -> dict[str, Any]:
    """按关键词召回长期记忆。"""
    import sqlite3
    conn = sqlite3.connect(str(db_path), timeout=5)
    conn.row_factory = sqlite3.Row
    try:
        limit = min(max(1, limit), 50)
        rows = conn.execute(
            """SELECT id, memory_type, content, confidence, created_at
               FROM agent_memories
               WHERE status = 'active' AND content LIKE '%' || ? || '%'
               ORDER BY confidence DESC, updated_at DESC
               LIMIT ?""",
            (query, limit),
        ).fetchall()
        memories = [
            {
                "id": row["id"],
                "memory_type": row["memory_type"],
                "content": row["content"],
                "confidence": row["confidence"],
                "created_at": row["created_at"],
            }
            for row in rows
        ]
        return {"success": True, "memories": memories}
    finally:
        conn.close()


TOOL_EXECUTORS.update({
    "update_candidate_stage": execute_update_candidate_stage,
    "record_communication": execute_record_communication,
    "generate_report": execute_generate_report,
    "memorize": execute_memorize,
    "recall": execute_recall,
})
