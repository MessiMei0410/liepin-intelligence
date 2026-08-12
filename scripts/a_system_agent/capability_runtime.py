from __future__ import annotations

import json
import hashlib
import os
import re
import secrets
import signal
import shutil
import sqlite3
import subprocess
import sys
import time
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Callable, TYPE_CHECKING

from .context import build_candidate_context
from .policy import is_stopped
from . import candidate_assessment, knowledge_base, negative_rules, query_builders, strategy_v2

if TYPE_CHECKING:
    from .service import AgentService


MULTICHANNEL = Path("/Users/messi/.codex/skills/multi-channel-search/scripts/a_system_multichannel.py")
LIEPIN_SEARCH = Path(__file__).resolve().parents[1] / "run_published_position_search.py"
RESUME_BACKFILL = Path(__file__).resolve().parents[1] / "backfill_pooled_resume_details.py"
XSAAS_SEARCH = Path(__file__).resolve().parents[1] / "xsaas_candidate_search.py"
OPENCLI_SHADOW = Path(__file__).resolve().parents[1] / "opencli_sourcing_shadow.py"
LIEPIN_OUTREACH = Path("/Users/messi/.codex/skills/liepin-cdp-search/scripts/liepin_im_followup.py")
LIEPIN_PUBLISH = Path("/Users/messi/.codex/skills/liepin-job-publish/scripts/liepin_publish_job.py")
MATCHING_REPORT = Path("/Users/messi/.codex/skills/candidate-matching-report/scripts/report_template.py")
JIASHI_REPORT = Path("/Users/messi/.codex/skills/jiashi-recommendation-report/scripts/fill_docx_template.py")
JIASHI_AUDIT = Path("/Users/messi/.codex/skills/jiashi-recommendation-report/scripts/audit_generated_report.py")
SALARY_REPORT = Path("/Users/messi/.codex/skills/candidate-salary-report/scripts/build_salary_report.py")
JIASHI_TEMPLATE = Path("/Users/messi/Desktop/嘉驰推荐报告/2026-06散落归档/嘉驰国际+客户名称--岗位名称--人选姓名（嘉驰模板）.docx")

DEFAULT_SOURCING_CELL_BATCH_SIZE = 8
MAX_SOURCING_CELL_BATCH_SIZE = 20
DEFAULT_PAGINATION_CONTINUATION_HEADROOM = 20
MAX_SOURCING_CONTINUATION_BATCHES = 256

# 注册在服务层（workflow_handler._execute_workflow_capability）实现、没有 run_* 确定性 Runner 的能力。
# 必须与 workflow_handler 的 locally_specialized 集合保持一致；注册期不变量
# assert_workflow_capabilities_resolvable 会在 AgentService 启动时校验漂移。
SERVICE_HANDLED_CAPABILITY_IDS: frozenset[str] = frozenset({
    "talent_pool_search",
    "candidate_batch_assessment",
    "candidate_pool_filter",
    "reply_triage",
    "communication_draft_batch",
})

# execute_external（后台渠道执行）当前支持的能力；只读/内部能力一律不允许后台渠道执行。
EXTERNAL_EXECUTION_CAPABILITY_IDS: frozenset[str] = frozenset({"multi_channel_sourcing"})


def assert_workflow_capabilities_resolvable(
    workflow_capabilities: list[Any],
    deterministic_runner_ids: set[str] | frozenset[str],
    service_handled_ids: set[str] | frozenset[str] = SERVICE_HANDLED_CAPABILITY_IDS,
) -> None:
    """注册期不变量：每个注册的工作流能力必须落到确定性 Runner（run_*）或服务层处理器。"""
    missing = sorted({
        str(item[0])
        for item in workflow_capabilities
        if str(item[0]) not in deterministic_runner_ids and str(item[0]) not in service_handled_ids
    })
    if missing:
        raise RuntimeError(
            "已注册工作流能力缺少执行实现（既无 run_* 确定性 Runner，也不在服务层处理器集合）："
            + "、".join(missing)
        )


def _loads(value: Any, default: Any) -> Any:
    if value in (None, ""):
        return default
    if isinstance(value, (dict, list)):
        return value
    try:
        return json.loads(str(value))
    except (TypeError, ValueError, json.JSONDecodeError):
        return default


def _row(row: sqlite3.Row | None) -> dict[str, Any]:
    return {key: row[key] for key in row.keys()} if row else {}


ZERO_RESULT_ATTRIBUTIONS = (
    "no_results",
    "session_expired",
    "compliance_wall",
    "parse_failure",
    "page_structure_changed",
    "loading_incomplete",
    "query_build_error",
    "pool_saturated",
    "unknown",
)

# 0 召回归因的中文解释（与 classify_zero_result 的枚举一一对应，单一来源）。
# 前端映射在仓外，含义必须与这里保持一致。
ZERO_RESULT_ATTRIBUTION_LABELS = {
    "session_expired": "登录态失效，需重新登录该渠道",
    "compliance_wall": "命中平台合规墙（需在浏览器里确认承诺函后重试）",
    "loading_incomplete": "页面加载未完成或查询未生效",
    "page_structure_changed": "页面结构变化，解析器需要适配",
    "parse_failure": "平台有结果但解析抓取失败",
    "no_results": "该渠道真实无匹配结果",
    "query_build_error": "查询构造异常",
    "pool_saturated": "本地人才库基本找遍了（重复率太高）",
    "unknown": "原因待排查",
}


def _round_int(entry: Any, key: str) -> int:
    if not isinstance(entry, dict):
        return 0
    try:
        return int(entry.get(key) or 0)
    except (TypeError, ValueError):
        return 0


def _trim_error(text: Any, limit: int = 1000) -> str:
    """错误文本截断：超长时保留尾部。

    traceback 的诊断行（异常类型 + 消息）在末尾，头部截断会把 ``X-SAAS_LOGIN_REQUIRED``
    这类归因信号切掉，使 classify_zero_result 系统性落到 unknown（T3 实战发现）。
    """
    value = str(text or "").strip()
    return value if len(value) <= limit else value[-limit:]


def _revision_consultant_evidence(context: dict[str, Any]) -> str:
    """Return only the consultant evidence embedded in a workflow revision instruction."""
    instruction = " ".join(str(context.get("revision_instruction") or "").split())
    if not instruction:
        return ""
    marker = "顾问已确认的原始条件："
    evidence = instruction.split(marker, 1)[1] if marker in instruction else instruction
    for suffix in ("。生成前必须逐项读取", "生成前必须逐项读取"):
        if suffix in evidence:
            evidence = evidence.split(suffix, 1)[0]
    return evidence.strip(" 。；;")


def _consultant_constraint_items(evidence: str) -> list[dict[str, str]]:
    """Extract modal constraints without asking the model to reinterpret their strength."""
    text = " ".join(str(evidence or "").split())
    if not text:
        return []
    candidates: list[tuple[str, str]] = []
    candidates.extend(("hard_requirement", value) for value in re.findall(r"必须[^，,；;。]+", text))
    for clause in re.split(r"[；;。]+", text):
        cleaned = clause.strip(" ，,")
        if not cleaned:
            continue
        if "优先" in cleaned or "更好" in cleaned:
            candidates.append(("preference", cleaned))
        if "可看" in cleaned or "可以看" in cleaned:
            candidates.append(("conditional_acceptance", cleaned))
        if any(token in cleaned for token in ("纠正", "术语", "不是", "不等于")):
            candidates.append(("consultant_wording", cleaned))
    rows: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()
    for constraint_type, rule in candidates:
        normalized = rule.strip(" ：:，,；;。")
        key = (constraint_type, normalized)
        if not normalized or key in seen:
            continue
        seen.add(key)
        rows.append({"type": constraint_type, "rule": normalized, "source": "consultant_revision"})
    return rows


def _lock_consultant_constraints(
    plan: dict[str, Any], strategy: dict[str, Any], evidence: str,
    locked_items: list[dict[str, str]] | None = None,
) -> list[dict[str, str]]:
    """Make revision semantics auditable and prevent model-generated constraint weakening."""
    constraints = _consultant_constraint_items(evidence)
    kind_map = {
        "must": "hard_requirement", "prefer": "preference",
        "allow": "conditional_acceptance", "exclude": "exclusion",
        "target_count": "target_count", "other": "consultant_wording",
    }
    seen = {(item["type"], item["rule"]) for item in constraints}
    for item in locked_items or []:
        if not isinstance(item, dict):
            continue
        rule = str(item.get("quote") or item.get("rule") or "").strip(" ：:，,；;。")
        constraint_type = kind_map.get(str(item.get("kind") or "other"), "consultant_wording")
        key = (constraint_type, rule)
        if not rule or key in seen:
            continue
        seen.add(key)
        constraints.append({"type": constraint_type, "rule": rule, "source": "copilot_verbatim"})
    if not constraints:
        return []
    strategy["consultant_constraints"] = constraints
    plan["consultant_constraints"] = constraints

    rules = [item["rule"] for item in constraints]
    essence = strategy.get("step1_job_essence")
    if isinstance(essence, dict):
        statement = str(essence.get("statement") or "").strip()
        locked = "顾问确认：" + "；".join(rules) + "。"
        essence["statement"] = f"{locked}{statement}" if locked not in statement else statement
        essence["confirmed_by"] = "consultant"

    hard_rules = [item["rule"] for item in constraints if item["type"] == "hard_requirement"]
    expectation = strategy.get("step5_expectation")
    if hard_rules and isinstance(expectation, dict):
        expectation["fallback_plan"] = (
            f"不得放宽顾问硬约束：{'；'.join(hard_rules)}；"
            "召回不足时只调整同义词、渠道或目标公司池，不降低门槛。"
        )

    trace = strategy.get("classification_trace")
    if isinstance(trace, list):
        trace.append(f"顾问修订约束已锁定 {len(constraints)} 项，模型不得弱化")
    summary = str(plan.get("strategy_summary") or "").strip()
    locked_summary = "顾问约束：" + "；".join(rules)
    plan["strategy_summary"] = f"{locked_summary}。{summary}" if locked_summary not in summary else summary
    return constraints


def _locked_constraint_conflicts(
    plan: dict[str, Any], strategy: dict[str, Any], constraints: list[dict[str, str]],
) -> list[str]:
    """Reject model output that changes the meaning of a locked consultant phrase."""
    if not constraints:
        return []
    serialized = json.dumps({"plan": plan, "strategy_v2": strategy}, ensure_ascii=False)
    errors: list[str] = []
    for item in constraints:
        rule = str(item.get("rule") or "").strip()
        if rule and rule not in serialized:
            errors.append(f"顾问原话约束未逐字保留：{rule}")
        if "三次电源" in rule and re.search(r"(?:三|3)次(?:以上|及以上|完整)", serialized):
            errors.append("领域术语“三次电源”被误解为次数条件")
    return list(dict.fromkeys(errors))


class CommandExecutionError(RuntimeError):
    def __init__(self, message: str, detail: dict[str, Any]):
        super().__init__(message)
        self.detail = detail


class ExternalPhaseError(RuntimeError):
    def __init__(
        self,
        message: str,
        *,
        phase: str,
        partial_result: dict[str, Any],
        detail: dict[str, Any],
    ):
        super().__init__(message)
        self.phase = phase
        self.partial_result = partial_result
        self.detail = detail


class ExternalExecutionCancelled(RuntimeError):
    """Raised after the workflow is stopped while a channel subprocess is active."""


def _json_object(value: Any) -> dict[str, Any]:
    try:
        parsed = json.loads(str(value or ""))
    except (TypeError, ValueError, json.JSONDecodeError):
        return {}
    return parsed if isinstance(parsed, dict) else {}


def _command_failure_summary(stdout: str, stderr: str, returncode: int) -> tuple[str, dict[str, Any]]:
    outer = _json_object(stdout) or _json_object(stderr)
    nested = _json_object(outer.get("stdout")) if outer else {}
    payload = nested or outer
    failed_checks = [
        check for check in payload.get("checks", [])
        if isinstance(check, dict) and check.get("ok") is False
    ] if isinstance(payload.get("checks"), list) else []
    labels = [
        str(check.get("message") or check.get("check") or "审计检查未通过").strip()
        for check in failed_checks[:3]
    ]
    if labels:
        summary = "；".join(dict.fromkeys(label for label in labels if label))
    else:
        raw = stderr or stdout or f"命令退出码 {returncode}"
        summary = _trim_error(raw, 600)
    detail = {
        "returncode": returncode,
        "command": outer.get("cmd") if outer else None,
        "failed_checks": [
            {
                "check": check.get("check"),
                "message": check.get("message"),
                "client": check.get("client"),
                "rows": check.get("rows", [])[:5] if isinstance(check.get("rows"), list) else [],
            }
            for check in failed_checks
        ],
        "stdout_tail": _trim_error(stdout, 2000),
        "stderr_tail": _trim_error(stderr, 2000),
    }
    return summary, detail


# 渠道查询方言层已模块化（S4-3c-2 / N1）：实现在 a_system_agent.query_builders。
# 以下常量与 adapt_channel_queries 是兼容别名/委托（既有测试与调用方签名不变），
# 新接线直接用 query_builders.build_liepin_queries / build_xsaas_queries。
XSAAS_QUERY_MAX_TERMS = query_builders.XSAAS_QUERY_MAX_TERMS
XSAAS_QUERY_MAX_COUNT = query_builders.XSAAS_QUERY_MAX_COUNT
LIEPIN_QUERY_MAX_TERMS = query_builders.LIEPIN_QUERY_MAX_TERMS
LIEPIN_QUERY_MAX_COUNT = query_builders.LIEPIN_QUERY_MAX_COUNT


def adapt_channel_queries(queries: list[Any], *, max_terms: int, max_count: int, company_terms: set[str] | None = None) -> list[str]:
    """兼容委托：实现已迁入 query_builders.adapt_queries（组合式方言，猎聘语法）。"""
    return query_builders.adapt_queries(queries, max_terms=max_terms, max_count=max_count, company_terms=company_terms)


# pool_saturated 阈值：dedupe_count/extracted_count 超过该比例视为本地池枯竭
# （F3 实证：猎聘 119/120=99% 排重仍原样重搜，系统无信号）。
POOL_SATURATED_DEDUPE_RATE = 0.9


def _has_query_build_error(rounds: list[dict[str, Any]], vocab: set[str]) -> bool:
    """查询构造错误信号（HTTP 成功但查询本身错了；任一 query 命中即成立）。

    a) "关键字："嵌套拼接（同一查询出现 ≥2 次，round7 真实形态：
       ``MPS 矽力杰 关键字：MPS 杰华特 关键字：MPS TME``），或查询间条件累加
       （后一条页面回显查询完整包含前一条且更长——条件未重置时 X-SaaS 的 selected_query 形态）；
    b) 单查询含 ≥2 个公司名（一人不可能同时在两家公司，组合语义必错；词表为空时该信号不启用）；
    c) 空查询，或 repr 残片（round6 真实形态：字典被 str() 当查询词，``{'evidence': ...`` 开头）。
    """
    selected_queries: list[str] = []
    for entry in rounds:
        query = str(entry.get("query") or "").strip()
        if "query" in entry and not query:
            return True
        selected = str(entry.get("selected_query") or "").strip()
        if selected:
            selected_queries.append(selected)
        for text in (query, selected):
            if not text:
                continue
            if text.count("关键字：") >= 2 or text.startswith("{'"):
                return True
            if vocab and sum(1 for token in text.split() if query_builders.is_company_token(token, vocab)) >= 2:
                return True
    return any(
        current != previous and current.startswith(previous)
        for previous, current in zip(selected_queries, selected_queries[1:])
    )


def classify_zero_result(
    channel: str,
    status: str,
    result: dict[str, Any],
    *,
    dedupe_count: int | None = None,
    company_vocab: set[str] | None = None,
) -> str:
    """用 runner 输出里的既有信号为 0 候选的渠道结果归因。

    信号来源（不改浏览器 runner 逻辑）：
    - 失败/阻断时 error 文本：X-SaaS 的 ``X-SAAS_LOGIN_REQUIRED``、猎聘的“登录已过期/登录态失效”
      → session_expired；“加载超时/未加载” → loading_incomplete。
    - per-query rounds：search_controls_missing（X-SaaS SEARCH_JS 返回的 reason）→ page_structure_changed；
      stale_query（查询未生效）→ loading_incomplete；平台有结果数但抓取 0 条 → parse_failure；
      全部查询平台返回 0 条，或抓到了但全部被评分门槛/排重过滤 → no_results。
    - query_build_error（S4-3c-1）：查询构造错误但 HTTP 成功——"关键字："嵌套拼接/查询间条件累加、
      单查询 ≥2 个公司名（company_vocab 判定）、空查询/repr 残片；任一 query 命中即归类该渠道。
    - pool_saturated（S4-3c-1）：召回正常（recall_count > 0）但排重率
      dedupe_count/extracted_count > 90%（调用方传入漏斗口径 dedupe_count，extracted 为 0 不计）。
    判定顺序：执行类（session/loading/结构）> query_build_error > pool_saturated
    > rounds 常规判定（parse_failure/no_results）> unknown；信号不足一律 unknown。
    """
    error = str(result.get("error") or "")
    if "LOGIN_REQUIRED" in error or "登录已过期" in error or "登录态失效" in error:
        return "session_expired"
    if "合规墙" in error or "合规承诺" in error or "compliancecommitment" in error.lower():
        return "compliance_wall"
    if "加载超时" in error or "未加载" in error:
        return "loading_incomplete"
    rounds = [entry for entry in result.get("rounds") or [] if isinstance(entry, dict)]
    if rounds:
        if any(str(entry.get("reason") or "") == "search_controls_missing" for entry in rounds):
            return "page_structure_changed"
        if any(str(entry.get("status") or "") == "stale_query" for entry in rounds):
            return "loading_incomplete"
        if any(str(entry.get("reason") or "") == "settle_timeout" for entry in rounds):
            return "loading_incomplete"
        if _has_query_build_error(rounds, company_vocab or set()):
            return "query_build_error"
        recall = sum(_round_int(entry, "result_count") for entry in rounds)
        extracted = sum(_round_int(entry, "extracted_count") for entry in rounds)
        if dedupe_count is not None and recall > 0 and extracted > 0 and dedupe_count / extracted > POOL_SATURATED_DEDUPE_RATE:
            return "pool_saturated"
        if recall > 0 and extracted <= 0:
            return "parse_failure"
        return "no_results"
    if status == "completed" and result.get("ok") is True and not error:
        # 查询成功但没有 per-query 明细（旧版 runner 输出），无法进一步区分。
        return "unknown"
    return "unknown"


def _slug(value: Any) -> str:
    text = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._-]+", "-", str(value or "").strip())
    return text.strip("-.")[:80] or "artifact"


def _list_text(value: Any) -> str:
    items = _loads(value, [])
    if not isinstance(items, list):
        return str(value or "")
    return "；".join(str(item).strip() for item in items if str(item).strip())


class RecruitingCapabilityRuntime:
    """Deterministic implementations behind the ASA capability registry."""

    def __init__(self, service: AgentService) -> None:
        self.service = service
        self.python = self._resolve_python()
        self.output_dir = service.db_path.parent / "asa_artifacts"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    @classmethod
    def deterministic_runner_ids(cls) -> frozenset[str]:
        """当前 run_* 确定性 Runner 覆盖的能力集合。"""
        return frozenset(
            name[len("run_"):]
            for name in dir(cls)
            if name.startswith("run_") and callable(getattr(cls, name, None))
        )

    def availability(self, capability_id: str | None = None) -> dict[str, Any]:
        """能力可用性与调用语义元数据：确定性 Runner / 服务层处理器 / 后台渠道执行支持。"""
        runner_ids = self.deterministic_runner_ids()
        known = sorted(runner_ids | set(SERVICE_HANDLED_CAPABILITY_IDS))
        selected = [str(capability_id)] if capability_id else known
        rows: list[dict[str, Any]] = []
        for cid in selected:
            spec = self.service.skills.get(cid) if hasattr(self.service, "skills") else None
            deterministic = cid in runner_ids
            service_handled = cid in SERVICE_HANDLED_CAPABILITY_IDS
            rows.append({
                "capability_id": cid,
                "registered": spec is not None,
                "deterministic_runner": deterministic,
                "service_handler": service_handled,
                "external_execution_supported": cid in EXTERNAL_EXECUTION_CAPABILITY_IDS,
                "execution_path": (
                    "deterministic_runner" if deterministic
                    else "service_handler" if service_handled
                    else "unsupported"
                ),
            })
        return {"ok": True, "capabilities": rows[0] if capability_id and len(rows) == 1 else rows}

    @staticmethod
    def _resolve_python() -> str:
        candidates = [
            os.environ.get("A_SYSTEM_PYTHON", ""),
            sys.executable,
            shutil.which("python3") or "",
            "/Library/Frameworks/Python.framework/Versions/3.11/bin/python3",
            "/opt/homebrew/bin/python3",
            "/usr/local/bin/python3",
        ]
        seen: set[str] = set()
        for candidate in candidates:
            if not candidate or candidate in seen or not Path(candidate).exists():
                continue
            seen.add(candidate)
            probe = subprocess.run(
                [candidate, "-c", "import openpyxl, docx"],
                capture_output=True, text=True, timeout=15,
            )
            if probe.returncode == 0:
                return candidate
        raise RuntimeError("ASA 找不到具备 openpyxl 和 python-docx 的 Python 运行环境")

    def execute(self, capability_id: str, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        handler = getattr(self, f"run_{capability_id}", None)
        if handler is None:
            spec = self.service.skills.get(capability_id) if capability_id in SERVICE_HANDLED_CAPABILITY_IDS else None
            if spec is not None:
                # 注册在服务层（workflow_handler）实现的能力：走已注册 handler，保证调用语义完整。
                return spec.handler(context, inputs)
            available = "、".join(sorted(self.deterministic_runner_ids() | set(SERVICE_HANDLED_CAPABILITY_IDS)))
            raise ValueError(f"能力没有可用的确定性 Runner 或服务层处理器：{capability_id}；可用能力：{available}")
        return handler(context, inputs)

    @staticmethod
    def _sourcing_score_thresholds(job: str, ability_terms: set[str]) -> tuple[int, int]:
        text = " ".join([str(job or ""), *(str(term or "") for term in ability_terms)]).casefold()
        markers = ("vpd", "vrm", "tlvr", "多相buck", "模块电源", "垂直供电")
        return (70, 80) if sum(marker.casefold() in text for marker in markers) >= 2 else (55, 65)

    @staticmethod
    def _channel_risk_stop_reason(result: dict[str, Any]) -> str:
        text = json.dumps(result, ensure_ascii=False).casefold()
        markers = (
            "安全风险", "风险提示", "安全合规", "合规墙", "compliancecommitment",
            "captcha", "人机验证", "操作频繁", "访问频繁", "账号风险",
        )
        return next((marker for marker in markers if marker.casefold() in text), "")

    @staticmethod
    def _channel_page_budget(request: dict[str, Any]) -> int:
        try:
            requested = int(request.get("max_pages_per_query") or 3)
        except (TypeError, ValueError):
            requested = 3
        return max(1, min(requested, 10))

    @staticmethod
    def _platform_capped_continuation_limit(request: dict[str, Any]) -> int:
        """Bound pagination retries so capped queries cannot keep re-entering the queue."""
        raw = request.get("max_platform_capped_retries")
        if raw not in (None, ""):
            try:
                return max(0, min(int(raw), 3))
            except (TypeError, ValueError):
                pass
        mode = str(request.get("liepin_risk_mode") or os.environ.get("ASA_LIEPIN_RISK_MODE", "low")).strip().lower()
        return 1 if mode in {"fast", "balanced_fast"} else 0

    @staticmethod
    def _liepin_detail_capture_options(request: dict[str, Any], target: int) -> tuple[int, list[str]]:
        """Low-risk fast mode: keep recall broad, throttle only resume detail pages."""
        mode = str(request.get("liepin_risk_mode") or os.environ.get("ASA_LIEPIN_RISK_MODE", "low")).strip().lower()
        if mode in {"fast", "balanced_fast"}:
            defaults = {"min": 1.8, "max": 4.0, "burst": 8, "cooldown": 8.0}
        elif mode in {"very_low", "safe", "conservative"}:
            defaults = {"min": 4.0, "max": 8.0, "burst": 4, "cooldown": 25.0}
        else:
            defaults = {"min": 2.8, "max": 6.2, "burst": 5, "cooldown": 16.0}

        def _int_value(key: str, fallback: int, minimum: int, maximum: int) -> int:
            raw = request.get(key) or os.environ.get(f"ASA_{key.upper()}")
            try:
                value = int(raw) if raw not in (None, "") else fallback
            except (TypeError, ValueError):
                value = fallback
            return max(minimum, min(value, maximum))

        def _float_value(key: str, fallback: float, minimum: float, maximum: float) -> float:
            raw = request.get(key) or os.environ.get(f"ASA_{key.upper()}")
            try:
                value = float(raw) if raw not in (None, "") else fallback
            except (TypeError, ValueError):
                value = fallback
            return max(minimum, min(value, maximum))

        fallback_limit = min(max(6, target), 12)
        detail_limit = _int_value("liepin_detail_limit", fallback_limit, 0, 40)
        min_delay = _float_value("liepin_detail_min_delay", float(defaults["min"]), 0.0, 30.0)
        max_delay = _float_value("liepin_detail_max_delay", float(defaults["max"]), min_delay, 45.0)
        burst_size = _int_value("liepin_detail_burst_size", int(defaults["burst"]), 1, 50)
        cooldown = _float_value("liepin_detail_burst_cooldown", float(defaults["cooldown"]), 0.0, 120.0)
        return detail_limit, [
            "--detail-min-delay", str(min_delay),
            "--detail-max-delay", str(max_delay),
            "--detail-burst-size", str(burst_size),
            "--detail-burst-cooldown", str(cooldown),
            "--stop-on-risk-page",
        ]

    def _external_request_cancelled(self, request: dict[str, Any]) -> bool:
        try:
            step_id = int(request.get("_workflow_step_id") or 0)
        except (TypeError, ValueError):
            return False
        execution_token = str(request.get("_workflow_execution_token") or "")
        return step_id > 0 and not self.service.workflow_engine.external_step_is_active(step_id, execution_token)

    def _ensure_external_request_active(self, request: dict[str, Any]) -> None:
        if self._external_request_cancelled(request):
            raise ExternalExecutionCancelled("工作流已停止，当前渠道执行已终止")

    def execute_external(self, capability_id: str, request: dict[str, Any]) -> dict[str, Any]:
        if capability_id not in EXTERNAL_EXECUTION_CAPABILITY_IDS:
            supported = "、".join(sorted(EXTERNAL_EXECUTION_CAPABILITY_IDS))
            raise ValueError(f"能力不支持后台渠道执行：{capability_id}；仅支持：{supported}")
        client, job = str(request.get("client") or ""), str(request.get("job") or "")
        if not client or not job:
            raise ValueError("寻访任务缺少客户或岗位")
        self._ensure_external_request_active(request)
        cancel_check = (
            (lambda: self._external_request_cancelled(request))
            if request.get("_workflow_step_id")
            else None
        )
        audit_only_result = request.get("_audit_only_result")
        if isinstance(audit_only_result, dict):
            sync_script = Path("/Users/messi/.codex/skills/a-system-workbench/scripts/a_system_sync.py")
            sync = self._run_external(
                [self.python, str(sync_script), "--client", client, "--job", job, "--no-open"],
                300,
                cancel_check=cancel_check,
            )
            return {
                **audit_only_result,
                "verified": True,
                "audit": {
                    "ok": True,
                    "summary": "A 系统收尾审计通过",
                    "returncode": sync.returncode,
                    "recovered_without_channel_rerun": True,
                },
            }
        approved_snapshot = request.get("strategy_snapshot") if isinstance(request.get("strategy_snapshot"), dict) else {}
        query_plan = request.get("query_plan_v1") if isinstance(request.get("query_plan_v1"), dict) else {}
        if not query_plan and isinstance(approved_snapshot.get("query_plan_v1"), dict):
            query_plan = approved_snapshot["query_plan_v1"]
        plan_ok, plan_errors = query_builders.validate_query_plan_v1(query_plan)
        approved_plan_hash = str(
            request.get("query_plan_hash") or approved_snapshot.get("query_plan_hash") or ""
        )
        if not plan_ok or not approved_plan_hash:
            detail = "；".join(plan_errors) if plan_errors else "缺少审批计划哈希"
            raise ValueError(f"缺少有效且批准的 query_plan_v1：{detail}")
        if not secrets.compare_digest(approved_plan_hash, str(query_plan.get("plan_hash") or "")):
            raise ValueError("批准的 query_plan_v1 哈希与执行请求不一致")
        target = max(1, min(int(request.get("target_count") or 10), 50))
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        run_id = str(request.get("resume_run_id") or f"asa-source-{stamp}")
        candidates_path = self.output_dir / "sourcing" / f"{_slug(client)}-{_slug(job)}-{stamp}.json"
        liepin_path = candidates_path.with_name(candidates_path.stem + "-liepin.json")
        xsaas_path = candidates_path.with_name(candidates_path.stem + "-xsaas.json")
        liepin_raw_path = candidates_path.with_name(candidates_path.stem + "-liepin-raw.json")
        xsaas_raw_path = candidates_path.with_name(candidates_path.stem + "-xsaas-raw.json")
        liepin_queries_path = candidates_path.with_name(candidates_path.stem + "-liepin-queries.json")
        xsaas_queries_path = candidates_path.with_name(candidates_path.stem + "-xsaas-queries.json")
        candidates_path.parent.mkdir(parents=True, exist_ok=True)
        strategy = request.get("strategy") if isinstance(request.get("strategy"), dict) else {}
        query_groups = approved_snapshot.get("query_groups") if isinstance(approved_snapshot.get("query_groups"), list) else []
        ability_terms = {
            str(term).strip()
            for group in query_groups
            if isinstance(group, dict)
            for term in (group.get("terms") or [])
            if str(term).strip()
        }
        quality_min_score, quality_recommend_score = self._sourcing_score_thresholds(job, ability_terms)
        page_budget = self._channel_page_budget(request)
        resume_requested = bool(request.get("resume_run_id"))
        all_runnable_cells = (
            self._resume_query_cells(
                run_id,
                query_plan,
                max_retries=int(request.get("max_query_retries") or 3),
                max_platform_capped_retries=self._platform_capped_continuation_limit(request),
            )
            if resume_requested
            else [cell for cell in query_plan.get("cells") or [] if isinstance(cell, dict)]
        )
        # Every approved cell may already be terminal when a paused workflow is
        # resumed. In that case, run the normal empty-batch intake/audit path so
        # the workflow can advance to assessment instead of failing at the cursor.
        try:
            cell_batch_size = max(
                1,
                min(
                    int(request.get("max_query_cells_per_batch") or DEFAULT_SOURCING_CELL_BATCH_SIZE),
                    MAX_SOURCING_CELL_BATCH_SIZE,
                ),
            )
        except (TypeError, ValueError):
            cell_batch_size = DEFAULT_SOURCING_CELL_BATCH_SIZE
        runnable_cells = all_runnable_cells[:cell_batch_size]
        executed_cell_ids = {
            str(cell.get("cell_id") or "") for cell in runnable_cells if isinstance(cell, dict)
        }
        execution_plan = {**query_plan, "cells": runnable_cells, "cell_count": len(runnable_cells)}
        liepin_queries = query_builders.query_plan_channel_entries(execution_plan, "liepin")
        xsaas_queries = query_builders.query_plan_channel_entries(execution_plan, "xsaas")
        company_terms = query_builders.query_plan_company_vocabulary(query_plan)
        liepin_queries_path.write_text(json.dumps({"queries": liepin_queries}, ensure_ascii=False, indent=2), encoding="utf-8")
        xsaas_queries_path.write_text(json.dumps({"queries": xsaas_queries}, ensure_ascii=False, indent=2), encoding="utf-8")
        # 并行跑两条渠道（OC1→production fallback），取代串行等待
        from concurrent.futures import ThreadPoolExecutor, as_completed
        _oc1 = self._opencli_primary_enabled(request)
        _cdp = int(request.get("cdp_port") or 9223)
        _lim = max(12, target * 2)
        _det, _detail_args = self._liepin_detail_capture_options(request, target)
        _report: dict[str, Any] = {}

        def _run_liepin() -> tuple[str, dict[str, Any] | None]:
            if not liepin_queries:
                liepin_path.write_text("[]", encoding="utf-8")
                liepin_raw_path.write_text("[]", encoding="utf-8")
                return "resume_skipped", {"ok": True, "status": "resume_skipped", "rounds": []}
            eng = "production_fallback" if _oc1 else "production"
            res: dict[str, Any] | None = None
            if _oc1:
                eng = self._attempt_opencli_primary(
                    channel="liepin", client=client, job=job, port=_cdp,
                    queries_path=liepin_queries_path, output_path=liepin_path,
                    raw_output_path=liepin_raw_path,
                    limit=min(_lim, 24), detail_limit=_det, report=_report,
                    detail_args=_detail_args,
                    cancel_check=cancel_check)
                if eng == "opencli":
                    res = {**_report.get("liepin", {}), "ok": True, "recall_engine": "opencli"}
                elif eng == "opencli_partial":
                    primary_summary = _report.get("liepin", {})
                    primary_rows = _loads(liepin_path.read_text(encoding="utf-8"), [])
                    primary_raw = _loads(liepin_raw_path.read_text(encoding="utf-8"), [])
                    fallback_entries = self._opencli_fallback_entries(
                        liepin_queries, primary_summary, primary_raw,
                    )
                    fallback_queries_path = liepin_queries_path.with_name(
                        liepin_queries_path.stem + "-paginated.json"
                    )
                    fallback_path = liepin_path.with_name(liepin_path.stem + "-paginated.json")
                    fallback_raw_path = liepin_raw_path.with_name(
                        liepin_raw_path.stem + "-paginated.json"
                    )
                    fallback_queries_path.write_text(
                        json.dumps({"queries": fallback_entries}, ensure_ascii=False, indent=2),
                        encoding="utf-8",
                    )
                    try:
                        fallback_result = self._run_external_json([
                            self.python, str(LIEPIN_SEARCH), "--client", client, "--position", job,
                            "--db", str(self.service.db_path), "--output-dir", str(candidates_path.parent),
                            "--port", str(_cdp), "--rounds", str(len(fallback_entries)),
                            "--max-cards", str(min(_lim, 24)), "--min-score", str(quality_min_score), "--recommend-score", str(quality_recommend_score),
                            "--max-pages", str(page_budget),
                            "--capture-links", "--capture-details", "--detail-limit", str(_det),
                            *_detail_args,
                            "--no-open-links", "--dry-run", "--json-output", str(fallback_path),
                            "--raw-json-output", str(fallback_raw_path),
                            "--queries-json", str(fallback_queries_path),
                        ], 900, cancel_check=cancel_check)
                    except ExternalExecutionCancelled:
                        raise
                    except Exception as exc:
                        fallback_result = {
                            "ok": False, "status": "blocked", "error": _trim_error(exc), "rounds": [],
                        }
                        fallback_path.write_text("[]", encoding="utf-8")
                        fallback_raw_path.write_text("[]", encoding="utf-8")
                    res, merged_rows, merged_raw = self._merge_opencli_completion(
                        channel="liepin",
                        primary_summary=primary_summary,
                        fallback_result=fallback_result,
                        primary_rows=primary_rows,
                        fallback_rows=_loads(fallback_path.read_text(encoding="utf-8"), []),
                        primary_raw=primary_raw,
                        fallback_raw=_loads(fallback_raw_path.read_text(encoding="utf-8"), []),
                    )
                    liepin_path.write_text(json.dumps(merged_rows, ensure_ascii=False, indent=2), encoding="utf-8")
                    liepin_raw_path.write_text(json.dumps(merged_raw, ensure_ascii=False, indent=2), encoding="utf-8")
                    eng = "opencli_paginated"
            if res is None:
                try:
                    res = self._run_external_json([
                        self.python, str(LIEPIN_SEARCH), "--client", client, "--position", job,
                        "--db", str(self.service.db_path), "--output-dir", str(candidates_path.parent),
                        "--port", str(_cdp), "--rounds", str(len(liepin_queries)),
                        "--max-cards", str(min(_lim, 24)), "--min-score", str(quality_min_score), "--recommend-score", str(quality_recommend_score),
                        "--max-pages", str(page_budget),
                        "--capture-links", "--capture-details", "--detail-limit", str(_det),
                        *_detail_args,
                        "--no-open-links", "--dry-run", "--json-output", str(liepin_path),
                        "--raw-json-output", str(liepin_raw_path),
                        "--queries-json", str(liepin_queries_path),
                    ], 900, cancel_check=cancel_check)
                except ExternalExecutionCancelled:
                    raise
                except Exception as exc:
                    res = {"ok": False, "status": "blocked", "error": _trim_error(exc), "rounds": []}
                    liepin_path.write_text("[]", encoding="utf-8")
                    liepin_raw_path.write_text("[]", encoding="utf-8")
            return eng, res

        def _run_xsaas() -> tuple[str, dict[str, Any] | None]:
            if not xsaas_queries:
                xsaas_path.write_text("[]", encoding="utf-8")
                xsaas_raw_path.write_text("[]", encoding="utf-8")
                return "resume_skipped", {"ok": True, "status": "resume_skipped", "rounds": []}
            eng = "production_fallback" if _oc1 else "production"
            res: dict[str, Any] | None = None
            if _oc1:
                eng = self._attempt_opencli_primary(
                    channel="xsaas", client=client, job=job, port=_cdp,
                    queries_path=xsaas_queries_path, output_path=xsaas_path,
                    raw_output_path=xsaas_raw_path,
                    limit=min(_lim, 100), detail_limit=_det, report=_report,
                    cancel_check=cancel_check)
                if eng == "opencli":
                    res = {**_report.get("xsaas", {}), "ok": True, "recall_engine": "opencli"}
                elif eng == "opencli_partial":
                    primary_summary = _report.get("xsaas", {})
                    primary_rows = _loads(xsaas_path.read_text(encoding="utf-8"), [])
                    primary_raw = _loads(xsaas_raw_path.read_text(encoding="utf-8"), [])
                    fallback_entries = self._opencli_fallback_entries(
                        xsaas_queries, primary_summary, primary_raw,
                    )
                    fallback_queries_path = xsaas_queries_path.with_name(
                        xsaas_queries_path.stem + "-paginated.json"
                    )
                    fallback_path = xsaas_path.with_name(xsaas_path.stem + "-paginated.json")
                    fallback_raw_path = xsaas_raw_path.with_name(
                        xsaas_raw_path.stem + "-paginated.json"
                    )
                    fallback_queries_path.write_text(
                        json.dumps({"queries": fallback_entries}, ensure_ascii=False, indent=2),
                        encoding="utf-8",
                    )
                    try:
                        fallback_result = self._run_external_json([
                            self.python, str(XSAAS_SEARCH), "--queries", str(fallback_queries_path),
                            "--output", str(fallback_path), "--port", str(_cdp),
                            "--raw-output", str(fallback_raw_path),
                            "--max-rows", str(min(_lim, 100)), "--db", str(self.service.db_path),
                            "--max-pages", str(page_budget),
                            "--client", client, "--job", job, "--min-score", str(quality_min_score),
                        ], 300, cancel_check=cancel_check)
                    except ExternalExecutionCancelled:
                        raise
                    except Exception as exc:
                        fallback_result = {"ok": False, "status": "blocked", "error": _trim_error(exc), "rounds": []}
                        fallback_path.write_text("[]", encoding="utf-8")
                        fallback_raw_path.write_text("[]", encoding="utf-8")
                    res, merged_rows, merged_raw = self._merge_opencli_completion(
                        channel="xsaas",
                        primary_summary=primary_summary,
                        fallback_result=fallback_result,
                        primary_rows=primary_rows,
                        fallback_rows=_loads(fallback_path.read_text(encoding="utf-8"), []),
                        primary_raw=primary_raw,
                        fallback_raw=_loads(fallback_raw_path.read_text(encoding="utf-8"), []),
                    )
                    xsaas_path.write_text(json.dumps(merged_rows, ensure_ascii=False, indent=2), encoding="utf-8")
                    xsaas_raw_path.write_text(json.dumps(merged_raw, ensure_ascii=False, indent=2), encoding="utf-8")
                    eng = "opencli_paginated"
            if res is None:
                try:
                    res = self._run_external_json([
                        self.python, str(XSAAS_SEARCH), "--queries", str(xsaas_queries_path),
                        "--output", str(xsaas_path), "--port", str(_cdp),
                        "--raw-output", str(xsaas_raw_path),
                        "--max-rows", str(min(_lim, 100)), "--db", str(self.service.db_path),
                        "--max-pages", str(page_budget),
                        "--client", client, "--job", job, "--min-score", str(quality_min_score),
                    ], 300, cancel_check=cancel_check)
                except ExternalExecutionCancelled:
                    raise
                except Exception:
                    res = {"ok": False, "status": "blocked", "error": _trim_error(sys.exc_info()[1])}
                    xsaas_path.write_text("[]", encoding="utf-8")
            return eng, res

        def _run_xsaas_guarded() -> tuple[str, dict[str, Any] | None]:
            # X-SaaS runners use isolated tabs marked with asa_search_runner=1.
            # A subprocess timeout cannot execute its own finally block, so the
            # parent removes only those owned tabs both before and after the run.
            from xsaas_candidate_search import close_runner_tabs

            try:
                close_runner_tabs(_cdp)
            except Exception:
                pass
            try:
                return _run_xsaas()
            finally:
                try:
                    close_runner_tabs(_cdp)
                except Exception:
                    pass

        with ThreadPoolExecutor(max_workers=2) as _pool:
            _futures = {
                _pool.submit(_run_liepin): "liepin",
                _pool.submit(_run_xsaas_guarded): "xsaas",
            }
            _results: dict[str, tuple[str, dict[str, Any] | None]] = {}
            for _fut in as_completed(_futures):
                _results[_futures[_fut]] = _fut.result()

        liepin_engine, search = _results["liepin"]
        xsaas_engine, xsaas = _results["xsaas"]
        self._ensure_external_request_active(request)
        primary_channels = _report
        risk_stop_reason = self._channel_risk_stop_reason(search)
        opencli_shadow = (
            {
                "enabled": False,
                "mode": "read_only_shadow",
                "affects_intake": False,
                "reason": "channel_risk_hard_stop",
            }
            if risk_stop_reason
            else self._run_opencli_shadow(
                request=request,
                client=client,
                job=job,
                port=int(request.get("cdp_port") or 9223),
                limit=max(12, target * 2),
                liepin_queries=liepin_queries,
                xsaas_queries=xsaas_queries,
                liepin_path=liepin_path,
                xsaas_path=xsaas_path,
                artifact_path=candidates_path.with_name(candidates_path.stem + "-opencli-shadow.json"),
                liepin_queries_path=liepin_queries_path,
                xsaas_queries_path=xsaas_queries_path,
                skip_channels={
                    channel
                    for channel, engine in (("liepin", liepin_engine), ("xsaas", xsaas_engine))
                    if engine.startswith("opencli")
                },
                cancel_check=cancel_check,
            )
        )
        liepin_candidates = _loads(liepin_path.read_text(encoding="utf-8"), [])
        xsaas_candidates = _loads(xsaas_path.read_text(encoding="utf-8"), [])
        if not liepin_raw_path.exists():
            liepin_raw_path.write_text(json.dumps(liepin_candidates, ensure_ascii=False, indent=2), encoding="utf-8")
        if not xsaas_raw_path.exists():
            xsaas_raw_path.write_text(json.dumps(xsaas_candidates, ensure_ascii=False, indent=2), encoding="utf-8")
        liepin_raw_candidates = _loads(liepin_raw_path.read_text(encoding="utf-8"), [])
        xsaas_raw_candidates = _loads(xsaas_raw_path.read_text(encoding="utf-8"), [])
        combined = liepin_candidates + xsaas_candidates
        quality_rejected: list[dict[str, Any]] = []
        if quality_min_score >= 70:
            def candidate_score(item: dict[str, Any]) -> int:
                try:
                    return int(float(item.get("fit_score") or 0))
                except (TypeError, ValueError):
                    return 0

            quality_rejected = [
                item for item in combined
                if candidate_score(item) < quality_min_score
            ]
            combined = [
                item for item in combined
                if candidate_score(item) >= quality_min_score
            ]
        candidates_path.write_text(json.dumps(combined, ensure_ascii=False, indent=2), encoding="utf-8")
        workflow_id = str(request.get("workflow_id") or "")

        def _normalize_run_result(raw: dict[str, Any] | None) -> dict[str, Any]:
            """把 runner 返回的 result.candidates 统一为数量，避免前端把列表当 0 处理。"""
            if not isinstance(raw, dict):
                return raw or {}
            normalized = dict(raw)
            candidates = normalized.get("candidates")
            if isinstance(candidates, list):
                normalized["candidates"] = len(candidates)
            return normalized

        channel_runs = [
            {"channel": "liepin", "status": "risk_stopped" if risk_stop_reason else "completed" if search.get("ok") else "blocked", "recall_engine": liepin_engine, "result": _normalize_run_result(search)},
            {"channel": "xsaas", "status": "completed" if xsaas.get("ok") else "blocked", "recall_engine": xsaas_engine, "result": _normalize_run_result(xsaas)},
        ]

        # Persist raw channel evidence before any formal candidate intake. The later
        # upserts only enrich these immutable occurrences with disposition receipts.
        raw_candidates = {"liepin": liepin_raw_candidates, "xsaas": xsaas_raw_candidates}
        recall_ledger = self._persist_candidate_recalls(
            run_id=run_id,
            workflow_id=workflow_id,
            client=client,
            job=job,
            query_plan=query_plan,
            raw_candidates=raw_candidates,
            applied={},
            min_score=quality_min_score,
        )
        query_cell_states = self._persist_query_cell_states(
            run_id=run_id,
            workflow_id=workflow_id,
            client=client,
            job=job,
            query_plan=query_plan,
            channel_runs=channel_runs,
            executed_cell_ids=executed_cell_ids,
        )
        coverage_certificate = self._build_coverage_certificate(
            run_id=run_id,
            workflow_id=workflow_id,
            client=client,
            job=job,
            query_plan=query_plan,
        )

        self._ensure_external_request_active(request)
        dry = self._run_external_json([self.python, str(MULTICHANNEL), "intake", "--db", str(self.service.db_path), "--client", client, "--job", job, "--input", str(candidates_path)], 120, cancel_check=cancel_check)
        recall_ledger = self._persist_candidate_recalls(
            run_id=run_id,
            workflow_id=workflow_id,
            client=client,
            job=job,
            query_plan=query_plan,
            raw_candidates=raw_candidates,
            applied=dry,
            min_score=quality_min_score,
        )
        self._ensure_external_request_active(request)
        applied = self._run_external_json([self.python, str(MULTICHANNEL), "intake", "--db", str(self.service.db_path), "--client", client, "--job", job, "--input", str(candidates_path), "--apply"], 180, cancel_check=cancel_check)
        attributions = self._persist_sourcing_attributions(
            applied, request.get("strategy") if isinstance(request.get("strategy"), dict) else {},
            workflow_id, client, job,
        )
        # 入池后补抓：对已入池但还没有完整猎聘履历的人选自动补抓一轮，写入 source_profiles，
        # 使人选详情页与后续 candidate_batch_assessment 能直接拿到完整履历。失败不阻断寻访。
        resume_backfill: dict[str, Any] = {"status": "skipped", "reason": ""}
        if risk_stop_reason:
            resume_backfill["reason"] = "channel_risk_stop"
        elif str(os.environ.get("ASA_RESUME_BACKFILL", "on")).strip().lower() in {"0", "off", "false"}:
            resume_backfill["reason"] = "disabled"
        else:
            self._ensure_external_request_active(request)
            _, detail_flags = self._liepin_detail_capture_options(request, target)
            try:
                backfill_limit = int(os.environ.get("ASA_RESUME_BACKFILL_LIMIT", "40") or 40)
            except (TypeError, ValueError):
                backfill_limit = 40
            backfill_limit = max(1, min(backfill_limit, 100))
            try:
                resume_backfill = self._run_external_json(
                    [
                        self.python, str(RESUME_BACKFILL),
                        "--db", str(self.service.db_path),
                        "--client", client, "--job", job,
                        "--port", str(int(request.get("cdp_port") or 9223)),
                        "--limit", str(backfill_limit),
                        *detail_flags,
                    ],
                    120 + 25 * backfill_limit,
                    cancel_check=cancel_check,
                )
            except ExternalExecutionCancelled:
                raise
            except Exception as exc:
                resume_backfill = {"status": "failed", "error": str(exc)[:300]}
        recall_ledger = self._persist_candidate_recalls(
            run_id=run_id,
            workflow_id=workflow_id,
            client=client,
            job=job,
            query_plan=query_plan,
            raw_candidates=raw_candidates,
            applied=applied,
            min_score=quality_min_score,
        )
        coverage_certificate = self._build_coverage_certificate(
            run_id=run_id,
            workflow_id=workflow_id,
            client=client,
            job=job,
            query_plan=query_plan,
        )
        try:
            funnel = self._persist_sourcing_funnel(
                run_id=run_id,
                workflow_id=workflow_id,
                client=client,
                job=job,
                channel_runs=channel_runs,
                channel_candidates={"liepin": liepin_candidates, "xsaas": xsaas_candidates},
                applied=applied,
                attributions=attributions,
                company_vocab=company_terms,
            )
        except Exception as exc:
            funnel = {"ok": False, "stored": 0, "error": str(exc)[:500]}
        partial_result = {
            "verified": True,
            "run_id": run_id,
            "channel_runs": channel_runs,
            "opencli_shadow": opencli_shadow,
            "opencli_primary": {"enabled": _oc1, "channels": primary_channels},
            "intake": {"dry_run": dry, "applied": applied, "source_file": str(candidates_path)},
            "attributions": attributions,
            "resume_backfill": resume_backfill,
            "candidate_recall_ledger": recall_ledger,
            "query_cell_states": query_cell_states,
            "coverage_certificate": coverage_certificate,
            "sourcing_funnel": funnel,
            "quality_gate": {
                "minimum_score": quality_min_score,
                "rejected_before_intake": len(quality_rejected),
                "policy": "specialized_power_minimum_score" if quality_min_score >= 70 else "channel_default",
            },
            "channel_risk_stop": {
                "active": bool(risk_stop_reason),
                "channel": "liepin" if risk_stop_reason else "",
                "signal": risk_stop_reason,
                "message": "猎聘命中安全风险提示，已停止猎聘及后续分页。" if risk_stop_reason else "",
            },
            "audit": {"ok": False, "summary": "等待 A 系统收尾审计"},
        }
        sync_script = Path("/Users/messi/.codex/skills/a-system-workbench/scripts/a_system_sync.py")
        try:
            sync = self._run_external([self.python, str(sync_script), "--client", client, "--job", job, "--no-open"], 300, cancel_check=cancel_check)
        except CommandExecutionError as exc:
            partial_result["audit"] = {
                "ok": False,
                "phase": "audit",
                "summary": str(exc),
                "detail": exc.detail,
            }
            raise ExternalPhaseError(
                f"寻访与入库已完成，但 A 系统收尾审计未通过：{exc}",
                phase="audit",
                partial_result=partial_result,
                detail=exc.detail,
            ) from exc
        learning = self._capture_search_learning(client, job, [*liepin_queries, *xsaas_queries])
        continuation = (
            {
                "summary": {
                    "scheduled": False,
                    "reason": "channel_risk_hard_stop",
                    "risk_signal": risk_stop_reason,
                    "remaining_cells": max(0, len(all_runnable_cells) - len(runnable_cells)),
                },
                "request": None,
            }
            if risk_stop_reason
            else self._sourcing_continuation(request=request, run_id=run_id, query_plan=query_plan)
        )
        final_result = {
            **partial_result,
            "audit": {
                "ok": True,
                "summary": "A 系统收尾审计通过",
                "returncode": sync.returncode,
            },
            "learning": learning,
            "continuation": continuation["summary"],
        }
        if continuation["request"] is not None:
            final_result["_continuation_request"] = continuation["request"]
        return final_result

    @staticmethod
    def _query_text(entries: list[Any]) -> str:
        if not entries:
            return ""
        first = entries[0]
        value = first.get("query") if isinstance(first, dict) else first
        return " ".join(str(value or "").split())

    @staticmethod
    def _query_entry_text(entry: Any) -> str:
        value = entry.get("query") if isinstance(entry, dict) else entry
        return " ".join(str(value or "").split())

    @classmethod
    def _opencli_fallback_entries(
        cls,
        entries: list[Any],
        summary: dict[str, Any],
        primary_rows: list[Any] | None = None,
    ) -> list[Any]:
        """Resume only OpenCLI query cells that did not prove exhaustion."""
        rounds = {
            cls._query_entry_text(item.get("query")): item
            for item in summary.get("rounds") or []
            if isinstance(item, dict) and cls._query_entry_text(item.get("query"))
        }
        fallback: list[Any] = []
        for entry in entries:
            query = cls._query_entry_text(entry)
            if not query:
                continue
            round_item = rounds.get(query)
            if round_item and round_item.get("terminal_state") == "exhausted":
                continue
            base = dict(entry) if isinstance(entry, dict) else {"query": query}
            extracted = _round_int(round_item, "extracted_count")
            cursor = round_item.get("cursor") if isinstance(round_item, dict) else None
            if (
                round_item
                and round_item.get("terminal_state") == "platform_capped"
                and extracted > 0
                and isinstance(cursor, dict)
                and int(cursor.get("page") or 0) > 1
            ):
                base["cursor"] = {"page": int(cursor["page"])}
                base["collected_before"] = extracted
                seen_keys = [
                    key
                    for item in primary_rows or []
                    if isinstance(item, dict)
                    and cls._query_entry_text(item.get("source_query") or item.get("query")) == query
                    and (key := cls._candidate_resume_key(item))
                ]
                if seen_keys:
                    base["seen_candidate_keys"] = list(dict.fromkeys(seen_keys))
            else:
                base.pop("cursor", None)
                base.pop("collected_before", None)
                base.pop("seen_candidate_keys", None)
            fallback.append(base)
        return fallback

    @staticmethod
    def _candidate_resume_key(item: dict[str, Any]) -> str:
        source_id = str(
            item.get("candidate_id") or item.get("resume_id") or item.get("res_id_encode")
            or item.get("xsaas_id") or ""
        ).strip()
        if source_id:
            return source_id
        return "|".join(
            " ".join(str(item.get(key) or "").split()).casefold()
            for key in ("name", "company", "title")
        )

    @staticmethod
    def _candidate_artifact_key(channel: str, item: dict[str, Any]) -> str:
        source_id = str(
            item.get("candidate_id") or item.get("resume_id") or item.get("res_id_encode")
            or item.get("xsaas_id") or item.get("resume_url") or item.get("source_url") or ""
        ).strip()
        if source_id:
            return f"{channel}:id:{source_id}"
        identity = "|".join(
            " ".join(str(item.get(key) or "").split()).casefold()
            for key in ("name", "company", "current_company", "title", "current_title")
        )
        return f"{channel}:identity:{identity}"

    @classmethod
    def _merge_opencli_completion(
        cls,
        *,
        channel: str,
        primary_summary: dict[str, Any],
        fallback_result: dict[str, Any],
        primary_rows: list[Any],
        fallback_rows: list[Any],
        primary_raw: list[Any],
        fallback_raw: list[Any],
    ) -> tuple[dict[str, Any], list[dict[str, Any]], list[dict[str, Any]]]:
        """Merge OpenCLI page 1 with paginated completion without losing occurrence evidence."""
        accepted_seen: set[str] = set()
        accepted: list[dict[str, Any]] = []
        for raw in [*primary_rows, *fallback_rows]:
            if not isinstance(raw, dict):
                continue
            key = cls._candidate_artifact_key(channel, raw)
            if key in accepted_seen:
                continue
            accepted_seen.add(key)
            accepted.append(raw)
        raw_rows = [item for item in [*primary_raw, *fallback_raw] if isinstance(item, dict)]

        fallback_rounds = {
            cls._query_entry_text(item.get("query")): item
            for item in fallback_result.get("rounds") or []
            if isinstance(item, dict) and cls._query_entry_text(item.get("query"))
        }
        merged_rounds: list[dict[str, Any]] = []
        merged_queries: set[str] = set()
        for raw_round in primary_summary.get("rounds") or []:
            if not isinstance(raw_round, dict):
                continue
            primary_round = dict(raw_round)
            query = cls._query_entry_text(primary_round.get("query"))
            merged_queries.add(query)
            fallback_round = fallback_rounds.get(query)
            if primary_round.get("terminal_state") == "exhausted" or not fallback_round:
                merged_rounds.append(primary_round)
                continue
            merged = dict(fallback_round)
            if (
                primary_round.get("terminal_state") == "platform_capped"
                and _round_int(primary_round, "extracted_count") > 0
            ):
                query_raw = [
                    item for item in raw_rows
                    if cls._query_entry_text(item.get("source_query") or item.get("query")) == query
                ]
                merged["extracted_count"] = len(query_raw)
                merged["unique_count"] = len({cls._candidate_artifact_key(channel, item) for item in query_raw})
                merged["pages_fetched"] = (
                    _round_int(primary_round, "pages_fetched")
                    + _round_int(fallback_round, "pages_fetched")
                )
                if fallback_round.get("result_count") is None:
                    merged["result_count"] = primary_round.get("result_count")
                merged["resumed_after_opencli"] = True
            merged_rounds.append(merged)
        merged_rounds.extend(
            dict(item)
            for query, item in fallback_rounds.items()
            if query not in merged_queries
        )
        result = {
            **fallback_result,
            "mode": "opencli_primary_with_paginated_completion",
            "opencli_primary": primary_summary,
            "rounds": merged_rounds,
            "candidates": len(accepted),
            "ok": bool(fallback_result.get("ok")),
        }
        return result, accepted, raw_rows

    @staticmethod
    def _opencli_primary_enabled(request: dict[str, Any]) -> bool:
        """OpenCLI 默认主召回；请求级参数或环境变量可显式关闭并回退生产 runner。"""
        configured = request.get("opencli_primary", os.environ.get("ASA_OPENCLI_PRIMARY", "1"))
        return str(configured).strip().lower() in {"1", "true", "yes", "on"}

    def _attempt_opencli_primary(
        self,
        *,
        channel: str,
        client: str,
        job: str,
        port: int,
        queries_path: Path,
        output_path: Path,
        raw_output_path: Path,
        limit: int,
        detail_limit: int,
        report: dict[str, Any],
        detail_args: list[str] | None = None,
        cancel_check: Callable[[], bool] | None = None,
    ) -> str:
        """OpenCLI 主渠道召回；失败、被阻断或无完整合格行时回退生产 runner。"""
        query_payload = _loads(queries_path.read_text(encoding="utf-8"), {})
        planned_queries = query_payload.get("queries") if isinstance(query_payload, dict) else []
        if any(
            isinstance(item, dict)
            and isinstance(item.get("cursor"), dict)
            and int(item["cursor"].get("page") or 0) > 1
            for item in (planned_queries if isinstance(planned_queries, list) else [])
        ):
            report[channel] = {
                "ok": False,
                "mode": "opencli_primary_recall",
                "channel": channel,
                "status": "production_fallback",
                "reason": "resume_cursor_requires_paginated_runner",
            }
            return "production_fallback"
        try:
            summary = self._run_external_json(
                [
                    self.python, str(OPENCLI_SHADOW), "--mode", "primary",
                    "--channel", channel, "--queries-json", str(queries_path),
                    "--output", str(output_path),
                    "--raw-output", str(raw_output_path),
                    "--client", client, "--job", job,
                    "--db", str(self.service.db_path), "--port", str(port),
                    "--limit", str(limit), "--detail-limit", str(detail_limit),
                    *(detail_args or []),
                    "--max-queries", str(max(1, len(planned_queries) if isinstance(planned_queries, list) else 0)),
                ],
                600,
                cancel_check=cancel_check,
            )
        except ExternalExecutionCancelled:
            raise
        except Exception as exc:
            summary = {
                "ok": False, "mode": "opencli_primary_recall", "channel": channel,
                "error": _trim_error(exc),
            }
        report[channel] = summary
        if summary.get("coverage_complete") or summary.get("ok"):
            return "opencli"
        if any(
            isinstance(item, dict)
            and item.get("terminal_state") in {"exhausted", "platform_capped"}
            for item in summary.get("rounds") or []
        ):
            return "opencli_partial"
        return "production_fallback"

    def _run_opencli_shadow(
        self,
        *,
        request: dict[str, Any],
        client: str,
        job: str,
        port: int,
        limit: int,
        liepin_queries: list[Any],
        xsaas_queries: list[Any],
        liepin_path: Path,
        xsaas_path: Path,
        artifact_path: Path,
        liepin_queries_path: Path | None = None,
        xsaas_queries_path: Path | None = None,
        skip_channels: set[str] | None = None,
        cancel_check: Callable[[], bool] | None = None,
    ) -> dict[str, Any]:
        configured = request.get("opencli_shadow", os.environ.get("ASA_OPENCLI_SHADOW", "1"))
        enabled = str(configured).strip().lower() not in {"0", "false", "no", "off"}
        if not enabled:
            return {"enabled": False, "mode": "read_only_shadow", "affects_intake": False, "channels": []}
        channels = []
        for channel, entries, baseline, queries_file in (
            ("liepin", liepin_queries, liepin_path, liepin_queries_path),
            ("xsaas", xsaas_queries, xsaas_path, xsaas_queries_path),
        ):
            if skip_channels and channel in skip_channels:
                channels.append({
                    "channel": channel, "status": "skipped",
                    "reason": "recall_engine_opencli", "affects_intake": False,
                })
                continue
            query = self._query_text(entries)
            if not query:
                channels.append({"channel": channel, "status": "skipped", "reason": "missing_query"})
                continue
            channel_limit = min(limit, 24 if channel == "liepin" else 100)
            command = [
                self.python,
                str(OPENCLI_SHADOW),
                "--channel", channel,
                "--query", query,
                "--baseline", str(baseline),
                "--client", client,
                "--job", job,
                "--db", str(self.service.db_path),
                "--port", str(port),
                "--limit", str(channel_limit),
            ]
            if queries_file is not None:
                command += ["--queries-json", str(queries_file)]
            try:
                result = self._run_external_json(command, 600, cancel_check=cancel_check)
                channels.append({"channel": channel, "status": "completed", **result})
            except ExternalExecutionCancelled:
                raise
            except Exception as exc:
                channels.append({
                    "channel": channel,
                    "status": "blocked",
                    "query": query,
                    "error": _trim_error(exc),
                    "affects_intake": False,
                })
        payload = {
            "enabled": True,
            "mode": "read_only_shadow",
            "affects_intake": False,
            "affects_outreach": False,
            "sample_policy": "first_nonempty_baseline_else_first",
            "channels": channels,
            "artifact": str(artifact_path),
        }
        artifact_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        history_path = artifact_path.parent / "opencli-shadow-history.jsonl"
        with history_path.open("a", encoding="utf-8") as handle:
            handle.write(json.dumps({
                "recorded_at": datetime.now().isoformat(timespec="seconds"),
                "client": client,
                "job": job,
                "workflow_id": str(request.get("workflow_id") or ""),
                "channels": channels,
            }, ensure_ascii=False) + "\n")
        payload["history"] = str(history_path)
        return payload

    def _resume_query_cells(
        self,
        run_id: str,
        query_plan: dict[str, Any],
        *,
        max_retries: int = 3,
        max_platform_capped_retries: int = 0,
    ) -> list[dict[str, Any]]:
        """Select only unfinished/retryable cells for an explicitly resumed run."""
        conn = self.service._connect()
        try:
            rows = conn.execute(
                "SELECT cell_id,plan_hash,status,retry_count,cursor_json,pages_fetched,terminal_reason,"
                "extracted_count,unique_count,updated_at FROM agent_sourcing_query_cells WHERE run_id=?",
                (run_id,),
            ).fetchall()
            recall_rows = conn.execute(
                "SELECT query_cell_id,source_candidate_id FROM agent_candidate_recalls "
                "WHERE run_id=? AND query_cell_id<>'' AND source_candidate_id<>''",
                (run_id,),
            ).fetchall()
        finally:
            conn.close()
        if not rows:
            return [cell for cell in query_plan.get("cells") or [] if isinstance(cell, dict)]
        plan_hash = str(query_plan.get("plan_hash") or "")
        if any(str(row["plan_hash"] or "") != plan_hash for row in rows):
            raise ValueError("断点续跑的 query_plan_v1 与原 run_id 不一致")
        states = {
            str(row["cell_id"]): (
                str(row["status"]), int(row["retry_count"] or 0), _loads(row["cursor_json"], {}),
                int(row["pages_fetched"] or 0), int(row["extracted_count"] or 0),
                int(row["unique_count"] or 0), str(row["terminal_reason"] or ""),
                str(row["updated_at"] or ""),
            )
            for row in rows
        }
        seen_keys_by_cell: dict[str, list[str]] = {}
        for row in recall_rows:
            cell_id = str(row["query_cell_id"] or "")
            source_id = str(row["source_candidate_id"] or "").strip()
            if cell_id and source_id and not source_id.startswith("anon_"):
                seen_keys_by_cell.setdefault(cell_id, []).append(source_id)
        blocked_families_by_channel: dict[str, set[str]] = {"liepin": set(), "xsaas": set()}
        for cell in query_plan.get("cells") or []:
            if not isinstance(cell, dict):
                continue
            cell_id = str(cell.get("cell_id") or "")
            status, _retries, cursor, _pages, _extracted, _unique, terminal_reason, _updated = states.get(
                cell_id, ("pending", 0, {}, 0, 0, 0, "", ""),
            )
            if status not in {"failed", "blocked"}:
                continue
            if status == "blocked" and cursor:
                continue
            if status == "blocked" and terminal_reason not in {
                "channel_blocked_before_query", "approved_cell_not_executed",
            }:
                continue
            channel = str(cell.get("channel") or "")
            blocked_families_by_channel.setdefault(channel, set()).update(
                str(value) for value in cell.get("query_family_ids") or [] if str(value).strip()
            )
        pending: list[dict[str, Any]] = []
        fallback_pending: list[dict[str, Any]] = []
        retryable: list[tuple[str, int, dict[str, Any]]] = []
        for plan_index, cell in enumerate(query_plan.get("cells") or []):
            if not isinstance(cell, dict):
                continue
            status, retries, cursor, pages_fetched, extracted_count, unique_count, terminal_reason, updated_at = states.get(
                str(cell.get("cell_id") or ""), ("pending", 0, {}, 0, 0, 0, "", ""),
            )
            if status == "pending":
                channel = str(cell.get("channel") or "")
                other_channel = "xsaas" if channel == "liepin" else "liepin"
                families = {str(value) for value in cell.get("query_family_ids") or [] if str(value).strip()}
                if families and families & blocked_families_by_channel.get(other_channel, set()):
                    fallback_pending.append({
                        **cell,
                        "execution_fallback_relay": {
                            "from_channel": other_channel,
                            "reason": "same_query_family_blocked",
                            "query_family_ids": sorted(families & blocked_families_by_channel[other_channel]),
                        },
                    })
                else:
                    pending.append(cell)
            elif (
                (status == "failed" and retries < max(1, max_retries))
                or (
                    status == "blocked"
                    and retries < max(1, max_retries)
                    and not cursor
                    and terminal_reason in {"channel_blocked_before_query", "approved_cell_not_executed"}
                )
            ):
                retryable.append((updated_at, plan_index, cell))
            elif (
                status in {"platform_capped", "blocked"}
                and (
                    (status == "platform_capped" and retries < max(0, max_platform_capped_retries))
                    or (status == "blocked" and retries < max(1, max_retries))
                )
                and isinstance(cursor, dict)
                and int(cursor.get("page") or 0) > 1
            ):
                retryable.append((
                    updated_at,
                    plan_index,
                    {
                        **cell,
                        "execution_cursor": {"page": int(cursor["page"])},
                        "execution_progress": {
                            "pages_fetched": pages_fetched,
                            "extracted_count": extracted_count,
                            "unique_count": unique_count,
                            "seen_candidate_keys": list(dict.fromkeys(seen_keys_by_cell.get(str(cell.get("cell_id") or ""), []))),
                        },
                    },
                ))
        retryable.sort(key=lambda item: (item[0], item[1]))
        return [*fallback_pending, *pending, *(item[2] for item in retryable)]

    def _sourcing_continuation(
        self,
        *,
        request: dict[str, Any],
        run_id: str,
        query_plan: dict[str, Any],
    ) -> dict[str, Any]:
        """Create a hash-bound next batch for retryable/cursor-bearing query cells."""
        try:
            index = max(0, int(request.get("_continuation_index") or 0))
        except (TypeError, ValueError):
            index = 0
        try:
            cell_batch_size = max(
                1,
                min(
                    int(request.get("max_query_cells_per_batch") or DEFAULT_SOURCING_CELL_BATCH_SIZE),
                    MAX_SOURCING_CELL_BATCH_SIZE,
                ),
            )
        except (TypeError, ValueError):
            cell_batch_size = DEFAULT_SOURCING_CELL_BATCH_SIZE
        approved_cell_count = sum(
            1 for cell in query_plan.get("cells") or [] if isinstance(cell, dict)
        )
        minimum_plan_continuations = max(
            0,
            (approved_cell_count + cell_batch_size - 1) // cell_batch_size - 1,
        )
        default_budget = min(
            MAX_SOURCING_CONTINUATION_BATCHES,
            minimum_plan_continuations + DEFAULT_PAGINATION_CONTINUATION_HEADROOM,
        )
        raw_budget = request.get("max_continuation_batches")
        if raw_budget in (None, ""):
            max_batches = default_budget
        else:
            try:
                requested_budget = max(
                    0,
                    min(int(raw_budget), MAX_SOURCING_CONTINUATION_BATCHES),
                )
            except (TypeError, ValueError):
                requested_budget = default_budget
            # A legacy/manual cap must not prevent every approved cell from receiving
            # its first execution attempt. Pagination still uses the bounded remainder.
            max_batches = max(minimum_plan_continuations, requested_budget)
        runnable = self._resume_query_cells(
            run_id,
            query_plan,
            max_retries=int(request.get("max_query_retries") or 3),
            max_platform_capped_retries=self._platform_capped_continuation_limit(request),
        )
        summary = {
            "scheduled": bool(runnable and index < max_batches),
            "run_id": run_id,
            "completed_batches": index + 1,
            "remaining_cells": len(runnable),
            "limit_reached": bool(runnable and index >= max_batches),
            "continuation_budget": max_batches,
            "minimum_plan_continuations": minimum_plan_continuations,
            "pagination_headroom": max(0, max_batches - minimum_plan_continuations),
        }
        if not summary["scheduled"]:
            return {"request": None, "summary": summary}
        next_request = {
            key: value
            for key, value in request.items()
            if key not in {"_audit_only_result", "resume_run_id", "_continuation_index"}
        }
        next_request.update({
            "resume_run_id": run_id,
            "_continuation_index": index + 1,
        })
        return {"request": next_request, "summary": summary}

    def _persist_query_cell_states(
        self,
        *,
        run_id: str,
        workflow_id: str,
        client: str,
        job: str,
        query_plan: dict[str, Any],
        channel_runs: list[dict[str, Any]],
        executed_cell_ids: set[str] | None = None,
    ) -> dict[str, Any]:
        """Persist per-query progress without turning truncation or unknown totals into completion."""
        def normalized(value: Any) -> str:
            return " ".join(str(value or "").split()).casefold()

        def integer(value: Any, default: int = 0) -> int:
            try:
                return int(value)
            except (TypeError, ValueError):
                return default

        rounds_by_query: dict[tuple[str, str], tuple[dict[str, Any], str, dict[str, Any]]] = {}
        channel_status: dict[str, tuple[str, dict[str, Any]]] = {}
        for run in channel_runs:
            if not isinstance(run, dict):
                continue
            channel = str(run.get("channel") or "").lower()
            result = run.get("result") if isinstance(run.get("result"), dict) else {}
            channel_status[channel] = (str(run.get("status") or ""), result)
            for round_item in result.get("rounds") or []:
                if isinstance(round_item, dict):
                    rounds_by_query[(channel, normalized(round_item.get("query")))] = (
                        round_item, str(run.get("status") or ""), result,
                    )

        job_id = self._job_id(client, job)
        terminal_counts: dict[str, int] = {}
        conn = self.service._connect()
        try:
            for cell in query_plan.get("cells") or []:
                if not isinstance(cell, dict):
                    continue
                conn.execute(
                    """
                    INSERT INTO agent_sourcing_query_cells
                    (run_id,workflow_id,job_id,plan_hash,cell_id,channel,query,priority,status)
                    VALUES (?,?,?,?,?,?,?,?, 'pending')
                    ON CONFLICT(run_id,cell_id) DO NOTHING
                    """,
                    (
                        run_id, workflow_id or None, job_id, str(query_plan.get("plan_hash") or ""),
                        str(cell.get("cell_id") or ""), str(cell.get("channel") or ""),
                        str(cell.get("query") or ""), integer(cell.get("priority")),
                    ),
                )

            for cell in query_plan.get("cells") or []:
                if not isinstance(cell, dict):
                    continue
                channel = str(cell.get("channel") or "").lower()
                cell_id = str(cell.get("cell_id") or "")
                match = rounds_by_query.get((channel, normalized(cell.get("query"))))
                existing = conn.execute(
                    "SELECT status,reported_total,pages_fetched,extracted_count,unique_count,cursor_json "
                    "FROM agent_sourcing_query_cells WHERE run_id=? AND cell_id=?",
                    (run_id, cell_id),
                ).fetchone()
                existing_status = str(existing["status"] or "") if existing else ""
                if executed_cell_ids is not None and cell_id not in executed_cell_ids:
                    if existing_status in {"exhausted", "platform_capped", "blocked", "failed"}:
                        terminal_counts[existing_status] = terminal_counts.get(existing_status, 0) + 1
                    continue
                reported_total: int | None = None
                extracted = 0
                unique_count = 0
                pages_fetched = 0
                cursor: Any = {}
                last_error = None
                if match:
                    round_item, run_status, channel_result = match
                    raw_total = round_item.get("result_count")
                    if raw_total is not None and str(raw_total).strip() != "":
                        reported_total = max(0, integer(raw_total))
                    extracted = max(0, integer(round_item.get("extracted_count")))
                    unique_count = max(0, integer(round_item.get("unique_count"), extracted))
                    pages_fetched = max(0, integer(round_item.get("pages_fetched"), 1))
                    cursor = round_item.get("cursor") or {}
                    round_status = str(round_item.get("status") or run_status or "completed")
                    explicit_terminal = str(round_item.get("terminal_state") or "")
                    if explicit_terminal in {"exhausted", "platform_capped", "blocked", "failed"}:
                        status = explicit_terminal
                        reason = str(round_item.get("terminal_reason") or explicit_terminal)
                    elif round_status == "failed":
                        status, reason = "failed", str(round_item.get("reason") or "query_failed")
                    elif round_status in {"blocked", "skipped", "stale_query"}:
                        status, reason = "blocked", str(round_item.get("reason") or round_status)
                    elif reported_total is None:
                        status, reason = "platform_capped", "reported_total_unknown"
                    elif extracted >= reported_total:
                        status, reason = "exhausted", "reported_total_exhausted"
                    else:
                        status, reason = "platform_capped", "reported_total_not_exhausted"
                    last_error = str(round_item.get("error") or channel_result.get("error") or "") or None
                    if existing_status in {"platform_capped", "blocked", "failed"}:
                        pages_fetched += max(0, integer(existing["pages_fetched"]))
                        extracted += max(0, integer(existing["extracted_count"]))
                        unique_count += max(0, integer(existing["unique_count"]))
                        if reported_total is None and existing["reported_total"] is not None:
                            reported_total = max(0, integer(existing["reported_total"]))
                    ledger = conn.execute(
                        "SELECT COUNT(*) AS occurrences,"
                        "COUNT(DISTINCT COALESCE(NULLIF(source_candidate_id,''),identity_key)) AS unique_count "
                        "FROM agent_candidate_recalls WHERE run_id=? AND query_cell_id=?",
                        (run_id, cell_id),
                    ).fetchone()
                    ledger_occurrences = int(ledger["occurrences"] or 0) if ledger else 0
                    ledger_unique = int(ledger["unique_count"] or 0) if ledger else 0
                    if ledger_occurrences or extracted == 0:
                        extracted = ledger_occurrences
                        unique_count = ledger_unique
                    if status == "exhausted" and reported_total is not None and unique_count < reported_total:
                        status = "platform_capped" if isinstance(cursor, dict) and cursor else "blocked"
                        reason = "duplicate_candidates_before_reported_total"
                else:
                    run_status, channel_result = channel_status.get(channel, ("", {}))
                    if existing is not None:
                        reported_total = (
                            max(0, integer(existing["reported_total"]))
                            if existing["reported_total"] is not None
                            else None
                        )
                        pages_fetched = max(0, integer(existing["pages_fetched"]))
                        extracted = max(0, integer(existing["extracted_count"]))
                        unique_count = max(0, integer(existing["unique_count"]))
                        cursor = _loads(existing["cursor_json"], {})
                    if run_status == "failed":
                        status, reason = "failed", "channel_failed_before_query"
                    elif run_status == "blocked":
                        status, reason = "blocked", "channel_blocked_before_query"
                    else:
                        status, reason = "blocked", "approved_cell_not_executed"
                    last_error = str(channel_result.get("error") or "") or None
                conn.execute(
                    """
                    UPDATE agent_sourcing_query_cells
                       SET status=?,reported_total=?,pages_fetched=?,extracted_count=?,unique_count=?,
                           cursor_json=?,retry_count=retry_count+?,
                           terminal_reason=?,last_error=?,started_at=COALESCE(started_at,datetime('now','localtime')),
                           finished_at=datetime('now','localtime'),updated_at=datetime('now','localtime')
                     WHERE run_id=? AND cell_id=?
                    """,
                    (
                        status, reported_total, pages_fetched, extracted, unique_count,
                        json.dumps(cursor, ensure_ascii=False),
                        1 if status in {"failed", "blocked"} or (
                            status == "platform_capped" and existing_status == "platform_capped"
                        ) else 0,
                        reason, last_error,
                        run_id, str(cell.get("cell_id") or ""),
                    ),
                )
                terminal_counts[status] = terminal_counts.get(status, 0) + 1
            conn.commit()
        finally:
            conn.close()
        return {"ok": True, "run_id": run_id, "stored": sum(terminal_counts.values()), "terminal_counts": terminal_counts}

    def _persist_candidate_recalls(
        self,
        *,
        run_id: str,
        workflow_id: str,
        client: str,
        job: str,
        query_plan: dict[str, Any],
        raw_candidates: dict[str, list[Any]],
        applied: dict[str, Any],
        min_score: int,
    ) -> dict[str, Any]:
        """Persist every extracted card before formal candidate intake or score filtering."""
        def normalized(value: Any) -> str:
            return re.sub(r"\s+", "", str(value or "")).casefold()

        def source_identifier(item: dict[str, Any]) -> str:
            return str(
                item.get("source_candidate_id") or item.get("candidate_id") or item.get("resume_id")
                or item.get("res_id_encode") or item.get("xsaas_id") or ""
            ).strip()

        def identity(item: dict[str, Any], channel: str, query: str) -> tuple[str, ...]:
            source_id = source_identifier(item)
            if source_id:
                return ("source", normalized(channel), normalized(query), normalized(source_id))
            return (
                "profile", normalized(channel), normalized(query), normalized(item.get("name")),
                normalized(item.get("company") or item.get("current_company")),
                normalized(item.get("title") or item.get("current_title")),
            )

        cell_by_query = {
            (str(cell.get("channel") or ""), normalized(cell.get("query"))): str(cell.get("cell_id") or "")
            for cell in query_plan.get("cells") or []
            if isinstance(cell, dict)
        }
        staged = applied.get("staged") if isinstance(applied.get("staged"), dict) else {}
        disposition: dict[tuple[str, ...], str] = {}
        staged_by_identity: dict[tuple[str, ...], dict[str, Any]] = {}
        for key, state in (("accepted", "accepted"), ("existing", "existing"), ("batch_duplicates", "batch_duplicate")):
            for raw in staged.get(key) or []:
                if not isinstance(raw, dict):
                    continue
                channel = str(raw.get("channel") or raw.get("source") or "").lower()
                query = str(raw.get("source_query") or raw.get("query") or "")
                item_identity = identity(raw, channel, query)
                disposition[item_identity] = state
                staged_by_identity[item_identity] = raw
        for error in staged.get("errors") or []:
            raw = error.get("raw") if isinstance(error, dict) and isinstance(error.get("raw"), dict) else {}
            channel = str(raw.get("channel") or raw.get("source") or "").lower()
            query = str(raw.get("source_query") or raw.get("query") or "")
            item_identity = identity(raw, channel, query)
            disposition[item_identity] = "invalid"
            staged_by_identity[item_identity] = raw

        receipts = (
            applied.get("intake", {}).get("receipts") or []
            if isinstance(applied.get("intake"), dict)
            else []
        )
        accepted = [item for item in staged.get("accepted") or [] if isinstance(item, dict)]
        receipt_by_identity: dict[tuple[str, ...], dict[str, Any]] = {}
        for accepted_item, receipt in zip(accepted, receipts, strict=False):
            if not isinstance(receipt, dict):
                continue
            accepted_channel = str(accepted_item.get("channel") or accepted_item.get("source") or "").lower()
            accepted_query = str(accepted_item.get("source_query") or accepted_item.get("query") or "")
            receipt_by_identity[identity(accepted_item, accepted_channel, accepted_query)] = receipt
        job_id = self._job_id(client, job)
        stored = 0
        by_state: dict[str, int] = {}
        conn = self.service._connect()
        try:
            for channel, values in raw_candidates.items():
                normalized_channel = str(channel or "unknown").lower()
                for index, raw in enumerate(values if isinstance(values, list) else [], 1):
                    if not isinstance(raw, dict):
                        continue
                    source_query = " ".join(str(raw.get("source_query") or raw.get("query") or "").split())
                    source_id = str(
                        raw.get("source_candidate_id") or raw.get("candidate_id") or raw.get("resume_id")
                        or raw.get("res_id_encode")
                        or raw.get("xsaas_id") or raw.get("resume_url") or raw.get("source_url") or ""
                    ).strip()
                    name = str(raw.get("name") or "").strip()
                    company = str(raw.get("company") or raw.get("current_company") or "").strip()
                    title = str(raw.get("title") or raw.get("current_title") or "").strip()
                    identity_key = "|".join((normalized(name), normalized(company), normalized(title)))
                    if not source_id:
                        source_id = "anon_" + hashlib.sha256(identity_key.encode("utf-8")).hexdigest()[:20]
                    try:
                        score = int(raw["fit_score"]) if raw.get("fit_score") is not None else None
                    except (TypeError, ValueError):
                        score = None
                    item_identity = identity(raw, normalized_channel, source_query)
                    state = disposition.get(item_identity, "not_intaked")
                    staged_item = staged_by_identity.get(item_identity, {})
                    exclusion_reason = None
                    if score is not None and score < min_score:
                        exclusion_reason = "score_below_threshold"
                    elif state == "existing":
                        exclusion_reason = "existing_candidate"
                    elif state == "batch_duplicate":
                        exclusion_reason = "same_batch_duplicate"
                    elif state == "invalid":
                        exclusion_reason = "normalization_error"
                    elif state == "not_intaked":
                        exclusion_reason = "not_in_intake_output"
                    receipt: dict[str, Any] = {}
                    if state == "accepted":
                        receipt = receipt_by_identity.get(
                            identity(raw, normalized_channel, source_query),
                            {},
                        )
                    receipt_status = str(receipt.get("status") or "")
                    if receipt_status in {"existing", "existing_relation"}:
                        state = receipt_status
                        exclusion_reason = "existing_candidate" if receipt_status == "existing" else "existing_relation"
                    page_number = max(1, int(raw.get("page_number") or raw.get("page") or 1))
                    position_index = max(0, int(raw.get("position_index") or index))
                    query_cell_id = cell_by_query.get((normalized_channel, normalized(source_query)), "")
                    recall_identity = "|".join((
                        run_id, normalized_channel, query_cell_id, source_id,
                        str(page_number), str(position_index), normalized(source_query),
                    ))
                    recall_id = "recall_" + hashlib.sha256(recall_identity.encode("utf-8")).hexdigest()[:24]
                    conn.execute(
                        """
                        INSERT INTO agent_candidate_recalls
                        (recall_id,run_id,workflow_id,job_id,query_cell_id,channel,source_candidate_id,
                         source_query,source_url,page_number,position_index,identity_key,candidate_name,
                         company,title,fit_score,fit_level,duplicate_state,exclusion_reason,detail_status,
                         candidate_id,job_candidate_id,raw_json)
                        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                        ON CONFLICT(recall_id) DO UPDATE SET
                          fit_score=excluded.fit_score,fit_level=excluded.fit_level,
                          duplicate_state=excluded.duplicate_state,exclusion_reason=excluded.exclusion_reason,
                          detail_status=excluded.detail_status,candidate_id=excluded.candidate_id,
                          job_candidate_id=excluded.job_candidate_id,raw_json=excluded.raw_json,
                          updated_at=datetime('now','localtime')
                        """,
                        (
                            recall_id, run_id, workflow_id or None, job_id, query_cell_id,
                            normalized_channel, source_id, source_query,
                            str(raw.get("resume_url") or raw.get("source_url") or raw.get("url") or ""),
                            page_number, position_index, identity_key, name, company, title, score,
                            str(raw.get("fit_level") or "") or None, state, exclusion_reason,
                            str(
                                staged_item.get("resume_capture_status")
                                or staged_item.get("detail_status")
                                or raw.get("resume_capture_status")
                                or raw.get("detail_status")
                                or "not_requested"
                            ),
                            int(receipt.get("candidate_id") or 0) or None,
                            int(receipt.get("job_candidate_id") or 0) or None,
                            json.dumps(raw, ensure_ascii=False),
                        ),
                    )
                    stored += 1
                    by_state[state] = by_state.get(state, 0) + 1
            conn.commit()
        finally:
            conn.close()
        return {"ok": True, "stored": stored, "run_id": run_id, "by_state": by_state}

    def _build_coverage_certificate(
        self,
        *,
        run_id: str,
        workflow_id: str,
        client: str,
        job: str,
        query_plan: dict[str, Any],
    ) -> dict[str, Any]:
        """Issue an auditable coverage certificate without claiming a hidden platform population."""
        job_id = self._job_id(client, job)
        conn = self.service._connect()
        try:
            cell_rows = conn.execute(
                "SELECT * FROM agent_sourcing_query_cells WHERE run_id=? ORDER BY priority,cell_id",
                (run_id,),
            ).fetchall()
            recall_row = conn.execute(
                """
                SELECT COUNT(*) AS raw_occurrences,
                       COUNT(DISTINCT channel || ':' || COALESCE(NULLIF(source_candidate_id,''),identity_key)) AS channel_unique_identities,
                       COUNT(DISTINCT CASE
                           WHEN REPLACE(identity_key,'|','')<>'' THEN identity_key
                           ELSE channel || ':' || source_candidate_id END) AS global_unique_identities,
                       SUM(CASE WHEN duplicate_state IN ('existing','existing_relation','batch_duplicate') THEN 1 ELSE 0 END) AS duplicate_occurrences,
                       SUM(CASE WHEN exclusion_reason='score_below_threshold' THEN 1 ELSE 0 END) AS below_threshold,
                       SUM(CASE WHEN query_cell_id='' THEN 1 ELSE 0 END) AS unmapped_occurrences,
                       COUNT(DISTINCT CASE WHEN job_candidate_id IS NOT NULL THEN job_candidate_id END) AS formally_intaked,
                       SUM(CASE WHEN detail_status='complete' THEN 1 ELSE 0 END) AS detail_complete,
                       SUM(CASE WHEN detail_status='partial' THEN 1 ELSE 0 END) AS detail_partial,
                       SUM(CASE WHEN detail_status='failed' THEN 1 ELSE 0 END) AS detail_failed
                FROM agent_candidate_recalls WHERE run_id=?
                """,
                (run_id,),
            ).fetchone()
            ledger_rows = conn.execute(
                """
                SELECT query_cell_id,COUNT(*) AS occurrences
                FROM agent_candidate_recalls
                WHERE run_id=? AND query_cell_id<>''
                GROUP BY query_cell_id
                """,
                (run_id,),
            ).fetchall()
            assessment_count = int(conn.execute(
                """
                SELECT COUNT(DISTINCT a.job_candidate_id)
                FROM agent_candidate_recalls r
                JOIN agent_candidate_assessments a ON a.job_candidate_id=r.job_candidate_id AND a.is_current=1
                WHERE r.run_id=?
                """,
                (run_id,),
            ).fetchone()[0])

            status_counts = {
                status: 0 for status in ("pending", "exhausted", "platform_capped", "blocked", "failed")
            }
            state_by_cell: dict[str, dict[str, Any]] = {}
            ledger_by_cell = {str(row["query_cell_id"]): int(row["occurrences"] or 0) for row in ledger_rows}
            executed = 0
            platform_totals: dict[str, dict[str, int | None]] = {}
            for row in cell_rows:
                item = _row(row)
                state_by_cell[str(item.get("cell_id") or "")] = item
                status = str(item.get("status") or "")
                if status in status_counts:
                    status_counts[status] += 1
                if status in {"exhausted", "platform_capped", "failed"} or int(item.get("pages_fetched") or 0) > 0:
                    executed += 1
                channel = str(item.get("channel") or "unknown")
                totals = platform_totals.setdefault(
                    channel,
                    {"reported_query_total": 0, "reported_total_known_cells": 0, "extracted_occurrences": 0},
                )
                if item.get("reported_total") is not None:
                    totals["reported_query_total"] = int(totals["reported_query_total"] or 0) + int(item["reported_total"])
                    totals["reported_total_known_cells"] = int(totals["reported_total_known_cells"] or 0) + 1
                totals["extracted_occurrences"] = int(totals["extracted_occurrences"] or 0) + int(item.get("extracted_count") or 0)

            expected_extracted = sum(int(item.get("extracted_count") or 0) for item in state_by_cell.values())
            mapped_occurrences = sum(ledger_by_cell.values())
            unmapped_occurrences = int(_row(recall_row).get("unmapped_occurrences") or 0)
            mismatched_cells = sum(
                int(item.get("extracted_count") or 0) != int(ledger_by_cell.get(cell_id, 0))
                for cell_id, item in state_by_cell.items()
            )
            evidence_integrity_passed = bool(
                unmapped_occurrences == 0
                and expected_extracted == mapped_occurrences
                and mismatched_cells == 0
            )

            all_companies: set[str] = set()
            all_groups: set[str] = set()
            executed_companies: set[str] = set()
            executed_groups: set[str] = set()
            for cell in query_plan.get("cells") or []:
                if not isinstance(cell, dict):
                    continue
                state = state_by_cell.get(str(cell.get("cell_id") or ""), {})
                was_executed = str(state.get("status") or "") in {"exhausted", "platform_capped", "failed"} or int(state.get("pages_fetched") or 0) > 0
                for ref in cell.get("provenance") or []:
                    if not isinstance(ref, dict):
                        continue
                    company = str(ref.get("company") or "").strip()
                    group = str(ref.get("group") or "").strip()
                    if company:
                        all_companies.add(company)
                        if was_executed:
                            executed_companies.add(company)
                    if group:
                        all_groups.add(group)
                        if was_executed:
                            executed_groups.add(group)

            approved = len([cell for cell in query_plan.get("cells") or [] if isinstance(cell, dict)])
            dimensions = query_plan.get("dimensions") if isinstance(query_plan.get("dimensions"), dict) else {}
            semantics = (
                query_plan.get("execution_semantics")
                if isinstance(query_plan.get("execution_semantics"), dict)
                else {}
            )
            evaluation_modes = (
                semantics.get("evaluation_constraints")
                if isinstance(semantics.get("evaluation_constraints"), dict)
                else {}
            )
            dimension_execution = {
                "retrieval_axes": semantics.get("retrieval_axes") or ["channel", "query"],
                "platform_filters_applied": semantics.get("platform_filters") or [],
                "dimensions": {
                    key: {
                        "approved_values": [str(value) for value in dimensions.get(key) or []],
                        "retrieval_filter_applied": False,
                        "evaluation_mode": str(evaluation_modes.get(key) or "post_recall_evaluation"),
                    }
                    for key in ("locations", "levels", "scenarios")
                },
            }
            if not evidence_integrity_passed:
                coverage_status = "coverage_unknown"
                defensible_claim = "查询执行记录与原始召回台账不一致，候选人覆盖未知"
            elif approved > 0 and status_counts["exhausted"] == approved:
                coverage_status = "approved_query_cells_exhausted"
                defensible_claim = "已穷尽批准的渠道关键词查询单元；地点、职级、场景未作为平台筛选执行"
            elif status_counts["platform_capped"]:
                coverage_status = "platform_truncated"
                defensible_claim = "已执行部分批准查询单元，但平台截断导致候选人总体覆盖未知"
            else:
                coverage_status = "coverage_unknown"
                defensible_claim = "批准的渠道关键词查询单元尚未完全执行，候选人总体覆盖未知"
            unknown_reasons: list[str] = []
            if status_counts["platform_capped"]:
                unknown_reasons.append("platform_truncated")
            if status_counts["blocked"]:
                unknown_reasons.append("blocked_query_cells")
            if status_counts["failed"]:
                unknown_reasons.append("failed_query_cells")
            if status_counts["pending"]:
                unknown_reasons.append("pending_query_cells")
            if not evidence_integrity_passed:
                unknown_reasons.append("recall_ledger_mismatch")
            unknown_reasons.append("platform_candidate_population_denominator_unavailable")

            recall = _row(recall_row)
            certificate_id = "coverage_" + hashlib.sha256(
                f"{run_id}|{query_plan.get('plan_hash') or ''}".encode("utf-8")
            ).hexdigest()[:24]
            certificate = {
                "schema_version": "coverage_certificate_v1",
                "certificate_id": certificate_id,
                "run_id": run_id,
                "workflow_id": workflow_id,
                "job_id": job_id,
                "plan_hash": str(query_plan.get("plan_hash") or ""),
                "issued_at": datetime.now().isoformat(timespec="seconds"),
                "coverage_status": coverage_status,
                "strategy_elements": {
                    "companies_approved": len(all_companies),
                    "companies_executed": len(executed_companies),
                    "keyword_groups_approved": len(all_groups),
                    "keyword_groups_executed": len(executed_groups),
                },
                "dimension_execution": dimension_execution,
                "query_cells": {
                    "approved": approved,
                    "executed": executed,
                    **status_counts,
                },
                "platform_query_totals": platform_totals,
                "candidate_recall": {
                    "raw_occurrences": int(recall.get("raw_occurrences") or 0),
                    "unique_identities": int(recall.get("global_unique_identities") or 0),
                    "global_unique_identities": int(recall.get("global_unique_identities") or 0),
                    "channel_unique_identities": int(recall.get("channel_unique_identities") or 0),
                    "duplicate_occurrences": int(recall.get("duplicate_occurrences") or 0),
                    "below_threshold": int(recall.get("below_threshold") or 0),
                    "formally_intaked": int(recall.get("formally_intaked") or 0),
                },
                "evidence_integrity": {
                    "passed": evidence_integrity_passed,
                    "expected_extracted_occurrences": expected_extracted,
                    "mapped_recall_occurrences": mapped_occurrences,
                    "unmapped_recall_occurrences": unmapped_occurrences,
                    "mismatched_query_cells": mismatched_cells,
                },
                "detail_completeness": {
                    "complete": int(recall.get("detail_complete") or 0),
                    "partial": int(recall.get("detail_partial") or 0),
                    "failed": int(recall.get("detail_failed") or 0),
                },
                "assessment": {"completed_unique_candidates": assessment_count},
                "claims": {
                    "all_candidates_covered": False,
                    "defensible_claim": defensible_claim,
                    "coverage_unknown_reasons": list(dict.fromkeys(unknown_reasons)),
                },
            }
            conn.execute(
                """
                INSERT INTO agent_sourcing_coverage_certificates
                (certificate_id,run_id,workflow_id,job_id,plan_hash,coverage_status,certificate_json)
                VALUES (?,?,?,?,?,?,?)
                ON CONFLICT(run_id) DO UPDATE SET
                  certificate_id=excluded.certificate_id,workflow_id=excluded.workflow_id,
                  job_id=excluded.job_id,plan_hash=excluded.plan_hash,
                  coverage_status=excluded.coverage_status,certificate_json=excluded.certificate_json,
                  issued_at=datetime('now','localtime')
                """,
                (
                    certificate_id, run_id, workflow_id or None, job_id,
                    str(query_plan.get("plan_hash") or ""), coverage_status,
                    json.dumps(certificate, ensure_ascii=False),
                ),
            )
            conn.commit()
            return certificate
        finally:
            conn.close()

    def _persist_sourcing_attributions(
        self, applied: dict[str, Any], strategy: dict[str, Any], workflow_id: str, client: str, job: str,
    ) -> dict[str, Any]:
        staged = applied.get("staged") if isinstance(applied.get("staged"), dict) else {}
        accepted = staged.get("accepted") if isinstance(staged.get("accepted"), list) else []
        intake = applied.get("intake") if isinstance(applied.get("intake"), dict) else {}
        receipts = intake.get("receipts") if isinstance(intake.get("receipts"), list) else []
        strategy_channels = strategy.get("channels") if isinstance(strategy.get("channels"), dict) else {}
        query_meta: dict[tuple[str, str], dict[str, Any]] = {}
        for channel, entries in strategy_channels.items():
            for entry in entries if isinstance(entries, list) else []:
                item = entry if isinstance(entry, dict) else {"query": entry}
                query = " ".join(str(item.get("query") or "").split())
                if query:
                    query_meta[(str(channel), query)] = item
        strategy_hash = hashlib.sha256(json.dumps(strategy, ensure_ascii=False, sort_keys=True).encode("utf-8")).hexdigest() if strategy else ""
        strategy_model = str((strategy.get("generation") or {}).get("model") or "") if isinstance(strategy.get("generation"), dict) else ""
        stored = 0
        channel_new: dict[str, int] = {}
        conn = self.service._connect()
        try:
            for index, candidate in enumerate(accepted):
                item = candidate if isinstance(candidate, dict) else {}
                receipt = (
                    receipts[index]
                    if index < len(receipts) and isinstance(receipts[index], dict)
                    else {}
                )
                job_candidate_id = int(receipt.get("job_candidate_id") or 0)
                if not job_candidate_id:
                    continue
                channel = str(item.get("channel") or item.get("source") or "unknown").lower()
                query = " ".join(str(item.get("source_query") or "").split()) or "未记录关键词"
                meta = query_meta.get((channel, query), {})
                conn.execute(
                    """
                    INSERT INTO agent_sourcing_attributions
                    (job_candidate_id,candidate_id,job_id,workflow_id,strategy_hash,strategy_model,
                     channel,source_query,source_round,source_purpose)
                    VALUES (?,?,?,?,?,?,?,?,?,?)
                    ON CONFLICT(job_candidate_id,channel,source_query) DO UPDATE SET
                      workflow_id=COALESCE(excluded.workflow_id,workflow_id),
                      strategy_hash=COALESCE(NULLIF(excluded.strategy_hash,''),strategy_hash),
                      strategy_model=COALESCE(NULLIF(excluded.strategy_model,''),strategy_model),
                      source_round=COALESCE(NULLIF(excluded.source_round,''),source_round),
                      source_purpose=COALESCE(NULLIF(excluded.source_purpose,''),source_purpose),
                      updated_at=datetime('now','localtime')
                    """,
                    (
                        job_candidate_id, int(receipt.get("candidate_id") or 0) or None,
                        self._job_id(client, job), workflow_id or None, strategy_hash or None, strategy_model or None,
                        channel, query, str(meta.get("round") or ""), str(meta.get("purpose") or ""),
                    ),
                )
                stored += 1
                channel_new[channel] = channel_new.get(channel, 0) + 1
            conn.commit()
        finally:
            conn.close()
        return {"stored": stored, "strategy_hash": strategy_hash, "workflow_id": workflow_id, "channel_new": channel_new}

    def _persist_sourcing_funnel(
        self,
        *,
        run_id: str,
        workflow_id: str,
        client: str,
        job: str,
        channel_runs: list[dict[str, Any]],
        channel_candidates: dict[str, list[Any]],
        applied: dict[str, Any],
        attributions: dict[str, Any],
        company_vocab: set[str] | None = None,
    ) -> dict[str, Any]:
        """每个 run×channel 落一行寻访漏斗，并把 0 结果归因回写到 channel_runs 条目上。"""
        staged = applied.get("staged") if isinstance(applied.get("staged"), dict) else {}
        intake_dups: dict[str, int] = {}
        for key in ("existing", "batch_duplicates"):
            entries = staged.get(key) if isinstance(staged.get(key), list) else []
            for entry in entries:
                item = entry if isinstance(entry, dict) else {}
                dup_channel = str(item.get("channel") or item.get("source") or "unknown").lower()
                intake_dups[dup_channel] = intake_dups.get(dup_channel, 0) + 1
        channel_new = attributions.get("channel_new") if isinstance(attributions.get("channel_new"), dict) else {}
        job_id = self._job_id(client, job)
        rows: list[dict[str, Any]] = []
        for run in channel_runs:
            channel = str(run.get("channel") or "unknown").lower()
            status = str(run.get("status") or "completed")
            result = run.get("result") if isinstance(run.get("result"), dict) else {}
            candidates = [item for item in channel_candidates.get(channel) or [] if isinstance(item, dict)]
            rounds = [entry for entry in result.get("rounds") or [] if isinstance(entry, dict)]
            detail = result.get("detail_capture") if isinstance(result.get("detail_capture"), dict) else {}
            recall = sum(_round_int(entry, "result_count") for entry in rounds)
            extracted = sum(_round_int(entry, "extracted_count") for entry in rounds)
            if extracted <= 0 and not rounds:
                extracted = len(candidates)
            scored = [item for item in candidates if item.get("fit_score") is not None]
            high_score = 0
            for item in scored:
                try:
                    if int(item.get("fit_score") or 0) >= 65:
                        high_score += 1
                except (TypeError, ValueError):
                    continue
            zero_attribution = ""
            if not candidates:
                zero_attribution = classify_zero_result(
                    channel,
                    status,
                    result,
                    dedupe_count=max(0, extracted - len(candidates)),
                    company_vocab=company_vocab,
                )
                run["zero_attribution"] = zero_attribution
            rows.append(
                {
                    "run_id": run_id,
                    "workflow_id": workflow_id or None,
                    "job_id": job_id,
                    "client": client,
                    "job": job,
                    "channel": channel,
                    "status": status,
                    "query_count": len(rounds),
                    "queries_json": json.dumps(rounds, ensure_ascii=False),
                    "recall_count": recall,
                    "extracted_count": extracted,
                    "dedupe_count": max(0, extracted - len(candidates)),
                    "unique_count": len(candidates),
                    "detail_complete": _round_int(detail, "complete"),
                    "detail_partial": _round_int(detail, "partial"),
                    "detail_failed": _round_int(detail, "failed"),
                    "intake_duplicate_count": int(intake_dups.get(channel, 0)),
                    "intake_new_count": int(channel_new.get(channel, 0) or 0),
                    "assessed_count": len(scored),
                    "high_score_count": high_score,
                    "zero_attribution": zero_attribution or None,
                    "error": _trim_error(result.get("error")) or None,
                }
            )
        conn = self.service._connect()
        try:
            for row in rows:
                cell_rows = conn.execute(
                    """
                    SELECT cell_id,query,status,reported_total,pages_fetched,extracted_count,
                           unique_count,terminal_reason,last_error
                    FROM agent_sourcing_query_cells
                    WHERE run_id=? AND channel=?
                    ORDER BY priority,cell_id
                    """,
                    (run_id, row["channel"]),
                ).fetchall()
                recall_stats = _row(conn.execute(
                    """
                    SELECT COUNT(*) AS raw_occurrences,
                           COUNT(DISTINCT CASE
                               WHEN REPLACE(identity_key,'|','')<>'' THEN identity_key
                               ELSE channel || ':' || source_candidate_id END) AS unique_identities,
                           SUM(CASE WHEN detail_status='complete' THEN 1 ELSE 0 END) AS detail_complete,
                           SUM(CASE WHEN detail_status='partial' THEN 1 ELSE 0 END) AS detail_partial,
                           SUM(CASE WHEN detail_status='failed' THEN 1 ELSE 0 END) AS detail_failed,
                           SUM(CASE WHEN duplicate_state IN ('existing','existing_relation','batch_duplicate') THEN 1 ELSE 0 END) AS intake_duplicates,
                           COUNT(DISTINCT CASE WHEN duplicate_state='accepted' AND job_candidate_id IS NOT NULL THEN job_candidate_id END) AS intake_new,
                           COUNT(DISTINCT CASE WHEN fit_score IS NOT NULL
                               THEN COALESCE(NULLIF(source_candidate_id,''),identity_key) END) AS assessed,
                           COUNT(DISTINCT CASE WHEN fit_score>=65
                               THEN COALESCE(NULLIF(source_candidate_id,''),identity_key) END) AS high_score
                    FROM agent_candidate_recalls
                    WHERE run_id=? AND channel=?
                    """,
                    (run_id, row["channel"]),
                ).fetchone())
                if cell_rows:
                    cell_items = [_row(item) for item in cell_rows]
                    statuses = {str(item.get("status") or "") for item in cell_items}
                    if "failed" in statuses:
                        row["status"] = "failed"
                    elif "blocked" in statuses:
                        row["status"] = "blocked"
                    elif "platform_capped" in statuses:
                        row["status"] = "platform_capped"
                    elif statuses == {"exhausted"}:
                        row["status"] = "completed"
                    row["query_count"] = len(cell_items)
                    row["queries_json"] = json.dumps(cell_items, ensure_ascii=False)
                    row["recall_count"] = sum(
                        int(item.get("reported_total") or 0)
                        for item in cell_items if item.get("reported_total") is not None
                    )
                    row["extracted_count"] = sum(int(item.get("extracted_count") or 0) for item in cell_items)
                raw_occurrences = int(recall_stats.get("raw_occurrences") or 0)
                unique_identities = int(recall_stats.get("unique_identities") or 0)
                row["dedupe_count"] = max(0, int(row["extracted_count"]) - unique_identities)
                row["unique_count"] = unique_identities
                row["detail_complete"] = max(
                    int(row["detail_complete"]), int(recall_stats.get("detail_complete") or 0),
                )
                row["detail_partial"] = max(
                    int(row["detail_partial"]), int(recall_stats.get("detail_partial") or 0),
                )
                row["detail_failed"] = max(
                    int(row["detail_failed"]), int(recall_stats.get("detail_failed") or 0),
                )
                row["intake_duplicate_count"] = max(
                    int(row["intake_duplicate_count"]), int(recall_stats.get("intake_duplicates") or 0),
                )
                row["intake_new_count"] = int(recall_stats.get("intake_new") or 0)
                row["assessed_count"] = int(recall_stats.get("assessed") or 0)
                row["high_score_count"] = int(recall_stats.get("high_score") or 0)
                if raw_occurrences > 0:
                    row["zero_attribution"] = None
                conn.execute(
                    """
                    INSERT INTO agent_sourcing_funnel
                    (run_id,workflow_id,job_id,client,job,channel,status,query_count,queries_json,
                     recall_count,extracted_count,dedupe_count,unique_count,
                     detail_complete,detail_partial,detail_failed,
                     intake_duplicate_count,intake_new_count,assessed_count,high_score_count,
                     zero_attribution,error)
                    VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                    ON CONFLICT(run_id,channel) DO UPDATE SET
                      workflow_id=COALESCE(excluded.workflow_id,workflow_id),
                      job_id=excluded.job_id,
                      status=excluded.status,
                      query_count=excluded.query_count,
                      queries_json=excluded.queries_json,
                      recall_count=excluded.recall_count,
                      extracted_count=excluded.extracted_count,
                      dedupe_count=excluded.dedupe_count,
                      unique_count=excluded.unique_count,
                      detail_complete=excluded.detail_complete,
                      detail_partial=excluded.detail_partial,
                      detail_failed=excluded.detail_failed,
                      intake_duplicate_count=excluded.intake_duplicate_count,
                      intake_new_count=excluded.intake_new_count,
                      assessed_count=excluded.assessed_count,
                      high_score_count=excluded.high_score_count,
                      zero_attribution=excluded.zero_attribution,
                      error=excluded.error,
                      updated_at=datetime('now','localtime')
                    """,
                    (
                        row["run_id"], row["workflow_id"], row["job_id"], row["client"], row["job"],
                        row["channel"], row["status"], row["query_count"], row["queries_json"],
                        row["recall_count"], row["extracted_count"], row["dedupe_count"], row["unique_count"],
                        row["detail_complete"], row["detail_partial"], row["detail_failed"],
                        row["intake_duplicate_count"], row["intake_new_count"],
                        row["assessed_count"], row["high_score_count"],
                        row["zero_attribution"], row["error"],
                    ),
                )
            conn.commit()
        finally:
            conn.close()
        return {"ok": True, "stored": len(rows), "run_id": run_id}

    def _record_sourcing_funnel_failure(
        self,
        *,
        run_id: str,
        workflow_id: str,
        client: str,
        job: str,
        channel: str,
        error: str,
    ) -> None:
        """渠道 runner 在合并前直接失败时，尽力留下一行失败漏斗（绝不掩盖原始异常）。"""
        try:
            conn = self.service._connect()
            try:
                conn.execute(
                    """
                    INSERT INTO agent_sourcing_funnel
                    (run_id,workflow_id,job_id,client,job,channel,status,zero_attribution,error)
                    VALUES (?,?,?,?,?,?,?,?,?)
                    ON CONFLICT(run_id,channel) DO UPDATE SET
                      status=excluded.status,
                      zero_attribution=excluded.zero_attribution,
                      error=excluded.error,
                      updated_at=datetime('now','localtime')
                    """,
                    (
                        run_id, workflow_id or None, self._job_id(client, job), client, job,
                        channel, "failed", classify_zero_result(channel, "failed", {"error": _trim_error(error)}), _trim_error(error),
                    ),
                )
                conn.commit()
            finally:
                conn.close()
        except Exception:
            pass

    @staticmethod
    def _run_command(
        command: list[str],
        timeout: int = 300,
        *,
        cancel_check: Callable[[], bool] | None = None,
        poll_interval: float = 0.25,
    ) -> subprocess.CompletedProcess[str]:
        if cancel_check is None:
            return subprocess.run(command, capture_output=True, text=True, timeout=timeout)
        proc = subprocess.Popen(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            start_new_session=True,
        )
        started = time.monotonic()
        try:
            while True:
                if cancel_check():
                    raise ExternalExecutionCancelled("工作流已停止，当前渠道执行已终止")
                remaining = timeout - (time.monotonic() - started)
                if remaining <= 0:
                    raise subprocess.TimeoutExpired(command, timeout)
                try:
                    stdout, stderr = proc.communicate(timeout=min(max(0.01, poll_interval), remaining))
                    return subprocess.CompletedProcess(command, proc.returncode, stdout, stderr)
                except subprocess.TimeoutExpired:
                    continue
        except (ExternalExecutionCancelled, subprocess.TimeoutExpired):
            try:
                os.killpg(proc.pid, signal.SIGTERM)
                proc.communicate(timeout=2)
            except (ProcessLookupError, subprocess.TimeoutExpired):
                try:
                    os.killpg(proc.pid, signal.SIGKILL)
                except ProcessLookupError:
                    pass
                proc.communicate()
            raise

    def _run(
        self,
        command: list[str],
        timeout: int = 300,
        *,
        cancel_check: Callable[[], bool] | None = None,
    ) -> subprocess.CompletedProcess[str]:
        proc = self._run_command(command, timeout, cancel_check=cancel_check)
        if proc.returncode != 0:
            message, detail = _command_failure_summary(proc.stdout, proc.stderr, proc.returncode)
            detail["command"] = command
            raise CommandExecutionError(message, detail)
        return proc

    def _run_external(
        self,
        command: list[str],
        timeout: int,
        *,
        cancel_check: Callable[[], bool] | None,
    ) -> subprocess.CompletedProcess[str]:
        if cancel_check is None:
            return self._run(command, timeout)
        return self._run(command, timeout, cancel_check=cancel_check)

    def _run_json(
        self,
        command: list[str],
        timeout: int = 300,
        *,
        cancel_check: Callable[[], bool] | None = None,
    ) -> dict[str, Any]:
        proc = self._run(command, timeout, cancel_check=cancel_check)
        try:
            payload = json.loads(proc.stdout)
        except json.JSONDecodeError as exc:
            raise RuntimeError("能力脚本未返回合法 JSON") from exc
        if not isinstance(payload, dict):
            raise RuntimeError("能力脚本必须返回 JSON 对象")
        return payload

    def _run_external_json(
        self,
        command: list[str],
        timeout: int,
        *,
        cancel_check: Callable[[], bool] | None,
    ) -> dict[str, Any]:
        if cancel_check is None:
            return self._run_json(command, timeout)
        return self._run_json(command, timeout, cancel_check=cancel_check)

    def _job(self, context: dict[str, Any]) -> dict[str, Any]:
        if context.get("type") != "job" or not context.get("id"):
            raise ValueError("该能力需要明确岗位上下文")
        conn = self.service._connect()
        try:
            columns = {str(row[1]) for row in conn.execute("PRAGMA table_info(positions)").fetchall()}
            position_join = ""
            position_fields = ""
            if columns:
                position_join = "LEFT JOIN positions p ON p.client=c.name AND p.title=j.title"
                position_fields = ",p.id AS position_id,p.location AS position_location,p.salary,p.education,p.experience,p.responsibilities,p.requirements,p.headcount,p.deadline,p.liepin_status"
            row = conn.execute(
                f"""
                SELECT j.*,c.name AS client{position_fields}
                FROM jobs j JOIN clients c ON c.id=j.client_id {position_join}
                WHERE j.id=? ORDER BY COALESCE(p.id,0) DESC LIMIT 1
                """ if columns else "SELECT j.*,c.name AS client FROM jobs j JOIN clients c ON c.id=j.client_id WHERE j.id=?",
                (int(context["id"]),),
            ).fetchone()
            if row is None:
                raise ValueError(f"找不到岗位：{context['id']}")
            result = _row(row)
            profile = conn.execute(
                "SELECT * FROM position_profiles WHERE client=? AND position=? ORDER BY id DESC LIMIT 1",
                (result.get("client"), result.get("title")),
            ).fetchone() if self._table(conn, "position_profiles") else None
            result["profile"] = _row(profile)
            return result
        finally:
            conn.close()

    @staticmethod
    def _table(conn: sqlite3.Connection, name: str) -> bool:
        return conn.execute("SELECT 1 FROM sqlite_master WHERE type IN ('table','view') AND name=?", (name,)).fetchone() is not None

    def _candidate(self, context: dict[str, Any]) -> dict[str, Any]:
        if context.get("type") != "candidate" or not context.get("id"):
            raise ValueError("该能力需要明确人选上下文")
        return build_candidate_context(self.service.db_path, int(context["id"]))

    def _latest_assessment(self, job_candidate_id: int) -> dict[str, Any]:
        conn = self.service._connect()
        try:
            row = conn.execute(
                "SELECT * FROM agent_candidate_assessments WHERE job_candidate_id=? AND is_current=1 ORDER BY id DESC LIMIT 1",
                (int(job_candidate_id),),
            ).fetchone()
            if row is None:
                result = self.service.submit_assessment(int(job_candidate_id), wait=True)
                return result.get("assessment") or {}
            item = _row(row)
            for key, target, default in (
                ("criteria_json", "criteria", {}), ("strengths_json", "strengths", []),
                ("gaps_json", "gaps", []), ("risks_json", "risks", []),
                ("verification_questions_json", "verification_questions", []),
                ("citations_json", "citations", []),
            ):
                item[target] = _loads(item.get(key), default)
            return item
        finally:
            conn.close()

    def _path(self, category: str, title: str, suffix: str) -> Path:
        folder = self.output_dir / category
        folder.mkdir(parents=True, exist_ok=True)
        return folder / f"{_slug(title)}-{datetime.now():%Y%m%d-%H%M%S}.{suffix.lstrip('.')}"

    @staticmethod
    def _artifact(kind: str, title: str, *, content: str = "", file_path: Path | None = None,
                  mime_type: str = "text/markdown", validation: str = "passed", metadata: dict[str, Any] | None = None) -> dict[str, Any]:
        return {
            "type": kind, "title": title, "content": content,
            "file_path": str(file_path) if file_path else "", "mime_type": mime_type,
            "validation_status": validation, "metadata": metadata or {},
        }

    @staticmethod
    def _blocked(summary: str, missing: list[str], references: list[dict[str, Any]] | None = None) -> dict[str, Any]:
        return {"summary": summary, "blocked": True, "missing_inputs": missing, "references": references or []}

    def _job_reference(self, job: dict[str, Any]) -> list[dict[str, Any]]:
        return [{"type": "job", "id": job.get("id"), "label": job.get("title"), "subtitle": job.get("client")}]

    def _candidate_reference(self, candidate: dict[str, Any]) -> list[dict[str, Any]]:
        identity, position, relation = candidate["identity"], candidate["position"], candidate["relation"]
        return [{"type": "candidate", "id": relation["job_candidate_id"], "label": identity.get("name"), "subtitle": f"{position.get('client','')} / {position.get('job','')}"}]

    def _candidate_event(self, candidate: dict[str, Any], event_type: str, status: str, summary: str, raw: dict[str, Any]) -> int:
        relation = candidate["relation"]
        conn = self.service._connect()
        try:
            cursor = conn.execute(
                """
                INSERT INTO candidate_events
                (job_candidate_id,person_id,job_id,event_type,event_status,event_time,summary,raw_json,source_table,source_id)
                VALUES (?,?,?,?,?,datetime('now','localtime'),?,?,?,?)
                """,
                (relation["job_candidate_id"], relation["person_id"], relation.get("job_id"), event_type, status,
                 summary[:1000], json.dumps(raw, ensure_ascii=False), "agent_workflow", str(raw.get("workflow_id") or "")),
            )
            conn.commit()
            return int(cursor.lastrowid)
        finally:
            conn.close()

    def _followup(self, candidate: dict[str, Any], task_type: str, reason: str, inputs: dict[str, Any], days: int = 2) -> int | None:
        conn = self.service._connect()
        try:
            if not self._table(conn, "followup_tasks"):
                return None
            columns = {str(row[1]) for row in conn.execute("PRAGMA table_info(followup_tasks)").fetchall()}
            next_id = int(conn.execute("SELECT COALESCE(MAX(id),0)+1 FROM followup_tasks").fetchone()[0])
            identity, position, relation = candidate["identity"], candidate["position"], candidate["relation"]
            values = {
                "id": next_id, "candidate_id": relation.get("source_candidate_id"), "candidate_name": identity.get("name"),
                "candidate_company": identity.get("company"), "client": position.get("client"), "position": position.get("job"),
                "task_type": task_type, "priority": int(inputs.get("priority") or 2),
                "due_at": (datetime.now() + timedelta(days=days)).strftime("%Y-%m-%d %H:%M:%S"),
                "status": "open", "reason": reason[:1000], "source_table": "agent_workflow",
                "source_id": inputs.get("step_id"), "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "job_candidate_id": relation["job_candidate_id"],
            }
            keys = [key for key in values if key in columns]
            conn.execute(
                f"INSERT INTO followup_tasks ({','.join(keys)}) VALUES ({','.join('?' for _ in keys)})",
                [values[key] for key in keys],
            )
            conn.commit()
            return next_id
        finally:
            conn.close()

    def run_job_intake(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        job = self._job(context)
        content = "\n".join([
            f"# 岗位接入 · {job['client']} / {job['title']}", "",
            f"- 状态：{job.get('status') or '未设置'}", f"- 地点：{job.get('position_location') or job.get('location') or '待补充'}",
            f"- 薪资：{job.get('salary') or '待补充'}", f"- 编制：{job.get('headcount') or '待补充'}",
            "", "## 岗位摘要", str(job.get("summary") or job.get("profile", {}).get("jd_analysis_summary") or "待补充"),
        ])
        return {"summary": "岗位接入信息已从 v3 事实源整理。", "references": self._job_reference(job),
                "artifacts": [self._artifact("job_brief", "岗位接入简报", content=content)]}

    def run_jd_calibration(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        job = self._job(context)
        profile = job.get("profile") or {}
        facts = {
            "客户": job.get("client"), "岗位": job.get("title"),
            "地点": job.get("position_location") or job.get("location"), "薪资": job.get("salary"),
            "学历": job.get("education") or profile.get("education_requirement"),
            "经验": job.get("experience") or profile.get("experience_requirement"),
            "硬门槛": _loads(profile.get("hard_requirements_json"), []) or _loads(job.get("hard_requirements"), []),
            "核心能力": _loads(profile.get("ability_keywords_json"), []) or _loads(job.get("ability_keywords"), []),
        }
        missing = [key for key in ("岗位", "地点", "硬门槛", "核心能力") if not facts.get(key)]
        content = "# JD 校准\n\n" + "\n".join(f"- {key}：{json.dumps(value, ensure_ascii=False) if isinstance(value, list) else value or '待补充'}" for key, value in facts.items())
        return {"summary": f"JD 校准完成；{len(missing)} 项待补充。", "missing_inputs": missing,
                "references": self._job_reference(job), "artifacts": [self._artifact("jd_calibration", "JD 校准结果", content=content, validation="needs_input" if missing else "passed")]}

    def run_job_library_update(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        objective = str(inputs.get("objective") or "")
        directions = self._job_update_directions(objective, inputs)
        conn = self.service._connect()
        try:
            client = self._job_update_client(conn, context, objective, inputs)
            if not client:
                return self._blocked("岗位库更新缺少客户名，未写入数据库。", ["client"])
            profiles = self._job_update_profiles(conn, client, objective, directions)
            if not profiles:
                return self._blocked("没有找到可用于更新岗位库的岗位画像，未写入数据库。", ["position_profiles"])
            client_row = conn.execute("SELECT id,name FROM clients WHERE name=?", (client,)).fetchone()
            if client_row is None:
                cursor = conn.execute("INSERT INTO clients(name) VALUES (?)", (client,))
                client_id = int(cursor.lastrowid)
            else:
                client_id = int(client_row["id"])

            changes: list[dict[str, Any]] = []
            now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            for profile in profiles:
                change = self._upsert_profile_job(conn, client_id, client, _row(profile), now)
                changes.append(change)
            archived = []
            if self._should_archive_legacy_update(objective, directions):
                archived = self._archive_legacy_split_jobs(conn, client, [item["title"] for item in changes], now)
            conn.commit()
        except Exception:
            conn.rollback()
            raise
        finally:
            conn.close()

        sync_result: dict[str, Any] = {"skipped": True}
        if not inputs.get("skip_sync"):
            sync_result = self._sync_job_library(client)
        receipt = {
            "client": client,
            "directions": directions,
            "changes": changes,
            "archived_legacy": archived,
            "sync": sync_result,
        }
        content = "# 岗位库更新回执\n\n```json\n" + json.dumps(receipt, ensure_ascii=False, indent=2) + "\n```"
        return {
            "summary": f"岗位库已更新：{client} 新增/更新 {len(changes)} 个岗位方向。",
            "references": [{"type": "job", "id": item["job_id"], "label": item["title"], "subtitle": client} for item in changes],
            "artifacts": [self._artifact("job_library_update_receipt", "岗位库更新回执", content=content, metadata=receipt)],
            "job_library_update": receipt,
        }

    def _job_update_client(self, conn: sqlite3.Connection, context: dict[str, Any], objective: str, inputs: dict[str, Any]) -> str:
        explicit = str(inputs.get("client") or "").strip()
        if explicit:
            return explicit
        if context.get("type") == "job" and context.get("id"):
            row = conn.execute(
                "SELECT c.name FROM jobs j JOIN clients c ON c.id=j.client_id WHERE j.id=?",
                (int(context["id"]),),
            ).fetchone()
            if row:
                return str(row["name"] or "")
        rows = conn.execute("SELECT name FROM clients ORDER BY LENGTH(name) DESC").fetchall()
        for row in rows:
            name = str(row["name"] or "")
            if name and name in objective:
                return name
        return ""

    @staticmethod
    def _job_update_directions(objective: str, inputs: dict[str, Any]) -> list[str]:
        raw = inputs.get("directions")
        if isinstance(raw, list):
            values = [str(item).strip() for item in raw if str(item).strip()]
            if values:
                return values
        mapping = (("PC", ("PC", "pc", "电脑")), ("服务器", ("服务器", "server", "Server")), ("ADAS", ("ADAS", "adas", "智驾", "辅助驾驶")))
        directions = [label for label, tokens in mapping if any(token in objective for token in tokens)]
        return directions

    def _job_update_profiles(self, conn: sqlite3.Connection, client: str, objective: str, directions: list[str]) -> list[sqlite3.Row]:
        if not self._table(conn, "position_profiles"):
            return []
        clauses = ["client=?"]
        params: list[Any] = [client]
        if "技术市场" in objective:
            clauses.append("position LIKE '%技术市场%'")
        if directions:
            direction_clauses = []
            for direction in directions:
                if direction == "PC":
                    direction_clauses.append("(position LIKE '%PC%' OR position LIKE '%电脑%')")
                elif direction == "服务器":
                    direction_clauses.append("position LIKE '%服务器%'")
                elif direction == "ADAS":
                    direction_clauses.append("position LIKE '%ADAS%'")
            if direction_clauses:
                clauses.append("(" + " OR ".join(direction_clauses) + ")")
        rows = conn.execute(
            f"SELECT * FROM position_profiles WHERE {' AND '.join(clauses)} ORDER BY id",
            params,
        ).fetchall()
        if directions:
            rows = [row for row in rows if "服务器或PC市场" not in str(row["position"] or "")]
        return rows

    def _upsert_profile_job(self, conn: sqlite3.Connection, client_id: int, client: str, profile: dict[str, Any], now: str) -> dict[str, Any]:
        title = str(profile.get("position") or "").strip()
        if not title:
            raise ValueError("岗位画像缺少 position")
        position = conn.execute(
            "SELECT * FROM positions WHERE client=? AND title=? ORDER BY id DESC LIMIT 1",
            (client, title),
        ).fetchone() if self._table(conn, "positions") else None
        position_data = _row(position)
        hard = _list_text(profile.get("hard_requirements_json")) or str(position_data.get("hard_requirements") or "")
        ability = _list_text(profile.get("ability_keywords_json")) or str(position_data.get("ability_keywords") or "")
        targets = _list_text(profile.get("target_companies_json")) or str(position_data.get("target_companies") or "")
        exclusions = _list_text(profile.get("exclusion_tags_json")) or str(position_data.get("exclusions") or "")
        search_words = _list_text(profile.get("search_keywords_json")) or str(position_data.get("search_words") or title)
        summary = str(profile.get("jd_analysis_summary") or position_data.get("summary") or f"{client}/{title} 岗位画像已按方向拆分。")
        status = str(position_data.get("status") or "待启动")
        existing = conn.execute("SELECT * FROM jobs WHERE client_id=? AND title=?", (client_id, title)).fetchone()
        if existing:
            job_id = int(existing["id"])
            lifecycle = str(existing["lifecycle_stage"] or "sourcing")
            if lifecycle in {"archived", "closed", "cancelled"}:
                lifecycle = "sourcing"
            conn.execute(
                """
                UPDATE jobs
                   SET status=?,lifecycle_stage=?,hard_requirements=?,ability_keywords=?,
                       target_companies=?,exclusions=?,search_words=?,summary=?,
                       source_layer='v3_position_profile',updated_at=datetime('now','localtime')
                 WHERE id=?
                """,
                (status, lifecycle, hard, ability, targets, exclusions, search_words, summary, job_id),
            )
            action = "updated"
        else:
            cursor = conn.execute(
                """
                INSERT INTO jobs
                (client_id,title,status,lifecycle_stage,source_layer,hard_requirements,ability_keywords,
                 target_companies,exclusions,search_words,summary,created_at,updated_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,datetime('now','localtime'),datetime('now','localtime'))
                """,
                (client_id, title, status, "sourcing", "v3_position_profile", hard, ability, targets, exclusions, search_words, summary),
            )
            job_id = int(cursor.lastrowid)
            action = "created"
        if position_data:
            conn.execute(
                """
                UPDATE positions
                   SET hard_requirements=?,ability_keywords=?,target_companies=?,exclusions=?,
                       search_words=?,summary=?,updated_at=datetime('now','localtime')
                 WHERE client=? AND title=?
                """,
                (hard, ability, targets, exclusions, search_words, summary, client, title),
            )
        self._upsert_job_metrics(conn, job_id, profile, now)
        return {"action": action, "job_id": job_id, "title": title, "position_id": position_data.get("id")}

    def _upsert_job_metrics(self, conn: sqlite3.Connection, job_id: int, profile: dict[str, Any], now: str) -> None:
        payload = {
            "metric_date": now[:10],
            "priority": "P0-最急｜用户指定最高优先级",
            "risk": _list_text(profile.get("risk_points_json")),
            "next_keywords_json": profile.get("search_keywords_json") or "[]",
            "target_companies_json": profile.get("target_companies_json") or "[]",
            "exclude_terms_json": profile.get("exclusion_tags_json") or "[]",
        }
        row = conn.execute("SELECT id FROM job_pipeline_metrics WHERE job_id=? ORDER BY id DESC LIMIT 1", (job_id,)).fetchone()
        if row:
            conn.execute(
                """
                UPDATE job_pipeline_metrics
                   SET metric_date=?,priority=?,risk=?,next_keywords_json=?,
                       target_companies_json=?,exclude_terms_json=?,data_gap=0
                 WHERE id=?
                """,
                (payload["metric_date"], payload["priority"], payload["risk"], payload["next_keywords_json"],
                 payload["target_companies_json"], payload["exclude_terms_json"], int(row["id"])),
            )
        else:
            conn.execute(
                """
                INSERT INTO job_pipeline_metrics
                (job_id,metric_date,a_count,b_count,p0_count,p1_count,published_count,under_review_count,
                 contacted_count,pending_followup_count,priority,risk,next_keywords_json,target_companies_json,
                 exclude_terms_json,data_gap)
                VALUES (?,?,0,0,0,0,0,0,0,0,?,?,?,?,?,0)
                """,
                (job_id, payload["metric_date"], payload["priority"], payload["risk"], payload["next_keywords_json"],
                 payload["target_companies_json"], payload["exclude_terms_json"]),
            )

    @staticmethod
    def _should_archive_legacy_update(objective: str, directions: list[str]) -> bool:
        return bool(directions) and any(token in objective for token in ("拆", "分成", "三个", "分别", "独立"))

    def _archive_legacy_split_jobs(self, conn: sqlite3.Connection, client: str, active_titles: list[str], now: str) -> list[dict[str, Any]]:
        rows = conn.execute(
            """
            SELECT j.id,j.title
            FROM jobs j JOIN clients c ON c.id=j.client_id
            WHERE c.name=? AND j.title LIKE '%技术市场%' AND j.title LIKE '%服务器或PC市场%'
            """,
            (client,),
        ).fetchall()
        archived = []
        for row in rows:
            if row["title"] in active_titles:
                continue
            conn.execute(
                """
                UPDATE jobs
                   SET status='已拆分为PC/服务器/ADAS-保留历史',
                       lifecycle_stage='archived',
                       closed_reason='岗位画像已拆分为独立方向，历史人选关系保留',
                       closed_at=datetime('now','localtime'),
                       updated_at=datetime('now','localtime')
                 WHERE id=?
                """,
                (int(row["id"]),),
            )
            archived.append({"job_id": int(row["id"]), "title": row["title"]})
        if self._table(conn, "positions"):
            conn.execute(
                """
                UPDATE positions
                   SET status='已拆分为PC/服务器/ADAS-保留历史',
                       updated_at=datetime('now','localtime')
                 WHERE client=? AND title LIKE '%技术市场%' AND title LIKE '%服务器或PC市场%'
                """,
                (client,),
            )
        return archived

    def _sync_job_library(self, client: str) -> dict[str, Any]:
        command = [
            self.python,
            "/Users/messi/.codex/skills/a-system-workbench/scripts/a_system_sync.py",
            "--client",
            client,
            "--db",
            str(self.service.db_path),
            "--no-open",
            "--no-backup",
        ]
        proc = subprocess.run(command, capture_output=True, text=True, timeout=300)
        return {
            "ok": proc.returncode == 0,
            "returncode": proc.returncode,
            "stdout": (proc.stdout or "")[-4000:],
            "stderr": (proc.stderr or "")[-2000:],
        }

    def _strategy_learning_context(self, job: dict[str, Any]) -> dict[str, Any]:
        conn = self.service._connect()
        try:
            experiments = [
                _row(row)
                for row in conn.execute(
                    """
                    SELECT channel,query,result_count,viewed_count,extracted_count,
                           recommended_count,reply_count,positive_reply_count,noise_notes,status,updated_at
                    FROM search_experiments
                    WHERE client=? AND position=?
                    ORDER BY datetime(COALESCE(updated_at,run_time,created_at)) DESC,id DESC LIMIT 24
                    """,
                    (job["client"], job["title"]),
                ).fetchall()
            ] if self._table(conn, "search_experiments") else []
            correction = _row(conn.execute(
                "SELECT * FROM strategy_corrections WHERE client=? AND position=? ORDER BY id DESC LIMIT 1",
                (job["client"], job["title"]),
            ).fetchone()) if self._table(conn, "strategy_corrections") else {}
            business_outcomes = [
                _row(row)
                for row in conn.execute(
                    """
                    SELECT sa.channel,sa.source_query,
                           COUNT(DISTINCT sa.job_candidate_id) AS attributed_candidates,
                           COUNT(sf.id) AS signal_count,ROUND(COALESCE(SUM(sf.weight),0),2) AS experience_score,
                           SUM(sf.signal_type='review_pass') AS review_pass,
                           SUM(sf.signal_type='contacted') AS contacted,
                           SUM(sf.signal_type='recommended') AS recommended,
                           SUM(sf.signal_type='stopped') AS stopped,
                           SUM(sf.signal_type IN ('client_approved','client_interview','client_offer','client_hired')) AS client_positive,
                           SUM(sf.signal_type='client_rejected') AS client_rejected
                    FROM agent_sourcing_attributions sa
                    LEFT JOIN agent_sourcing_feedback sf ON sf.attribution_id=sa.id
                    WHERE sa.job_id=?
                    GROUP BY sa.channel,sa.source_query
                    ORDER BY experience_score DESC,signal_count DESC,sa.id DESC LIMIT 30
                    """,
                    (job["id"],),
                ).fetchall()
            ] if self._table(conn, "agent_sourcing_attributions") and self._table(conn, "agent_sourcing_feedback") else []
            pending_adjustments: list[dict[str, Any]] = []
            if self._table(conn, "agent_sourcing_adjustments"):
                pending_adjustments = [
                    _row(row)
                    for row in conn.execute(
                        """
                        SELECT a.*, p.display_name AS candidate_name
                        FROM agent_sourcing_adjustments a
                        LEFT JOIN job_candidates jc ON jc.id=a.candidate_id
                        LEFT JOIN people p ON p.id=jc.person_id
                        WHERE a.job_id=? AND a.status='pending'
                        ORDER BY a.id
                        """,
                        (job["id"],),
                    ).fetchall()
                ]
        finally:
            conn.close()
        memories = self.service.search_memories(
            f"{job['client']} {job['title']} 寻访关键词 搜索效果",
            context_type="job", context_id=job["id"], client=str(job["client"]), job=str(job["title"]), limit=8,
        )
        return {
            "historical_experiments": experiments,
            "explicit_corrections": correction,
            "business_outcomes": business_outcomes,
            "approved_memories": memories.get("memories") or [] if memories.get("mode") == "active" else [],
            "memory_mode": memories.get("mode") or "off",
            "memory_hits": len(memories.get("memories") or []),
            "stop_note_adjustments": pending_adjustments,
        }

    _STOP_NOTE_ADJUSTMENT_LABELS: dict[str, str] = {
        "add_keyword": "补充关键词",
        "remove_keyword": "移除关键词",
        "exclude_company": "排除公司",
        "add_company": "补充公司",
        "add_filter": "添加过滤",
        "adjust_salary_range": "调整薪资区间",
    }

    @staticmethod
    def _stop_note_adjustments_summary(adjustments: list[dict[str, Any]]) -> str:
        """把 pending 调整格式化为策略 prompt 可消费的文本摘要。"""
        if not adjustments:
            return ""
        parts: list[str] = []
        for item in adjustments:
            if not isinstance(item, dict):
                continue
            adjust_type = str(item.get("adjust_type") or item.get("type") or "").strip()
            value = str(item.get("value") or "").strip()
            if not adjust_type or not value:
                continue
            label = RecruitingCapabilityRuntime._STOP_NOTE_ADJUSTMENT_LABELS.get(adjust_type, adjust_type)
            rationale = str(item.get("rationale") or "")[:60]
            parts.append(f"{label}：{value}" + (f"（来源：{rationale}）" if rationale else ""))
        return "；".join(parts)

    def _apply_stop_note_adjustments(self, job_id: int, workflow_id: str) -> None:
        """策略生成成功后，将本轮消费的 pending 调整标记为 applied，并记录应用时候选池基线。失败静默。"""
        if not job_id:
            return
        try:
            conn = self.service._connect()
            try:
                # 轮次 = 本工作流已完成 search_strategy 步数（含当前步），无工作流时默认 1。
                if workflow_id:
                    round_row = conn.execute(
                        "SELECT COUNT(*) AS n FROM agent_workflow_steps WHERE workflow_id=? AND capability_id='search_strategy' AND status='completed'",
                        (workflow_id,),
                    ).fetchone()
                    applied_round = (int(round_row["n"] if round_row else 0) or 0) + 1
                else:
                    applied_round = 1
                # 应用时候选池基线快照（供"调整前后效果对比"追踪）。
                baseline = self._candidate_pool_baseline(conn, job_id)
                conn.execute(
                    """
                    UPDATE agent_sourcing_adjustments
                       SET status='applied', applied_at=datetime('now','localtime'), applied_round=?,
                           baseline_json=?
                     WHERE job_id=? AND status='pending'
                    """,
                    (applied_round, json.dumps(baseline, ensure_ascii=False), job_id),
                )
                conn.commit()
            finally:
                conn.close()
        except Exception:
            pass

    @staticmethod
    def _candidate_pool_baseline(conn: sqlite3.Connection, job_id: int) -> dict[str, int]:
        """候选池质量基线：总池 / 待复核 / 已触达 / 已停止（按 clean_stage 口径）。"""
        rows = conn.execute(
            """
            SELECT
              COUNT(*) AS total,
              SUM(clean_stage LIKE 'S1%' OR clean_stage LIKE 'X1%' OR clean_stage LIKE 'H1%' OR clean_stage LIKE 'H5%' OR clean_stage LIKE 'S2%') AS pending_review,
              SUM(clean_stage LIKE 'S3%' OR clean_stage LIKE 'S4%' OR clean_stage LIKE 'S5%' OR clean_stage LIKE 'S6%' OR clean_stage LIKE 'S7%' OR clean_stage LIKE 'S8%' OR clean_stage LIKE 'S9%' OR clean_stage LIKE 'S10%' OR clean_stage LIKE 'S11%' OR clean_stage LIKE 'S12%' OR clean_stage LIKE 'S13%' OR clean_stage LIKE 'X2%' OR clean_stage LIKE 'X3%' OR clean_stage LIKE 'X4%' OR clean_stage LIKE 'X5%') AS contacted,
              SUM(clean_stage LIKE '%停止%' OR clean_stage LIKE '%淘汰%' OR clean_stage LIKE '%不通过%') AS stopped
            FROM job_candidates WHERE job_id=?
            """,
            (job_id,),
        ).fetchone()
        return {
            "total": int(rows["total"] or 0),
            "pending_review": int(rows["pending_review"] or 0),
            "contacted": int(rows["contacted"] or 0),
            "stopped": int(rows["stopped"] or 0),
        }

    @staticmethod
    def _normalize_strategy_entries(values: Any) -> list[dict[str, str]]:
        rows: list[dict[str, str]] = []
        for index, value in enumerate(values if isinstance(values, list) else []):
            item = value if isinstance(value, dict) else {"query": value}
            query = " ".join(str(item.get("query") or "").split())
            if not query:
                continue
            rows.append({
                "round": str(item.get("round") or f"model-{index + 1}"),
                "query": query,
                "purpose": str(item.get("purpose") or "大模型生成的岗位寻访查询"),
                "evidence": str(item.get("evidence") or "岗位事实"),
            })
        return rows

    @staticmethod
    def _strategy_term_grounded(term: str, canonical_text: str) -> bool:
        canonical = re.sub(r"\s+", "", canonical_text.lower())
        normalized = re.sub(r"\s+", "", term.lower())
        if normalized and normalized in canonical:
            return True
        parts = [value for value in re.split(r"[\s/、,，;；]+", term) if value]
        if len(parts) > 1:
            return all(RecruitingCapabilityRuntime._strategy_term_grounded(value, canonical_text) for value in parts)
        ascii_tokens = re.findall(r"[A-Za-z][A-Za-z0-9+-]*", term)
        if ascii_tokens and not all(token.lower() in canonical for token in ascii_tokens):
            return False
        chinese = "".join(re.findall(r"[\u4e00-\u9fff]", term))
        if not chinese:
            return bool(ascii_tokens)
        if len(chinese) <= 2:
            return chinese.lower() in canonical
        bigrams = [chinese[index:index + 2] for index in range(len(chinese) - 1)]
        matched = sum(value.lower() in canonical for value in bigrams)
        return matched / max(1, len(bigrams)) >= 0.5

    def _validate_model_strategy(
        self, raw: dict[str, Any], fallback: dict[str, Any], job: dict[str, Any], learning: dict[str, Any], max_queries: int,
    ) -> dict[str, Any]:
        profile = job.get("profile") or {}
        canonical_text = " ".join(
            str(value or "")
            for value in (job.get("title"), job.get("responsibilities"), job.get("requirements"), job.get("hard_requirements"), job.get("exclusions"))
        ).lower()
        legacy_terms = [
            str(value).strip()
            for key in ("ability_keywords_json", "search_keywords_json")
            for value in _loads(profile.get(key), [])
            if str(value).strip()
        ]
        unsupported = {
            term for term in legacy_terms
            if len(re.sub(r"\W+", "", term)) >= 3 and not self._strategy_term_grounded(term, canonical_text)
        }
        raw_channels = raw.get("channels") if isinstance(raw.get("channels"), dict) else {}
        fallback_channels = fallback.get("channels") if isinstance(fallback.get("channels"), dict) else {}
        removed: list[str] = []
        channels: dict[str, list[dict[str, str]]] = {}
        for channel in ("liepin", "xsaas"):
            accepted: list[dict[str, str]] = []
            seen: set[str] = set()
            for item in [*self._normalize_strategy_entries(raw_channels.get(channel)), *self._normalize_strategy_entries(fallback_channels.get(channel))]:
                normalized = re.sub(r"\s+", "", item["query"].lower())
                bad = next((term for term in unsupported if re.sub(r"\s+", "", term.lower()) in normalized), "")
                if bad:
                    removed.append(f"{item['query']}（无岗位依据：{bad}）")
                    continue
                if normalized in seen:
                    continue
                seen.add(normalized)
                accepted.append(item)
                if len(accepted) >= max_queries:
                    break
            channels[channel] = accepted
        model_used = any(self._normalize_strategy_entries(raw_channels.get(channel)) for channel in ("liepin", "xsaas"))
        return {
            **{key: value for key, value in fallback.items() if key not in {"channels"}},
            "channels": channels,
            "target_companies": raw.get("target_companies") or fallback.get("target_companies") or [],
            "strategy_summary": str(raw.get("strategy_summary") or "围绕岗位硬门槛、应用场景与目标公司分层寻访。"),
            "learning_notes": raw.get("learning_notes") if isinstance(raw.get("learning_notes"), list) else [],
            "generation": {
                "mode": "llm" if model_used else "deterministic_fallback",
                "model": self.service.llm.model,
                "consultant_mode": "senior_consultant_v1",
                "memory_mode": learning["memory_mode"],
                "memory_hits": learning["memory_hits"],
                "experiment_count": len(learning["historical_experiments"]),
                "removed_unsupported_queries": removed,
            },
        }

    def _capture_search_learning(self, client: str, job: str, queries: list[Any]) -> dict[str, Any]:
        query_values = [
            " ".join(str((item or {}).get("query") if isinstance(item, dict) else item or "").split())
            for item in queries
        ]
        query_values = list(dict.fromkeys(value for value in query_values if value))
        if not query_values:
            return {"stored_memories": 0, "queries": 0}
        conn = self.service._connect()
        try:
            placeholders = ",".join("?" for _ in query_values)
            rows = conn.execute(
                f"""
                SELECT channel,query,result_count,viewed_count,recommended_count,positive_reply_count,noise_notes
                FROM search_experiments
                WHERE client=? AND position=? AND query IN ({placeholders})
                ORDER BY datetime(COALESCE(updated_at,run_time,created_at)) DESC,id DESC
                """,
                [client, job, *query_values],
            ).fetchall()
        finally:
            conn.close()
        stored = 0
        seen: set[tuple[str, str]] = set()
        for row in rows:
            key = (str(row["channel"] or ""), str(row["query"] or ""))
            if key in seen:
                continue
            seen.add(key)
            recommended = int(row["recommended_count"] or 0)
            viewed = int(row["viewed_count"] or 0)
            outcome = "有效" if recommended > 0 else "待降权"
            content = (
                f"{client}/{job} 搜索经验：{key[0] or '未知渠道'} 关键词“{key[1]}”{outcome}；"
                f"查看 {viewed}，推荐 {recommended}，正向回复 {int(row['positive_reply_count'] or 0)}。"
            )
            if row["noise_notes"]:
                content += f" 噪音：{str(row['noise_notes'])[:160]}"
            self.service.store_memory(
                scope_type="job", scope_id=str(self._job_id(client, job)), memory_type="search_outcome",
                content=content, source_type="search_experiment", source_id=f"{key[0]}:{key[1]}",
                confidence=0.9 if recommended > 0 else 0.72,
            )
            stored += 1
        return {"stored_memories": stored, "queries": len(query_values)}

    def _query_plan_learning_metrics(self, job_id: int) -> list[dict[str, Any]]:
        """Aggregate prior marginal yield, overlap and downstream business feedback per query."""
        conn = self.service._connect()
        try:
            rows = conn.execute(
                """
                SELECT qc.channel,qc.query,COUNT(DISTINCT qc.run_id) AS runs,
                       COUNT(r.id) AS raw_occurrences,
                       COUNT(DISTINCT r.channel || ':' || COALESCE(NULLIF(r.source_candidate_id,''),r.identity_key)) AS unique_identities,
                       COUNT(DISTINCT CASE WHEN r.detail_status='complete'
                           THEN r.channel || ':' || COALESCE(NULLIF(r.source_candidate_id,''),r.identity_key) END) AS detail_complete,
                       COUNT(DISTINCT CASE WHEN r.job_candidate_id IS NOT NULL
                           THEN r.job_candidate_id END) AS intake_count
                FROM agent_sourcing_query_cells qc
                LEFT JOIN agent_candidate_recalls r
                  ON r.run_id=qc.run_id AND r.query_cell_id=qc.cell_id
                WHERE qc.job_id=?
                GROUP BY qc.channel,qc.query
                """,
                (job_id,),
            ).fetchall()
            feedback_rows = conn.execute(
                """
                SELECT sa.channel,sa.source_query,COALESCE(SUM(sf.weight),0) AS business_score
                FROM agent_sourcing_attributions sa
                LEFT JOIN agent_sourcing_feedback sf ON sf.attribution_id=sa.id
                WHERE sa.job_id=?
                GROUP BY sa.channel,sa.source_query
                """,
                (job_id,),
            ).fetchall()
        finally:
            conn.close()
        feedback = {
            (str(row["channel"] or ""), " ".join(str(row["source_query"] or "").split()).casefold()): float(row["business_score"] or 0)
            for row in feedback_rows
        }
        metrics: list[dict[str, Any]] = []
        for row in rows:
            runs = max(1, int(row["runs"] or 0))
            raw = max(0, int(row["raw_occurrences"] or 0))
            unique = max(0, int(row["unique_identities"] or 0))
            complete = max(0, int(row["detail_complete"] or 0))
            intake = max(0, int(row["intake_count"] or 0))
            channel = str(row["channel"] or "")
            query = str(row["query"] or "")
            metrics.append({
                "channel": channel,
                "query": query,
                "runs": runs,
                "raw_occurrences": raw,
                "unique_identities": unique,
                "detail_complete": complete,
                "intake_count": intake,
                "unique_yield_per_run": round(unique / runs, 4),
                "overlap_rate": round(1 - unique / raw, 4) if raw else 0.0,
                "business_score": feedback.get((channel, " ".join(query.split()).casefold()), 0.0),
            })
        return metrics

    def _job_id(self, client: str, job: str) -> int:
        conn = self.service._connect()
        try:
            row = conn.execute(
                "SELECT j.id FROM jobs j JOIN clients c ON c.id=j.client_id WHERE c.name=? AND j.title=? ORDER BY j.id LIMIT 1",
                (client, job),
            ).fetchone()
            return int(row["id"]) if row else 0
        finally:
            conn.close()

    def run_search_strategy(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        job = self._job(context)
        knowledge_health = strategy_v2.knowledge_base_health()
        max_queries = min(12, int(inputs.get("max_queries") or 6))
        command = [self.python, str(MULTICHANNEL), "plan", "--db", str(self.service.db_path), "--client", str(job["client"]), "--job", str(job["title"]), "--max-queries", str(max_queries)]
        try:
            fallback = self._run_json(command, 60)
        except Exception:
            profile = job.get("profile") or {}
            fallback = {
                "client": job["client"], "job": job["title"],
                "queries": _loads(profile.get("search_keywords_json"), [])[:6],
                "target_companies": _loads(profile.get("target_companies_json"), [])[:20],
                "exclusions": _loads(profile.get("exclusion_tags_json"), [])[:20],
                "source": "v3_fallback",
            }
        learning = self._strategy_learning_context(job)
        # S4-1：策略生成前先定级（L1/L2/L3）并盘点四锚点；顾问在 Copilot 侧的
        # 放行/锚点回复经工作流上下文 strategy_clarification 传入（ consultant_override /
        # consultant_answers ），推断项按 PRD §1 保持 inferred:true + confidence。
        clarification = context.get("strategy_clarification") if isinstance(context.get("strategy_clarification"), dict) else {}
        revision_evidence = _revision_consultant_evidence(context)
        understanding = context.get("intent_understanding") if isinstance(context.get("intent_understanding"), dict) else {}
        locked_items = [
            {"quote": str(item.get("quote") or "").strip(), "kind": str(item.get("kind") or "other")}
            for item in (understanding.get("constraints") or [])
            if isinstance(item, dict) and str(item.get("quote") or "").strip()
        ]
        for quote in context.get("locked_constraints") or []:
            normalized_quote = str(quote or "").strip()
            if normalized_quote and not any(item["quote"] == normalized_quote for item in locked_items):
                locked_items.append({"quote": normalized_quote, "kind": "other"})
        consultant_answers = "；".join(
            value
            for value in (
                str(clarification.get("consultant_answers") or "").strip(),
                revision_evidence,
                "；".join(item["quote"] for item in locked_items),
            )
            if value
        )
        consultant = {
            "consultant_override": bool(clarification.get("consultant_override")),
            "consultant_answers": consultant_answers,
        }
        archetype, archetype_trace = strategy_v2.match_job_archetype(job.get("client"), job.get("title"))
        classification = strategy_v2.classify_strategy_input(
            job, archetype=archetype, consultant_answers=consultant["consultant_answers"]
        )
        classification["trace"] = [*archetype_trace, *classification["trace"]]
        # S4-2：知识库消费 —— 客户画像挂载（精确/别名/模糊需确认）、公司图谱召回、
        # restricted 白名单约束。画像与图谱进 LLM 输入；restricted 键值不进任何生成
        # 上下文，仅由运行时并入 negative_rules（source=restricted_client）。
        profile_match, profile_trace = knowledge_base.match_client_profile(job.get("client"))
        restricted_info, restricted_trace = knowledge_base.load_restricted_constraints(job.get("client"))
        graph, graph_trace = knowledge_base.load_company_graph()
        # 知识飞轮二期：公司校准覆盖层（company_calibrations 表，仅 status='calibrated'）
        # 合并进图谱后再推导公司池；命中公司标注 source=consultant_calibrated 并留痕。
        # 覆盖层每次运行直接读库（表很小，不做进程内缓存，校准提交即时生效）；
        # 覆盖层为空/加载失败时不改图谱、不动 trace，输出与现状逐字节一致。
        calibration_overlay, _overlay_load_trace = knowledge_base.load_calibration_overlay(self.service.db_path)
        if calibration_overlay:
            graph, overlay_trace = knowledge_base.apply_calibration_overlay(graph, calibration_overlay)
            graph_trace = [*graph_trace, *overlay_trace]
        graph_query = " ".join(
            part
            for part in (
                str(job.get("title") or ""),
                str(job.get("ability_keywords") or ""),
                knowledge_base.profile_context(profile_match["profile"]).get("track", "") if profile_match else "",
            )
            if part
        )
        graph_pool, graph_pool_trace = knowledge_base.derive_graph_pool(graph, query_text=graph_query)
        # 目标池不得含客户本公司（图谱按赛道召回时本公司常高分命中）
        client_raw = " ".join(str(job.get("client") or "").split())
        client_norm = knowledge_base.normalize_client_name(client_raw)
        before = len(graph_pool)
        graph_pool = [
            company
            for company in graph_pool
            if not knowledge_base.name_match_rule(client_raw, client_norm, company["name"])[0]
        ]
        if len(graph_pool) < before:
            graph_pool_trace.append(f"已从图谱池剔除客户本公司 {before - len(graph_pool)} 家")
        # S4-3：排除规则引擎 —— 第 4 步之后强制过五类检查清单（PRD §4），
        # 逐类输出 适用/不适用+依据；禁挖名单/竞业从 restricted 层按客户继承。
        # 只在运行时并入 negative_rules，绝不进入 LLM 输入。
        negative_checklist, checklist_trace = negative_rules.build_negative_rule_checklist(
            job,
            restricted_info=restricted_info,
            archetype=archetype,
            consultant_answers=consultant["consultant_answers"],
        )
        classification["trace"] = [
            *classification["trace"], *profile_trace, *restricted_trace, *graph_trace, *graph_pool_trace, *checklist_trace,
        ]
        # 知识飞轮二期：技能本体（step4 别名归一/相关词提示，source=kb_skill）与
        # 职级映射（step3 优先查 kb_level_mapping，source=kb_level，查不到走 LLM/原型路径）。
        skill_ontology, skill_ontology_trace = knowledge_base.load_skill_ontology()
        level_mapping, level_mapping_trace = knowledge_base.load_level_mapping()
        level_hit, level_hit_trace = knowledge_base.map_level(job.get("title"), level_mapping)
        classification["trace"] = [
            *classification["trace"], *skill_ontology_trace, *level_mapping_trace, *level_hit_trace,
        ]
        client_profile_payload: dict[str, Any] = {"matched": False}
        if profile_match:
            client_profile_payload = {
                "matched": True,
                "name": profile_match["name"],
                "rule": profile_match["rule"],
                "needs_confirmation": profile_match["needs_confirmation"],
                "context": knowledge_base.profile_context(profile_match["profile"]),
            }
        payload = {
            "canonical_position": {
                "client": job["client"], "job": job["title"],
                "requirements": job.get("requirements") or job.get("hard_requirements") or "",
                "responsibilities": job.get("responsibilities") or "",
                "education": job.get("education") or "", "experience": job.get("experience") or "",
                "hard_requirements": job.get("hard_requirements") or "",
                "exclusions": job.get("exclusions") or "", "location": job.get("position_location") or job.get("location") or "",
                "objective": inputs.get("objective") or "",
            },
            "legacy_profile_suggestions": {
                "ability_keywords": _loads((job.get("profile") or {}).get("ability_keywords_json"), []),
                "search_keywords": _loads((job.get("profile") or {}).get("search_keywords_json"), []),
                "target_companies": _loads((job.get("profile") or {}).get("target_companies_json"), []),
            },
            "input_classification": {
                "input_level": classification["input_level"],
                "anchors": classification["anchors"],
                "missing_anchors": classification["missing_anchors"],
            },
            "job_archetype": {
                key: archetype[key]
                for key in ("archetype_id", "title", "client", "essence", "directions", "target_functions", "level_mapping", "keyword_groups")
            } if archetype else {},
            "consultant_input": consultant,
            "client_profile": client_profile_payload,
            "kb_graph_candidates": graph_pool,
            **learning,
            "stop_note_adjustments_summary": self._stop_note_adjustments_summary(learning.get("stop_note_adjustments") or []),
            "deterministic_fallback": fallback,
        }
        try:
            generated = self.service.llm.generate_search_strategy(payload)
        except Exception:
            generated = {}
        llm_v2_fragment = generated.pop("strategy_v2", None) if isinstance(generated, dict) else None
        plan = self._validate_model_strategy(generated, fallback, job, learning, max_queries)
        plan["generation"]["input_level"] = classification["input_level"]
        v2 = strategy_v2.build_strategy_v2(
            plan, classification, archetype=archetype, consultant=consultant, llm_fragment=llm_v2_fragment,
            profile_match=knowledge_base.profile_matched_info(profile_match),
            graph_pool=graph_pool,
            restricted_rules=knowledge_base.restricted_negative_rules(restricted_info),
            negative_checklist=negative_checklist,
            canonical_position=payload["canonical_position"],
            skill_ontology=skill_ontology,
            level_hit=level_hit,
            profile_context=(
                knowledge_base.profile_context(profile_match["profile"])
                if profile_match and isinstance(profile_match.get("profile"), dict)
                else {}
            ),
            learning=learning,
        )
        consultant_constraints = _lock_consultant_constraints(plan, v2, revision_evidence, locked_items)
        strategy_v2.refresh_consultant_judgement(v2)
        v2_ok, v2_errors = strategy_v2.validate_strategy_v2(v2)
        constraint_errors = _locked_constraint_conflicts(plan, v2, consultant_constraints)
        if constraint_errors:
            v2_ok = False
            v2_errors = [*v2_errors, *constraint_errors]
        # S4-3c-4（N6）：策略全要素消费检查 —— 对照命中原型的种子要素清单（T1/T2/T3 各层
        # 公司池、地点策略、排除规则、有效关键词组）核对 strategy_v2 是否全部消费，未使用项
        # 显式列出供顾问确认页展示；种子未命中（无原型岗位）coverage_report=None 留痕不算缺失。
        coverage_report = strategy_v2.build_coverage_report(archetype, v2)
        v2["coverage_report"] = coverage_report
        if coverage_report:
            existing_trace = v2.get("classification_trace")
            v2["classification_trace"] = [
                *(existing_trace if isinstance(existing_trace, list) else []),
                f"N6 要素消费检查：消费 {coverage_report['consumed_count']}/{coverage_report['element_count']} 项"
                + (
                    f"，未使用：{'、'.join(item['element'] for item in coverage_report['unused'])}"
                    if coverage_report["unused"]
                    else "，种子要素全部消费"
                ),
            ]
        result: dict[str, Any] = {
            "summary": "已基于岗位事实、客户画像、岗位原型和历史反馈生成寻访策略，并补齐资深顾问判断简报。",
            "strategy": plan,
            "input_level": classification["input_level"],
            "knowledge_health": knowledge_health,
            "references": self._job_reference(job),
        }
        if consultant_constraints:
            result["consultant_constraints"] = consultant_constraints
        if v2_ok:
            query_plan = query_builders.schedule_query_plan_v1(
                query_builders.compile_query_plan_v1(v2),
                self._query_plan_learning_metrics(int(job.get("id") or 0)),
            )
            golden_replay = strategy_v2.build_golden_candidate_replay(archetype, query_plan)
            result["strategy_v2"] = v2
            result["query_plan_v1"] = query_plan
            result["golden_candidate_replay_v1"] = golden_replay
            # 策略生成成功：消费本轮 stop_note_adjustments。
            self._apply_stop_note_adjustments(int(job.get("id") or 0), str(inputs.get("workflow_id") or ""))
            content = "# 多渠道寻访策略（strategy_v2）\n\n```json\n" + json.dumps(v2, ensure_ascii=False, indent=2) + "\n```"
            result["artifacts"] = [
                self._artifact(
                    "search_strategy", "多渠道寻访策略", content=content,
                    metadata={
                        "plan": plan,
                        "strategy_v2": v2,
                        "query_plan_v1": query_plan,
                        "golden_candidate_replay_v1": golden_replay,
                        "schema_version": strategy_v2.STRATEGY_V2_VERSION,
                        "coverage_report": coverage_report,
                    },
                )
            ]
        else:
            # 即使策略未通过 schema 校验，脚本能力也必须返回一个可审计的结果。
            # 这个产物是诊断记录，不是可执行策略；后置条件保持失败，避免无效策略
            # 被工作流误判为完成，同时把真实校验错误传给工作流和前端。
            error_payload = {
                "errors": v2_errors,
                "trace": classification["trace"][-12:],
                "input_level": classification["input_level"],
                "schema_version": strategy_v2.STRATEGY_V2_VERSION,
                "client": str(job.get("client") or ""),
                "job": str(job.get("title") or ""),
            }
            result["strategy_v2_error"] = error_payload
            readable_errors = v2_errors or ["未提供具体校验错误"]
            diagnostic_lines = [
                "# 多渠道寻访策略校验诊断",
                "",
                f"岗位：{job.get('client') or ''}｜{job.get('title') or ''}",
                f"输入分级：{classification['input_level']}",
                "",
                "## 校验错误",
                *[f"- {error}" for error in readable_errors],
                "",
                "## 执行追踪",
                *[f"- {item}" for item in error_payload["trace"]],
                "",
                "该记录仅用于审计和定位问题，未生成可执行寻访策略。",
            ]
            result["summary"] = "寻访策略未通过 strategy_v2 校验，已生成诊断记录。"
            result["artifacts"] = [
                self._artifact(
                    "search_strategy",
                    "寻访策略校验诊断",
                    content="\n".join(diagnostic_lines),
                    validation="failed",
                    metadata={
                        "diagnostic": True,
                        "schema_version": strategy_v2.STRATEGY_V2_VERSION,
                        "strategy_v2_error": error_payload,
                    },
                )
            ]
            result["postcondition"] = {
                "verified": False,
                "recoverable": True,
                "reason": "寻访策略未通过 strategy_v2 校验：" + "；".join(readable_errors[:4]),
            }
        return result

    def run_multi_channel_sourcing(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        job = self._job(context)
        strategy = self._workflow_strategy(inputs)
        approved_snapshot = self._approved_sourcing_snapshot(str(inputs.get("workflow_id") or ""))
        if approved_snapshot:
            strategy = {
                "strategy_summary": approved_snapshot.get("summary") or "",
                "channels": approved_snapshot.get("channels") or {},
                "target_companies": [
                    item.get("name") for item in (approved_snapshot.get("company_pool") or [])
                    if isinstance(item, dict) and item.get("name")
                ],
                "consultant_constraints": approved_snapshot.get("locked_constraints") or [],
            }
        try:
            preflight = self._run_json([self.python, str(MULTICHANNEL), "preflight", "--db", str(self.service.db_path), "--client", str(job["client"]), "--job", str(job["title"]), "--port", str(int(inputs.get("cdp_port") or 9223))], 90)
        except Exception as exc:
            preflight = {"ok": False, "status": "preflight_unavailable", "error": str(exc)[:1000]}
        channels = preflight.get("channels") or preflight.get("preflight") or {}
        ticket = {
            "client": job["client"], "job": job["title"], "preflight": preflight,
            "workflow_id": str(inputs.get("workflow_id") or ""),
            "target_count": int(approved_snapshot.get("target_count") or inputs.get("target_count") or self._target_count(inputs.get("objective")) or 10),
            "cdp_port": int(inputs.get("cdp_port") or 9223),
            "strategy": strategy,
            "strategy_snapshot": approved_snapshot,
            "strategy_hash": str(approved_snapshot.get("strategy_hash") or ""),
            "query_plan_v1": approved_snapshot.get("query_plan_v1") or {},
            "query_plan_hash": str(approved_snapshot.get("query_plan_hash") or ""),
            "required_result": {"verified": True, "channel_runs": [], "intake": {}, "audit": {}},
        }
        return {
            "summary": "渠道预检完成，已生成受约束寻访任务；渠道执行完成并读回前不会进入评估。",
            "references": self._job_reference(job), "external_action_executed": False,
            "external_request": ticket,
            "auto_execute_request": ticket if preflight.get("ok") is not False else None,
            "artifacts": [self._artifact("sourcing_ticket", "多渠道寻访执行任务", content=json.dumps(ticket, ensure_ascii=False, indent=2), validation="pending_execution", metadata={"channels": channels})],
        }

    def _workflow_strategy(self, inputs: dict[str, Any]) -> dict[str, Any]:
        workflow_id = str(inputs.get("workflow_id") or "")
        if not workflow_id:
            return {}
        conn = self.service._connect()
        try:
            row = conn.execute(
                "SELECT output_json FROM agent_workflow_steps WHERE workflow_id=? AND capability_id='search_strategy' AND status='completed' ORDER BY sequence DESC LIMIT 1",
                (workflow_id,),
            ).fetchone()
            output = _loads(row["output_json"], {}) if row else {}
            return output.get("strategy") if isinstance(output.get("strategy"), dict) else {}
        finally:
            conn.close()

    def _approved_sourcing_snapshot(self, workflow_id: str) -> dict[str, Any]:
        if not workflow_id:
            return {}
        conn = self.service._connect()
        try:
            row = conn.execute(
                """
                SELECT preflight_json FROM agent_approvals
                WHERE workflow_id=? AND action_type='multi_channel_sourcing' AND status='approved'
                ORDER BY id DESC LIMIT 1
                """,
                (workflow_id,),
            ).fetchone()
            preflight = _loads(row["preflight_json"], {}) if row else {}
            snapshot = preflight.get("strategy_snapshot") if isinstance(preflight.get("strategy_snapshot"), dict) else {}
            if snapshot and preflight.get("strategy_hash") == snapshot.get("strategy_hash"):
                return snapshot
            return {}
        finally:
            conn.close()

    @staticmethod
    def _target_count(objective: Any) -> int:
        match = re.search(r"(\d+)\s*(?:位|个|人)", str(objective or ""))
        return min(100, int(match.group(1))) if match else 0

    def run_job_publish_prepare(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        job = self._job(context)
        overrides = inputs.get("publish_fields") if isinstance(inputs.get("publish_fields"), dict) else {}
        draft = {
            "client_company": overrides.get("client_company") or job.get("client"),
            "job_title": overrides.get("job_title") or job.get("title"),
            "city_keyword": overrides.get("city_keyword") or job.get("position_location") or job.get("location"),
            "city_choice": overrides.get("city_choice") or job.get("position_location") or job.get("location"),
            "salary_low_k": overrides.get("salary_low_k"), "salary_high_k": overrides.get("salary_high_k"),
            "salary_months": overrides.get("salary_months") or 12,
            "job_category_keyword": overrides.get("job_category_keyword"), "job_category_choice": overrides.get("job_category_choice"),
            "industry_keyword": overrides.get("industry_keyword"), "industry_choice": overrides.get("industry_choice"),
            "work_year_keyword": overrides.get("work_year_keyword") or job.get("experience"),
            "work_year_choice": overrides.get("work_year_choice"), "work_year_low": overrides.get("work_year_low"), "work_year_high": overrides.get("work_year_high"),
            "education_choice": overrides.get("education_choice") or job.get("education"), "education_tongzhao": bool(overrides.get("education_tongzhao", False)),
            "private_job": bool(overrides.get("private_job", False)), "recruit_count": int(overrides.get("recruit_count") or job.get("headcount") or 1),
            "close_date": overrides.get("close_date") or job.get("deadline"),
            "description": overrides.get("description") or "\n".join(filter(None, [str(job.get("responsibilities") or ""), str(job.get("requirements") or job.get("summary") or "")])),
        }
        salary = str(job.get("salary") or "")
        numbers = [int(value) for value in re.findall(r"\d+", salary)]
        if not draft["salary_low_k"] and numbers:
            draft["salary_low_k"] = numbers[0]
        if not draft["salary_high_k"] and len(numbers) > 1:
            draft["salary_high_k"] = numbers[1]
        required = ["client_company", "job_title", "city_keyword", "salary_low_k", "salary_high_k", "job_category_choice", "industry_choice", "description", "close_date"]
        missing = [key for key in required if draft.get(key) in (None, "")]
        path = self._path("job_publish", f"{job['client']}-{job['title']}-draft", "json")
        path.write_text(json.dumps(draft, ensure_ascii=False, indent=2), encoding="utf-8")
        artifacts = [
            self._artifact(
                "job_publish_draft", "猎聘岗位发布草稿", file_path=path, mime_type="application/json",
                content=json.dumps(draft, ensure_ascii=False, indent=2),
                validation="needs_input" if missing else "passed",
                metadata={"client": job.get("client"), "job": job.get("title"), "job_id": job.get("id")},
            )
        ]
        result = {"summary": "岗位发布草稿已生成。" if not missing else "岗位发布草稿缺少关键字段，已阻塞正式发布。",
                  "references": self._job_reference(job), "draft": draft, "missing_inputs": missing,
                  "artifacts": artifacts}
        if missing:
            result["blocked"] = True
            return result
        prepare_log = self._path("job_publish", "liepin-publish-prepare-readback", "json")
        payload = self._run_json([
            self.python, str(LIEPIN_PUBLISH), "--mode", "prepare",
            "--port", str(int(inputs.get("cdp_port") or 9223)), "--draft", str(path), "--log", str(prepare_log),
        ], 180)
        failed = payload.get("ok") is False or str(payload.get("status") or "").lower() in {"blocked", "failed", "error"}
        artifacts.append(self._artifact(
            "job_publish_prepare_readback", "猎聘岗位发布预检读回", file_path=prepare_log,
            mime_type="application/json", content=json.dumps(payload, ensure_ascii=False, indent=2),
            validation="blocked" if failed else "passed",
            metadata={"draft_path": str(path), "client": job.get("client"), "job": job.get("title"), "job_id": job.get("id")},
        ))
        result["prepare_readback"] = payload
        result["summary"] = "岗位发布草稿已填入猎聘发布表单并完成读回预检。" if not failed else "猎聘岗位发布预检未通过，已阻塞正式发布。"
        if failed:
            result["blocked"] = True
            result["missing_inputs"] = ["修正猎聘发布预检问题"]
        return result

    def run_job_publish_execute(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        draft = self._dependency_file(inputs, "job_publish_draft")
        readback = self._dependency_file(inputs, "job_publish_prepare_readback")
        if not draft or not readback:
            return self._blocked("没有通过校验和读回的岗位发布草稿。", ["job_publish_draft", "job_publish_prepare_readback"])
        log = self._path("job_publish", "liepin-publish-readback", "json")
        payload = self._run_json([self.python, str(LIEPIN_PUBLISH), "--mode", "publish", "--confirm", "PUBLISH", "--port", str(int(inputs.get("cdp_port") or 9223)), "--draft", str(draft), "--log", str(log)], 180)
        verified = bool(payload.get("verified") or payload.get("status") in {"published", "submitted", "auditing"})
        if not verified:
            raise RuntimeError("猎聘发布动作未通过结果页或职位列表读回验证")
        job = self._job(context)
        conn = self.service._connect()
        try:
            if self._table(conn, "positions"):
                columns = {str(row[1]) for row in conn.execute("PRAGMA table_info(positions)").fetchall()}
                assignments = ["status='已发布'", "updated_at=datetime('now','localtime')"]
                if "liepin_status" in columns:
                    assignments.append("liepin_status='已发布/已验证'")
                if "liepin_published_at" in columns:
                    assignments.append("liepin_published_at=datetime('now','localtime')")
                if "liepin_verify_log" in columns:
                    assignments.append("liepin_verify_log=?")
                    conn.execute(f"UPDATE positions SET {','.join(assignments)} WHERE client=? AND title=?", (str(log), job["client"], job["title"]))
                else:
                    conn.execute(f"UPDATE positions SET {','.join(assignments)} WHERE client=? AND title=?", (job["client"], job["title"]))
                conn.commit()
        finally:
            conn.close()
        return {"summary": "猎聘岗位已发布并完成页面读回验证。", "external_action_executed": True,
                "external_result": payload, "artifacts": [self._artifact("external_action_receipt", "猎聘岗位发布回执", file_path=log, mime_type="application/json", content=json.dumps(payload, ensure_ascii=False, indent=2))]}

    def _dependency_file(self, inputs: dict[str, Any], artifact_type: str) -> Path | None:
        workflow_id = str(inputs.get("workflow_id") or "")
        if not workflow_id:
            return None
        conn = self.service._connect()
        try:
            row = conn.execute("SELECT file_path,validation_status FROM agent_artifacts WHERE workflow_id=? AND artifact_type=? ORDER BY id DESC LIMIT 1", (workflow_id, artifact_type)).fetchone()
            if row and row["file_path"] and row["validation_status"] == "passed" and Path(row["file_path"]).exists():
                return Path(row["file_path"])
            return None
        finally:
            conn.close()

    def run_resume_export(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        candidate = self._candidate(context)
        from docx import Document
        identity, position = candidate["identity"], candidate["position"]
        doc = Document()
        doc.add_heading(str(identity.get("name") or "候选人简历"), 0)
        for label, value in (("当前公司", identity.get("company")), ("当前职位", identity.get("title")), ("城市", identity.get("city")), ("学历", identity.get("education")), ("经验", identity.get("experience")), ("目标岗位", f"{position.get('client','')} / {position.get('job','')}")):
            doc.add_paragraph(f"{label}：{value or '不详'}")
        doc.add_heading("履历原始证据", level=1)
        profiles = candidate.get("source_profiles") or []
        if not profiles:
            doc.add_paragraph("暂无来源简历正文。")
        for profile in profiles:
            raw = _loads(profile.get("raw_json"), {})
            text = raw.get("profile_text") or raw.get("resume_text") or raw.get("text") or ""
            if text:
                doc.add_paragraph(str(text))
        path = self._path("resumes", f"{identity.get('name')}-{position.get('job')}-结构化简历", "docx")
        doc.save(path)
        return {"summary": "结构化简历 DOCX 已生成。", "references": self._candidate_reference(candidate),
                "artifacts": [self._artifact("resume_document", "结构化简历", file_path=path, mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")]}

    def run_matching_report(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        candidate = self._candidate(context)
        relation, identity, position = candidate["relation"], candidate["identity"], candidate["position"]
        assessment = self._latest_assessment(int(relation["job_candidate_id"]))
        criteria = assessment.get("criteria") or {}
        status_icon = {"met": "通过", "partial": "部分", "not_met": "不通过", "unknown": "待核验"}
        hard = [[item.get("criterion"), "；".join(item.get("evidence") or []) or item.get("reason") or "无证据", status_icon.get(item.get("status"), "待核验")] for item in criteria.get("hard_requirements") or []]
        star_map = {"met": 5, "partial": 3, "not_met": 1, "unknown": 2}
        matches = [{"duty": item.get("criterion"), "evidence": "；".join(item.get("evidence") or []) or "待核验", "stars": star_map.get(item.get("status"), 2)} for item in criteria.get("core_abilities") or []]
        risks = [{"level": "高" if "关键" in str(value) else "中", "title": str(value), "description": str(value), "verify": "在下一次沟通中核验并记录证据"} for value in assessment.get("risks") or assessment.get("gaps") or []]
        data = {
            "candidate": identity.get("name"), "company": position.get("client"), "position": position.get("job"),
            "hard_gates": hard, "responsibility_matches": matches,
            "bonus_items": [[value, "来自 ASA 当前评估", "通过"] for value in (assessment.get("strengths") or [])],
            "risks": risks, "scores": {"综合匹配": int(assessment.get("fit_score") or 0)},
            "total_score": int(assessment.get("fit_score") or 0),
            "interview_suggestions": [["证据核验", value] for value in (assessment.get("verification_questions") or [])],
            "verdict": str(assessment.get("recommendation") or "待复核"), "conclusion_summary": str(assessment.get("next_action") or "人工复核"),
        }
        data_path = self._path("reports", f"{identity.get('name')}-matching-data", "json")
        output = self._path("reports", f"人岗匹配-{position.get('client')}-{position.get('job')}-{identity.get('name')}", "docx")
        data_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        self._run([self.python, str(MATCHING_REPORT), "--data", f"@{data_path}", "--output", str(output)], 120)
        return {"summary": "人岗匹配分析报告已生成。", "references": self._candidate_reference(candidate),
                "artifacts": [self._artifact(
                    "matching_report", "人岗匹配分析报告", file_path=output,
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    metadata={
                        "assessment_id": assessment.get("id"), "job_candidate_id": relation.get("job_candidate_id"),
                        "person_id": relation.get("person_id"), "candidate_id": relation.get("source_candidate_id"),
                        "job_id": relation.get("job_id"), "client": position.get("client"), "job": position.get("job"),
                    },
                )]}

    def _s6_assessment_doc(self, job_candidate_id: int, job_id: Any) -> dict[str, Any] | None:
        """S6 判人评估 artifact（candidate_assessment，人×岗）；不存在/岗位号缺失 → None。"""
        try:
            jid = int(job_id or 0)
        except (TypeError, ValueError):
            return None
        if not jid:
            return None
        conn = self.service._connect()
        try:
            payload = candidate_assessment.get_assessment(conn, int(job_candidate_id), jid)
            return payload["assessment"] if payload else None
        finally:
            conn.close()

    def run_recommendation_report(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        candidate = self._candidate(context)
        relation, identity, position = candidate["relation"], candidate["identity"], candidate["position"]
        # S6-3：推荐报告强制引用判人评估块（PRD §1 ②）——无评估不许退回纯简历罗列。
        s6 = self._s6_assessment_doc(int(relation["job_candidate_id"]), relation.get("job_id"))
        if s6 is None:
            return self._blocked(
                "推荐报告必须引用判人评估结论：该人选还没有判人评估，请先在人选卡「评估」区生成判人评估"
                "（职业轨迹/在同龄人里的位置/动机与时机/需要核实的问题），再生成推荐报告。",
                ["先完成判人评估（candidate_assessment）"],
                self._candidate_reference(candidate),
            )
        assessment_block = candidate_assessment.report_reference_block(s6)
        assessment = self._latest_assessment(int(relation["job_candidate_id"]))
        if float(assessment.get("evidence_coverage") or 0) < 0.75:
            return self._blocked("证据覆盖不足，不能生成可发送推荐报告。", ["核验问题完成", "证据覆盖率>=0.75"], self._candidate_reference(candidate))
        if not JIASHI_TEMPLATE.exists():
            return self._blocked("嘉驰标准模板不存在。", [str(JIASHI_TEMPLATE)], self._candidate_reference(candidate))
        profile_summary = str((candidate.get("candidate_profile") or {}).get("profile_summary") or "履历职责待进一步结构化核验")
        data = {
            "customer": position.get("client"), "position": position.get("job"), "name": identity.get("name"),
            "current_location": identity.get("city") or "不详", "expected_location": "不详",
            "consultant_comments": assessment_block["lines"] + (assessment.get("strengths") or [])[:6],
            "education": [str(identity.get("education") or "不详")],
            "work_experience": [f"时间不详 {identity.get('company') or '公司待核验'}\n担任职位：{identity.get('title') or '职位待核验'}\n工作职责：{profile_summary}"],
            "project_experience": ["暂无明确项目经历"], "motivation": "待核验", "leaving_reason": "待核验",
        }
        data_path = self._path("reports", f"{identity.get('name')}-jiashi-data", "json")
        output = self._path("reports", f"嘉驰国际-{position.get('client')}-{position.get('job')}-{identity.get('name')}", "docx")
        data_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        self._run([self.python, str(JIASHI_REPORT), "--template", str(JIASHI_TEMPLATE), "--data", str(data_path), "--output", str(output)], 180)
        self._sanitize_docx_privacy(output)
        audit = self._run([self.python, str(JIASHI_AUDIT), str(output)], 120)
        return {"summary": "嘉驰推荐报告草稿已生成并通过模板审计，已引用判人评估块（评估只辅助判断，发送前请顾问复核）。", "references": self._candidate_reference(candidate),
                "artifacts": [self._artifact(
                    "recommendation_report", "嘉驰推荐报告", file_path=output,
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    content=audit.stdout[-4000:],
                    metadata={
                        "assessment_id": assessment.get("id"), "job_candidate_id": relation.get("job_candidate_id"),
                        "person_id": relation.get("person_id"), "candidate_id": relation.get("source_candidate_id"),
                        "job_id": relation.get("job_id"), "client": position.get("client"), "job": position.get("job"),
                        "attached_to_candidate": True, "external_submitted": False,
                        "s6_assessment": {
                            "artifact_id": f"candidate_assessment_{int(relation['job_candidate_id'])}_{int(relation.get('job_id') or 0)}",
                            "as_of": assessment_block["as_of"],
                            "assessor_version": assessment_block["assessor_version"],
                            "trajectory_verdict": assessment_block["trajectory_verdict"],
                            "percentile_band": assessment_block["percentile_band"],
                            "percentile_band_label": assessment_block["percentile_band_label"],
                            "reference_n": assessment_block["reference_n"],
                            "top_risks": assessment_block["top_risks"],
                            "risks_pending": assessment_block["risks_pending"],
                        },
                    },
                )]}

    @staticmethod
    def _sanitize_docx_privacy(path: Path) -> None:
        from docx import Document
        doc = Document(path)
        private = re.compile(r"(?:手机号|手机号码|联系电话|电话|微信|WeChat|微信号|邮箱|Email|E-mail)|(?:\+?86[-\s]?)?1[3-9]\d{9}|[A-Za-z0-9_.+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", re.I)
        for paragraph in doc.paragraphs:
            if private.search(paragraph.text or ""):
                paragraph.text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if private.search(paragraph.text or ""):
                            paragraph.text = ""
        doc.save(path)

    def run_interview_followup(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        return self._lifecycle_note(context, inputs, "interview_followup", "面试跟进", "interview_followup", "interview_note")

    def run_salary_negotiation(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        return self._lifecycle_note(context, inputs, "salary_negotiation", "谈薪跟进", "salary_negotiation", "salary_negotiation_note")

    def run_decision_coaching(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        return self._lifecycle_note(context, inputs, "decision_coaching", "候选人决策辅导", "decision_coaching", "decision_coaching")

    def run_onboarding_followup(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        return self._lifecycle_note(context, inputs, "onboarding_followup", "入职跟进", "onboarding_followup", "onboarding_note", days=7)

    def _lifecycle_note(self, context: dict[str, Any], inputs: dict[str, Any], event_type: str, label: str, task_type: str, artifact_type: str, days: int = 2) -> dict[str, Any]:
        candidate = self._candidate(context)
        objective = str(inputs.get("objective") or "").strip()
        if not objective:
            return self._blocked(f"{label}缺少业务事实。", ["objective"], self._candidate_reference(candidate))
        event_id = self._candidate_event(candidate, event_type, "recorded", objective, inputs)
        task_id = self._followup(candidate, task_type, objective, inputs, days)
        content = f"# {label}\n\n- 记录：{objective}\n- 事件 ID：{event_id}\n- 跟进任务 ID：{task_id or '未创建'}\n"
        return {"summary": f"{label}事实已记录，并创建后续任务。", "references": self._candidate_reference(candidate),
                "artifacts": [self._artifact(artifact_type, label, content=content, metadata={"event_id": event_id, "task_id": task_id})]}

    def run_salary_verification(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        candidate = self._candidate(context)
        data = inputs.get("salary_data") if isinstance(inputs.get("salary_data"), dict) else None
        if not data or not data.get("records"):
            return self._blocked("薪资核验需要结构化流水证据，未找到时不会生成虚假报告。", ["salary_data.records"], self._candidate_reference(candidate))
        data = {**data, "candidate_name": data.get("candidate_name") or candidate["identity"].get("name"), "report_date": data.get("report_date") or datetime.now().strftime("%Y-%m-%d")}
        data_path = self._path("salary", f"{candidate['identity'].get('name')}-salary-data", "json")
        out_dir = self.output_dir / "salary"
        data_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        proc = self._run([self.python, str(SALARY_REPORT), "--input", str(data_path), "--output-dir", str(out_dir)], 180)
        files = [Path(line.strip()) for line in proc.stdout.splitlines() if line.strip().endswith(".docx")]
        return {"summary": f"薪资证据报告已生成，共 {len(files)} 个文件。", "references": self._candidate_reference(candidate),
                "artifacts": [self._artifact("salary_report", path.stem, file_path=path, mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document") for path in files]}

    @staticmethod
    def _message_hash(job_candidate_id: Any, message: str) -> str:
        raw = f"{job_candidate_id}|{message}".encode("utf-8")
        return hashlib.sha256(raw).hexdigest()[:24]

    def _outreach_message(self, candidate: dict[str, Any], inputs: dict[str, Any]) -> str:
        explicit = " ".join(str(inputs.get("message") or "").split())
        if explicit:
            return explicit
        identity, position = candidate["identity"], candidate["position"]
        return (
            f"{identity.get('name') or '你好'}，你好。我这边有一个{position.get('client') or ''}"
            f"{position.get('job') or '岗位'}机会，和你{identity.get('company') or ''}"
            f"{identity.get('title') or ''}经历比较相关，方便了解一下吗？"
        )

    def _outreach_targets(self, context: dict[str, Any], inputs: dict[str, Any]) -> list[dict[str, Any]]:
        context_type = str(context.get("type") or "")
        if context_type == "candidate":
            return [self._candidate(context)]
        filters = context.get("filters") if isinstance(context.get("filters"), dict) else {}
        queue = str(inputs.get("queue") or filters.get("queue") or "待联系")
        limit = max(1, min(int(inputs.get("limit") or filters.get("limit") or 20), 20))
        inbox = self.service.get_flow_inbox(queue=queue, limit=limit)
        targets: list[dict[str, Any]] = []
        for item in inbox.get("items") or []:
            job_candidate_id = item.get("job_candidate_id")
            if not job_candidate_id:
                continue
            try:
                targets.append(build_candidate_context(self.service.db_path, int(job_candidate_id)))
            except Exception:
                continue
        return targets

    def run_outreach_prepare(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        targets = self._outreach_targets(context, inputs)
        if not targets:
            return self._blocked("当前目标或队列没有可触达的人选。", ["选择当前目标、当前队列或具体人选"])
        items: list[dict[str, Any]] = []
        skipped: list[dict[str, Any]] = []
        for candidate in targets[:20]:
            relation, identity, position = candidate["relation"], candidate["identity"], candidate["position"]
            if is_stopped(candidate):
                skipped.append({"job_candidate_id": relation["job_candidate_id"], "candidate": identity.get("name"), "reason": "关系已停止"})
                continue
            message = self._outreach_message(candidate, inputs)
            message_hash = self._message_hash(relation["job_candidate_id"], message)
            items.append({
                "job_candidate_id": relation["job_candidate_id"], "person_id": relation.get("person_id"),
                "job_id": relation.get("job_id"), "candidate": identity.get("name"),
                "company": identity.get("company"), "title": identity.get("title"),
                "client": position.get("client"), "job": position.get("job"), "channel": "猎聘职聊",
                "message": message, "message_hash": message_hash,
                "before": relation.get("clean_stage") or relation.get("raw_status") or "未触达",
                "after": "猎聘消息发送并读回后进入已触达/待回复",
                "status": "pending",
            })
        if not items:
            return self._blocked("本批人选均不可触达。", ["选择未停止且有明确人岗关系的人选"], [])
        payload = {
            "version": "1.0", "prepared_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "batch_limit": 20, "item_count": len(items), "items": items, "skipped": skipped,
        }
        path = self._path("outreach", "猎聘触达锁定草稿", "json")
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        references = [
            {"type": "candidate", "id": item["job_candidate_id"], "label": item["candidate"], "subtitle": f"{item['client']} / {item['job']}"}
            for item in items[:8]
        ]
        return {
            "summary": f"已锁定 {len(items)} 条猎聘触达草稿，正式发送前需要批量确认。",
            "references": references,
            "outreach_draft_batch": payload,
            "artifacts": [self._artifact(
                "outreach_draft_batch", "猎聘触达锁定草稿", file_path=path, mime_type="application/json",
                content=json.dumps(payload, ensure_ascii=False, indent=2),
                metadata={"item_count": len(items), "batch_limit": 20},
            )],
        }

    def _sent_message_exists(self, job_candidate_id: int, message_hash: str) -> bool:
        conn = self.service._connect()
        try:
            row = conn.execute(
                """
                SELECT 1 FROM candidate_events
                WHERE job_candidate_id=? AND event_type='candidate_outreach'
                  AND event_status='sent_verified' AND raw_json LIKE ?
                LIMIT 1
                """,
                (int(job_candidate_id), f"%{message_hash}%"),
            ).fetchone()
            return row is not None
        finally:
            conn.close()

    def _load_outreach_batch(self, inputs: dict[str, Any]) -> dict[str, Any] | None:
        path = self._dependency_file(inputs, "outreach_draft_batch")
        if not path:
            return None
        return _loads(path.read_text(encoding="utf-8"), {})

    def run_outreach_execute(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        batch = self._load_outreach_batch(inputs)
        if not batch:
            return self._blocked("没有已锁定并通过审批的触达草稿。", ["outreach_draft_batch"])
        results: list[dict[str, Any]] = []
        references: list[dict[str, Any]] = []
        for item in (batch.get("items") or [])[:20]:
            job_candidate_id = int(item.get("job_candidate_id") or 0)
            message = str(item.get("message") or "").strip()
            message_hash = str(item.get("message_hash") or self._message_hash(job_candidate_id, message))
            if not job_candidate_id or not message:
                results.append({**item, "status": "failed", "error": "缺少人岗关系或文案"})
                continue
            if self._sent_message_exists(job_candidate_id, message_hash):
                results.append({**item, "status": "skipped", "reason": "同一锁定文案已发送并验证"})
                continue
            candidate = build_candidate_context(self.service.db_path, job_candidate_id)
            if is_stopped(candidate):
                results.append({**item, "status": "skipped", "reason": "关系已停止"})
                continue
            identity, position = candidate["identity"], candidate["position"]
            base = [
                self.python, str(LIEPIN_OUTREACH), "--port", str(int(inputs.get("cdp_port") or 9223)),
                "--candidate", str(identity.get("name")), "--message", message, "--check", str(position.get("job")),
            ]
            try:
                dry = self._run_json(base, 90)
                if dry.get("status") != "dry_run_ok":
                    raise RuntimeError("猎聘触达预检未通过")
                sent = self._run_json(base + ["--send"], 120)
                if sent.get("status") != "sent_verified":
                    raise RuntimeError("猎聘消息点击后未通过会话读回验证")
                event_id = self._candidate_event(
                    candidate, "candidate_outreach", "sent_verified", f"猎聘触达已验证：{message[:80]}",
                    {**inputs, "message_hash": message_hash, "locked_message": message, "channel_result": sent},
                )
                conn = self.service._connect()
                try:
                    conn.execute(
                        "UPDATE job_candidates SET raw_status='contacted',clean_stage='已触达',flow_bucket='已触达/待回复',updated_at=datetime('now','localtime') WHERE id=?",
                        (job_candidate_id,),
                    )
                    conn.commit()
                finally:
                    conn.close()
                result_item = {**item, "status": "sent_verified", "event_id": event_id, "dry_run": dry, "receipt": sent}
            except Exception as exc:
                result_item = {**item, "status": "failed", "error": str(exc)[:1000]}
            results.append(result_item)
            references.append({"type": "candidate", "id": job_candidate_id, "label": item.get("candidate"), "subtitle": f"{item.get('client','')} / {item.get('job','')}"})
        success = [item for item in results if item.get("status") in {"sent_verified", "skipped"}]
        failed = [item for item in results if item.get("status") == "failed"]
        if not success and failed:
            raise RuntimeError("本批猎聘触达全部失败：" + "；".join(str(item.get("error") or "") for item in failed[:3]))
        payload = {
            "verified": bool(success), "batch_status": "partial_failed" if failed else "completed",
            "sent_count": len([item for item in results if item.get("status") == "sent_verified"]),
            "skipped_count": len([item for item in results if item.get("status") == "skipped"]),
            "failed_count": len(failed), "items": results,
        }
        return {
            "summary": f"猎聘触达完成：发送 {payload['sent_count']} 人，跳过 {payload['skipped_count']} 人，失败 {payload['failed_count']} 人。",
            "references": references[:8], "external_action_executed": True, "external_result": payload,
            "artifacts": [self._artifact(
                "external_action_receipt", "猎聘触达批量回执", content=json.dumps(payload, ensure_ascii=False, indent=2),
                metadata={"sent_count": payload["sent_count"], "failed_count": payload["failed_count"]},
            )],
        }

    def run_client_recommendation(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        candidate = self._candidate(context)
        report = self._dependency_file(inputs, "recommendation_report")
        if not report:
            return self._blocked("客户推荐缺少通过审计的推荐报告。", ["recommendation_report"], self._candidate_reference(candidate))
        request = {"candidate": candidate["identity"].get("name"), "client": candidate["position"].get("client"), "job": candidate["position"].get("job"), "report": str(report), "channel": inputs.get("channel") or "manual_client_channel"}
        return {"summary": "客户推荐材料已锁定，等待指定客户渠道完成发送并读回。", "references": self._candidate_reference(candidate),
                "external_action_executed": False, "external_request": request,
                "artifacts": [self._artifact("external_action_ticket", "客户推荐执行任务", content=json.dumps(request, ensure_ascii=False, indent=2), validation="pending_execution")]}

    def run_offer_confirmation(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        candidate = self._candidate(context)
        terms = inputs.get("offer_terms") if isinstance(inputs.get("offer_terms"), dict) else None
        if not terms:
            return self._blocked("Offer 确认缺少明确条件，不改变候选人阶段。", ["offer_terms"], self._candidate_reference(candidate))
        event_id = self._candidate_event(candidate, "offer_confirmation", "confirmed", "Offer 条件已人工确认", {**inputs, "offer_terms": terms})
        return {"summary": "Offer 条件已记录；未代表候选人已经接受。", "references": self._candidate_reference(candidate), "external_action_executed": True,
                "artifacts": [self._artifact("offer_confirmation", "Offer 条件确认", content=json.dumps(terms, ensure_ascii=False, indent=2), metadata={"event_id": event_id})]}

    def run_project_retrospective(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        if context.get("type") == "candidate":
            candidate = self._candidate(context)
            refs = self._candidate_reference(candidate)
            relation_id = candidate["relation"]["job_candidate_id"]
            conn = self.service._connect()
            try:
                events = [_row(row) for row in conn.execute("SELECT event_type,event_status,event_time,summary FROM candidate_events WHERE job_candidate_id=? ORDER BY event_time,id", (relation_id,)).fetchall()]
            finally:
                conn.close()
            content = "# 项目复盘\n\n" + "\n".join(f"- {item.get('event_time') or ''} · {item.get('event_type')} · {item.get('summary') or ''}" for item in events)
        else:
            job = self._job(context)
            refs = self._job_reference(job)
            conn = self.service._connect()
            try:
                counts = {row["clean_stage"] or "未整理": int(row["total"]) for row in conn.execute("SELECT clean_stage,COUNT(*) total FROM job_candidates WHERE job_id=? GROUP BY clean_stage", (job["id"],)).fetchall()}
            finally:
                conn.close()
            content = "# 项目复盘\n\n```json\n" + json.dumps(counts, ensure_ascii=False, indent=2) + "\n```"
        return {"summary": "已基于 v3 事件与漏斗生成项目复盘。", "references": refs,
                "artifacts": [self._artifact("project_retrospective", "项目复盘", content=content)]}

    def run_memory_capture(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        content = str(inputs.get("confirmed_memory") or "").strip()
        if not content:
            return self._blocked("长期记忆只沉淀明确确认的信息。", ["confirmed_memory"])
        scope = context.get("type") if context.get("type") in {"job", "candidate"} else "global"
        result = self.service.store_memory(scope, str(context.get("id") or ""), "workflow_outcome", content, "agent_workflow", str(inputs.get("workflow_id") or ""), 1.0)
        return {"summary": "经确认的业务经验已写入 ASA 长期记忆。", "memory": result}

    def run_identity_merge_preflight(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        candidate = self._candidate(context)
        other_id = inputs.get("other_job_candidate_id")
        if not other_id:
            return self._blocked("身份合并预检需要另一条候选人关系。", ["other_job_candidate_id"], self._candidate_reference(candidate))
        other = build_candidate_context(self.service.db_path, int(other_id))
        comparison = {"left": candidate["identity"], "right": other["identity"], "same_name": candidate["identity"].get("name") == other["identity"].get("name"), "same_company": candidate["identity"].get("company") == other["identity"].get("company"), "same_title": candidate["identity"].get("title") == other["identity"].get("title")}
        comparison["allowed"] = comparison["same_name"] and (comparison["same_company"] or comparison["same_title"])
        return {"summary": "身份对比完成；该步骤不会执行合并。", "comparison": comparison,
                "references": self._candidate_reference(candidate) + self._candidate_reference(other),
                "artifacts": [self._artifact("identity_comparison", "候选人身份对比", content=json.dumps(comparison, ensure_ascii=False, indent=2), validation="passed" if comparison["allowed"] else "blocked")]}
