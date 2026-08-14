from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from . import candidate_assessment, knowledge_base, negative_rules, query_builders, strategy_v2
from .capability_runtime_base import (
    JIASHI_AUDIT,
    JIASHI_REPORT,
    JIASHI_TEMPLATE,
    LIEPIN_OUTREACH,
    LIEPIN_PUBLISH,
    MATCHING_REPORT,
    MULTICHANNEL,
    SALARY_REPORT,
    _loads,
    _lock_consultant_constraints,
    _locked_constraint_conflicts,
    _revision_consultant_evidence,
    _row,
)
from .context import build_candidate_context
from .policy import is_stopped


_JIASHI_RESUME_LABELS = {
    "推荐岗位": "position",
    "所属中心": "customer",
    "姓名": "name",
    "出生年月": "birth",
    "性别": "gender",
    "婚育": "marital_status",
    "工作地址": "current_location",
    "综合年薪": "current_salary",
    "期望薪酬": "expected_salary",
    "推荐理由": "consultant_comments",
    "教育背景": "education",
    "工作经历": "work_experience",
    "项目经历": "project_experience",
}
_JIASHI_RESUME_LIST_FIELDS = {"consultant_comments", "education", "work_experience", "project_experience"}


def _parse_jiashi_resume_fields(text: str) -> dict[str, Any]:
    """解析嘉驰报告结构的上传简历文本（推荐岗位/所属中心/姓名/…标签行）。

    标量字段取标签行“标签：值”的值；段落字段（推荐理由/教育背景/工作经历/
    项目经历）收集标签行之后的连续非标签行。
    """
    fields: dict[str, Any] = {}
    current = ""
    for raw_line in str(text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        match = re.match(r"^([一-龥]{2,6})\s*[：:]\s*(.*)$", line)
        if match and match.group(1) in _JIASHI_RESUME_LABELS:
            current = _JIASHI_RESUME_LABELS[match.group(1)]
            value = match.group(2).strip()
            if current in _JIASHI_RESUME_LIST_FIELDS:
                fields.setdefault(current, [])
                if value:
                    fields[current].append(value)
            elif value:
                fields[current] = value
            continue
        if current in _JIASHI_RESUME_LIST_FIELDS:
            fields.setdefault(current, []).append(line)
    return fields


class RunnerDeliveryMixin:
    """交付能力：策略学习/停止笔记、多渠道寻访、岗位发布、报告与候选人触达。"""

    def _strategy_learning_context(self, job: dict[str, Any]) -> dict[str, Any]:
        conn = self.service._connect()
        try:
            experiments = [
                _row(row)
                for row in conn.execute(
                    """
                    SELECT channel,query,result_count,viewed_count,extracted_count,
                           recommended_count,reply_count,positive_reply_count,noise_notes,status,updated_at
                    FROM search_experiments
                    WHERE client=? AND position=?
                    ORDER BY datetime(COALESCE(updated_at,run_time,created_at)) DESC,id DESC LIMIT 24
                    """,
                    (job["client"], job["title"]),
                ).fetchall()
            ] if self._table(conn, "search_experiments") else []
            correction = _row(conn.execute(
                "SELECT * FROM strategy_corrections WHERE client=? AND position=? ORDER BY id DESC LIMIT 1",
                (job["client"], job["title"]),
            ).fetchone()) if self._table(conn, "strategy_corrections") else {}
            business_outcomes = [
                _row(row)
                for row in conn.execute(
                    """
                    SELECT sa.channel,sa.source_query,
                           COUNT(DISTINCT sa.job_candidate_id) AS attributed_candidates,
                           COUNT(sf.id) AS signal_count,ROUND(COALESCE(SUM(sf.weight),0),2) AS experience_score,
                           SUM(sf.signal_type='review_pass') AS review_pass,
                           SUM(sf.signal_type='contacted') AS contacted,
                           SUM(sf.signal_type='recommended') AS recommended,
                           SUM(sf.signal_type='stopped') AS stopped,
                           SUM(sf.signal_type IN ('client_approved','client_interview','client_offer','client_hired')) AS client_positive,
                           SUM(sf.signal_type='client_rejected') AS client_rejected
                    FROM agent_sourcing_attributions sa
                    LEFT JOIN agent_sourcing_feedback sf ON sf.attribution_id=sa.id
                    WHERE sa.job_id=?
                    GROUP BY sa.channel,sa.source_query
                    ORDER BY experience_score DESC,signal_count DESC,sa.id DESC LIMIT 30
                    """,
                    (job["id"],),
                ).fetchall()
            ] if self._table(conn, "agent_sourcing_attributions") and self._table(conn, "agent_sourcing_feedback") else []
            accepted_adjustments: list[dict[str, Any]] = []
            if self._table(conn, "agent_sourcing_adjustments"):
                accepted_adjustments = [
                    _row(row)
                    for row in conn.execute(
                        """
                        SELECT a.*, p.display_name AS candidate_name
                        FROM agent_sourcing_adjustments a
                        LEFT JOIN job_candidates jc ON jc.id=a.candidate_id
                        LEFT JOIN people p ON p.id=jc.person_id
                        WHERE a.job_id=? AND a.status='accepted'
                        ORDER BY a.id
                        """,
                        (job["id"],),
                    ).fetchall()
                ]
        finally:
            conn.close()
        memories = self.service.search_memories(
            f"{job['client']} {job['title']} 寻访关键词 搜索效果",
            context_type="job", context_id=job["id"], client=str(job["client"]), job=str(job["title"]), limit=8,
        )
        return {
            "historical_experiments": experiments,
            "explicit_corrections": correction,
            "business_outcomes": business_outcomes,
            "approved_memories": memories.get("memories") or [] if memories.get("mode") == "active" else [],
            "memory_mode": memories.get("mode") or "off",
            "memory_hits": len(memories.get("memories") or []),
            "stop_note_adjustments": accepted_adjustments,
        }

    _STOP_NOTE_ADJUSTMENT_LABELS: dict[str, str] = {
        "add_keyword": "补充关键词",
        "remove_keyword": "移除关键词",
        "exclude_company": "排除公司",
        "add_company": "补充公司",
        "add_filter": "添加过滤",
        "adjust_salary_range": "调整薪资区间",
    }

    @staticmethod
    def _stop_note_adjustments_summary(adjustments: list[dict[str, Any]]) -> str:
        """把顾问已采纳的调整格式化为策略 prompt 可消费的文本摘要。"""
        # 延迟导入：组合类在 capability_runtime 模块完成定义，避免模块级循环依赖。
        from .capability_runtime import RecruitingCapabilityRuntime
        if not adjustments:
            return ""
        parts: list[str] = []
        for item in adjustments:
            if not isinstance(item, dict):
                continue
            adjust_type = str(item.get("adjust_type") or item.get("type") or "").strip()
            value = str(item.get("value") or "").strip()
            if not adjust_type or not value:
                continue
            label = RecruitingCapabilityRuntime._STOP_NOTE_ADJUSTMENT_LABELS.get(adjust_type, adjust_type)
            rationale = str(item.get("rationale") or "")[:60]
            parts.append(f"{label}：{value}" + (f"（来源：{rationale}）" if rationale else ""))
        return "；".join(parts)

    @staticmethod
    def _normalize_strategy_entries(values: Any) -> list[dict[str, str]]:
        rows: list[dict[str, str]] = []
        for index, value in enumerate(values if isinstance(values, list) else []):
            item = value if isinstance(value, dict) else {"query": value}
            query = " ".join(str(item.get("query") or "").split())
            if not query:
                continue
            rows.append({
                "round": str(item.get("round") or f"model-{index + 1}"),
                "query": query,
                "purpose": str(item.get("purpose") or "大模型生成的岗位寻访查询"),
                "evidence": str(item.get("evidence") or "岗位事实"),
            })
        return rows

    @staticmethod
    def _strategy_term_grounded(term: str, canonical_text: str) -> bool:
        # 延迟导入：组合类在 capability_runtime 模块完成定义，避免模块级循环依赖。
        from .capability_runtime import RecruitingCapabilityRuntime
        canonical = re.sub(r"\s+", "", canonical_text.lower())
        normalized = re.sub(r"\s+", "", term.lower())
        if normalized and normalized in canonical:
            return True
        parts = [value for value in re.split(r"[\s/、,，;；]+", term) if value]
        if len(parts) > 1:
            return all(RecruitingCapabilityRuntime._strategy_term_grounded(value, canonical_text) for value in parts)
        ascii_tokens = re.findall(r"[A-Za-z][A-Za-z0-9+-]*", term)
        if ascii_tokens and not all(token.lower() in canonical for token in ascii_tokens):
            return False
        chinese = "".join(re.findall(r"[\u4e00-\u9fff]", term))
        if not chinese:
            return bool(ascii_tokens)
        if len(chinese) <= 2:
            return chinese.lower() in canonical
        bigrams = [chinese[index:index + 2] for index in range(len(chinese) - 1)]
        matched = sum(value.lower() in canonical for value in bigrams)
        return matched / max(1, len(bigrams)) >= 0.5

    def _validate_model_strategy(
        self, raw: dict[str, Any], fallback: dict[str, Any], job: dict[str, Any], learning: dict[str, Any], max_queries: int,
    ) -> dict[str, Any]:
        profile = job.get("profile") or {}
        canonical_text = " ".join(
            str(value or "")
            for value in (job.get("title"), job.get("responsibilities"), job.get("requirements"), job.get("hard_requirements"), job.get("exclusions"))
        ).lower()
        legacy_terms = [
            str(value).strip()
            for key in ("ability_keywords_json", "search_keywords_json")
            for value in _loads(profile.get(key), [])
            if str(value).strip()
        ]
        unsupported = {
            term for term in legacy_terms
            if len(re.sub(r"\W+", "", term)) >= 3 and not self._strategy_term_grounded(term, canonical_text)
        }
        raw_channels = raw.get("channels") if isinstance(raw.get("channels"), dict) else {}
        fallback_channels = fallback.get("channels") if isinstance(fallback.get("channels"), dict) else {}
        removed: list[str] = []
        channels: dict[str, list[dict[str, str]]] = {}
        for channel in ("liepin", "xsaas"):
            accepted: list[dict[str, str]] = []
            seen: set[str] = set()
            for item in [*self._normalize_strategy_entries(raw_channels.get(channel)), *self._normalize_strategy_entries(fallback_channels.get(channel))]:
                normalized = re.sub(r"\s+", "", item["query"].lower())
                bad = next((term for term in unsupported if re.sub(r"\s+", "", term.lower()) in normalized), "")
                if bad:
                    removed.append(f"{item['query']}（无岗位依据：{bad}）")
                    continue
                if normalized in seen:
                    continue
                seen.add(normalized)
                accepted.append(item)
                if len(accepted) >= max_queries:
                    break
            channels[channel] = accepted
        model_used = any(self._normalize_strategy_entries(raw_channels.get(channel)) for channel in ("liepin", "xsaas"))
        return {
            **{key: value for key, value in fallback.items() if key not in {"channels"}},
            "channels": channels,
            "target_companies": raw.get("target_companies") or fallback.get("target_companies") or [],
            "strategy_summary": str(raw.get("strategy_summary") or "围绕岗位硬门槛、应用场景与目标公司分层寻访。"),
            "learning_notes": raw.get("learning_notes") if isinstance(raw.get("learning_notes"), list) else [],
            "generation": {
                "mode": "llm" if model_used else "deterministic_fallback",
                "model": self.service.llm.model,
                "consultant_mode": "senior_consultant_v1",
                "memory_mode": learning["memory_mode"],
                "memory_hits": learning["memory_hits"],
                "experiment_count": len(learning["historical_experiments"]),
                "removed_unsupported_queries": removed,
            },
        }

    def _capture_search_learning(self, client: str, job: str, queries: list[Any]) -> dict[str, Any]:
        query_values = [
            " ".join(str((item or {}).get("query") if isinstance(item, dict) else item or "").split())
            for item in queries
        ]
        query_values = list(dict.fromkeys(value for value in query_values if value))
        if not query_values:
            return {"stored_memories": 0, "queries": 0}
        conn = self.service._connect()
        try:
            placeholders = ",".join("?" for _ in query_values)
            rows = conn.execute(
                f"""
                SELECT channel,query,result_count,viewed_count,recommended_count,positive_reply_count,noise_notes
                FROM search_experiments
                WHERE client=? AND position=? AND query IN ({placeholders})
                ORDER BY datetime(COALESCE(updated_at,run_time,created_at)) DESC,id DESC
                """,
                [client, job, *query_values],
            ).fetchall()
        finally:
            conn.close()
        stored = 0
        seen: set[tuple[str, str]] = set()
        for row in rows:
            key = (str(row["channel"] or ""), str(row["query"] or ""))
            if key in seen:
                continue
            seen.add(key)
            recommended = int(row["recommended_count"] or 0)
            viewed = int(row["viewed_count"] or 0)
            outcome = "有效" if recommended > 0 else "待降权"
            content = (
                f"{client}/{job} 搜索经验：{key[0] or '未知渠道'} 关键词“{key[1]}”{outcome}；"
                f"查看 {viewed}，推荐 {recommended}，正向回复 {int(row['positive_reply_count'] or 0)}。"
            )
            if row["noise_notes"]:
                content += f" 噪音：{str(row['noise_notes'])[:160]}"
            self.service.store_memory(
                scope_type="job", scope_id=str(self._job_id(client, job)), memory_type="search_outcome",
                content=content, source_type="search_experiment", source_id=f"{key[0]}:{key[1]}",
                confidence=0.9 if recommended > 0 else 0.72,
            )
            stored += 1
        return {"stored_memories": stored, "queries": len(query_values)}

    def _query_plan_learning_metrics(self, job_id: int) -> list[dict[str, Any]]:
        """Aggregate prior marginal yield, overlap and downstream business feedback per query."""
        conn = self.service._connect()
        try:
            rows = conn.execute(
                """
                SELECT qc.channel,qc.query,COUNT(DISTINCT qc.run_id) AS runs,
                       COUNT(r.id) AS raw_occurrences,
                       COUNT(DISTINCT r.channel || ':' || COALESCE(NULLIF(r.source_candidate_id,''),r.identity_key)) AS unique_identities,
                       COUNT(DISTINCT CASE WHEN r.detail_status='complete'
                           THEN r.channel || ':' || COALESCE(NULLIF(r.source_candidate_id,''),r.identity_key) END) AS detail_complete,
                       COUNT(DISTINCT CASE WHEN r.job_candidate_id IS NOT NULL
                           THEN r.job_candidate_id END) AS intake_count
                FROM agent_sourcing_query_cells qc
                LEFT JOIN agent_candidate_recalls r
                  ON r.run_id=qc.run_id AND r.query_cell_id=qc.cell_id
                WHERE qc.job_id=?
                GROUP BY qc.channel,qc.query
                """,
                (job_id,),
            ).fetchall()
            feedback_rows = conn.execute(
                """
                SELECT sa.channel,sa.source_query,COALESCE(SUM(sf.weight),0) AS business_score
                FROM agent_sourcing_attributions sa
                LEFT JOIN agent_sourcing_feedback sf ON sf.attribution_id=sa.id
                WHERE sa.job_id=?
                GROUP BY sa.channel,sa.source_query
                """,
                (job_id,),
            ).fetchall()
        finally:
            conn.close()
        feedback = {
            (str(row["channel"] or ""), " ".join(str(row["source_query"] or "").split()).casefold()): float(row["business_score"] or 0)
            for row in feedback_rows
        }
        metrics: list[dict[str, Any]] = []
        for row in rows:
            runs = max(1, int(row["runs"] or 0))
            raw = max(0, int(row["raw_occurrences"] or 0))
            unique = max(0, int(row["unique_identities"] or 0))
            complete = max(0, int(row["detail_complete"] or 0))
            intake = max(0, int(row["intake_count"] or 0))
            channel = str(row["channel"] or "")
            query = str(row["query"] or "")
            metrics.append({
                "channel": channel,
                "query": query,
                "runs": runs,
                "raw_occurrences": raw,
                "unique_identities": unique,
                "detail_complete": complete,
                "intake_count": intake,
                "unique_yield_per_run": round(unique / runs, 4),
                "overlap_rate": round(1 - unique / raw, 4) if raw else 0.0,
                "business_score": feedback.get((channel, " ".join(query.split()).casefold()), 0.0),
            })
        return metrics

    def _job_id(self, client: str, job: str) -> int:
        conn = self.service._connect()
        try:
            row = conn.execute(
                "SELECT j.id FROM jobs j JOIN clients c ON c.id=j.client_id WHERE c.name=? AND j.title=? ORDER BY j.id LIMIT 1",
                (client, job),
            ).fetchone()
            return int(row["id"]) if row else 0
        finally:
            conn.close()

    def run_search_strategy(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        job = self._job(context)
        knowledge_health = strategy_v2.knowledge_base_health()
        max_queries = min(12, int(inputs.get("max_queries") or 6))
        command = [self.python, str(MULTICHANNEL), "plan", "--db", str(self.service.db_path), "--client", str(job["client"]), "--job", str(job["title"]), "--max-queries", str(max_queries)]
        try:
            fallback = self._run_json(command, 60)
        except Exception:
            profile = job.get("profile") or {}
            fallback = {
                "client": job["client"], "job": job["title"],
                "queries": _loads(profile.get("search_keywords_json"), [])[:6],
                "target_companies": _loads(profile.get("target_companies_json"), [])[:20],
                "exclusions": _loads(profile.get("exclusion_tags_json"), [])[:20],
                "source": "v3_fallback",
            }
        learning = self._strategy_learning_context(job)
        sourcing_adjustment_input = [
            {
                "id": int(item["id"]),
                "adjust_type": str(item.get("adjust_type") or ""),
                "value": str(item.get("value") or ""),
                "rationale": str(item.get("rationale") or ""),
            }
            for item in (learning.get("stop_note_adjustments") or [])
            if isinstance(item, dict) and item.get("id") is not None
        ]
        # S4-1：策略生成前先定级（L1/L2/L3）并盘点四锚点；顾问在 Copilot 侧的
        # 放行/锚点回复经工作流上下文 strategy_clarification 传入（ consultant_override /
        # consultant_answers ），推断项按 PRD §1 保持 inferred:true + confidence。
        clarification = context.get("strategy_clarification") if isinstance(context.get("strategy_clarification"), dict) else {}
        revision_evidence = _revision_consultant_evidence(context)
        understanding = context.get("intent_understanding") if isinstance(context.get("intent_understanding"), dict) else {}
        locked_items = [
            {"quote": str(item.get("quote") or "").strip(), "kind": str(item.get("kind") or "other")}
            for item in (understanding.get("constraints") or [])
            if isinstance(item, dict) and str(item.get("quote") or "").strip()
        ]
        for quote in context.get("locked_constraints") or []:
            normalized_quote = str(quote or "").strip()
            if normalized_quote and not any(item["quote"] == normalized_quote for item in locked_items):
                locked_items.append({"quote": normalized_quote, "kind": "other"})
        consultant_answers = "；".join(
            value
            for value in (
                str(clarification.get("consultant_answers") or "").strip(),
                revision_evidence,
                "；".join(item["quote"] for item in locked_items),
            )
            if value
        )
        consultant = {
            "consultant_override": bool(clarification.get("consultant_override")),
            "consultant_answers": consultant_answers,
        }
        archetype, archetype_trace = strategy_v2.match_job_archetype(job.get("client"), job.get("title"))
        classification = strategy_v2.classify_strategy_input(
            job, archetype=archetype, consultant_answers=consultant["consultant_answers"]
        )
        classification["trace"] = [*archetype_trace, *classification["trace"]]
        # S4-2：知识库消费 —— 客户画像挂载（精确/别名/模糊需确认）、公司图谱召回、
        # restricted 白名单约束。画像与图谱进 LLM 输入；restricted 键值不进任何生成
        # 上下文，仅由运行时并入 negative_rules（source=restricted_client）。
        profile_match, profile_trace = knowledge_base.match_client_profile(job.get("client"))
        restricted_info, restricted_trace = knowledge_base.load_restricted_constraints(job.get("client"))
        graph, graph_trace = knowledge_base.load_company_graph()
        # 知识飞轮二期：公司校准覆盖层（company_calibrations 表，仅 status='calibrated'）
        # 合并进图谱后再推导公司池；命中公司标注 source=consultant_calibrated 并留痕。
        # 覆盖层每次运行直接读库（表很小，不做进程内缓存，校准提交即时生效）；
        # 覆盖层为空/加载失败时不改图谱、不动 trace，输出与现状逐字节一致。
        calibration_overlay, _overlay_load_trace = knowledge_base.load_calibration_overlay(self.service.db_path)
        if calibration_overlay:
            graph, overlay_trace = knowledge_base.apply_calibration_overlay(graph, calibration_overlay)
            graph_trace = [*graph_trace, *overlay_trace]
        graph_query = " ".join(
            part
            for part in (
                str(job.get("title") or ""),
                str(job.get("ability_keywords") or ""),
                knowledge_base.profile_context(profile_match["profile"]).get("track", "") if profile_match else "",
            )
            if part
        )
        graph_pool, graph_pool_trace = knowledge_base.derive_graph_pool(graph, query_text=graph_query)
        # 目标池不得含客户本公司（图谱按赛道召回时本公司常高分命中）
        client_raw = " ".join(str(job.get("client") or "").split())
        client_norm = knowledge_base.normalize_client_name(client_raw)
        before = len(graph_pool)
        graph_pool = [
            company
            for company in graph_pool
            if not knowledge_base.name_match_rule(client_raw, client_norm, company["name"])[0]
        ]
        if len(graph_pool) < before:
            graph_pool_trace.append(f"已从图谱池剔除客户本公司 {before - len(graph_pool)} 家")
        # S4-3：排除规则引擎 —— 第 4 步之后强制过五类检查清单（PRD §4），
        # 逐类输出 适用/不适用+依据；禁挖名单/竞业从 restricted 层按客户继承。
        # 只在运行时并入 negative_rules，绝不进入 LLM 输入。
        negative_checklist, checklist_trace = negative_rules.build_negative_rule_checklist(
            job,
            restricted_info=restricted_info,
            archetype=archetype,
            consultant_answers=consultant["consultant_answers"],
        )
        classification["trace"] = [
            *classification["trace"], *profile_trace, *restricted_trace, *graph_trace, *graph_pool_trace, *checklist_trace,
        ]
        # 知识飞轮二期：技能本体（step4 别名归一/相关词提示，source=kb_skill）与
        # 职级映射（step3 优先查 kb_level_mapping，source=kb_level，查不到走 LLM/原型路径）。
        skill_ontology, skill_ontology_trace = knowledge_base.load_skill_ontology()
        level_mapping, level_mapping_trace = knowledge_base.load_level_mapping()
        level_hit, level_hit_trace = knowledge_base.map_level(job.get("title"), level_mapping)
        classification["trace"] = [
            *classification["trace"], *skill_ontology_trace, *level_mapping_trace, *level_hit_trace,
        ]
        client_profile_payload: dict[str, Any] = {"matched": False}
        if profile_match:
            client_profile_payload = {
                "matched": True,
                "name": profile_match["name"],
                "rule": profile_match["rule"],
                "needs_confirmation": profile_match["needs_confirmation"],
                "context": knowledge_base.profile_context(profile_match["profile"]),
            }
        payload = {
            "canonical_position": {
                "client": job["client"], "job": job["title"],
                "requirements": job.get("requirements") or job.get("hard_requirements") or "",
                "responsibilities": job.get("responsibilities") or "",
                "education": job.get("education") or "", "experience": job.get("experience") or "",
                "hard_requirements": job.get("hard_requirements") or "",
                "exclusions": job.get("exclusions") or "", "location": job.get("position_location") or job.get("location") or "",
                "objective": inputs.get("objective") or "",
            },
            "legacy_profile_suggestions": {
                "ability_keywords": _loads((job.get("profile") or {}).get("ability_keywords_json"), []),
                "search_keywords": _loads((job.get("profile") or {}).get("search_keywords_json"), []),
                "target_companies": _loads((job.get("profile") or {}).get("target_companies_json"), []),
            },
            "input_classification": {
                "input_level": classification["input_level"],
                "anchors": classification["anchors"],
                "missing_anchors": classification["missing_anchors"],
            },
            "job_archetype": {
                key: archetype[key]
                for key in ("archetype_id", "title", "client", "essence", "directions", "target_functions", "level_mapping", "keyword_groups")
            } if archetype else {},
            "consultant_input": consultant,
            "client_profile": client_profile_payload,
            "kb_graph_candidates": graph_pool,
            **learning,
            "stop_note_adjustments_summary": self._stop_note_adjustments_summary(learning.get("stop_note_adjustments") or []),
            "deterministic_fallback": fallback,
        }
        try:
            generated = self.service.llm.generate_search_strategy(payload)
        except Exception:
            generated = {}
        llm_v2_fragment = generated.pop("strategy_v2", None) if isinstance(generated, dict) else None
        plan = self._validate_model_strategy(generated, fallback, job, learning, max_queries)
        plan["generation"]["input_level"] = classification["input_level"]
        v2 = strategy_v2.build_strategy_v2(
            plan, classification, archetype=archetype, consultant=consultant, llm_fragment=llm_v2_fragment,
            profile_match=knowledge_base.profile_matched_info(profile_match),
            graph_pool=graph_pool,
            restricted_rules=knowledge_base.restricted_negative_rules(restricted_info),
            banned_companies=(
                (restricted_info.get("constraints") or {}).get("banned_companies")
                if isinstance(restricted_info, dict) and isinstance(restricted_info.get("constraints"), dict)
                else []
            ),
            negative_checklist=negative_checklist,
            canonical_position=payload["canonical_position"],
            skill_ontology=skill_ontology,
            level_hit=level_hit,
            profile_context=(
                knowledge_base.profile_context(profile_match["profile"])
                if profile_match and isinstance(profile_match.get("profile"), dict)
                else {}
            ),
            learning=learning,
        )
        consultant_constraints = _lock_consultant_constraints(plan, v2, revision_evidence, locked_items)
        strategy_v2.refresh_consultant_judgement(v2)
        v2_ok, v2_errors = strategy_v2.validate_strategy_v2(v2)
        constraint_errors = _locked_constraint_conflicts(plan, v2, consultant_constraints)
        if constraint_errors:
            v2_ok = False
            v2_errors = [*v2_errors, *constraint_errors]
        # S4-3c-4（N6）：策略全要素消费检查 —— 对照命中原型的种子要素清单（T1/T2/T3 各层
        # 公司池、地点策略、排除规则、有效关键词组）核对 strategy_v2 是否全部消费，未使用项
        # 显式列出供顾问确认页展示；种子未命中（无原型岗位）coverage_report=None 留痕不算缺失。
        coverage_report = strategy_v2.build_coverage_report(archetype, v2)
        v2["coverage_report"] = coverage_report
        if coverage_report:
            existing_trace = v2.get("classification_trace")
            v2["classification_trace"] = [
                *(existing_trace if isinstance(existing_trace, list) else []),
                f"N6 要素消费检查：消费 {coverage_report['consumed_count']}/{coverage_report['element_count']} 项"
                + (
                    f"，未使用：{'、'.join(item['element'] for item in coverage_report['unused'])}"
                    if coverage_report["unused"]
                    else "，种子要素全部消费"
                ),
            ]
        result: dict[str, Any] = {
            "summary": "已基于岗位事实、客户画像、岗位原型和历史反馈生成寻访策略，并补齐资深顾问判断简报。",
            "strategy": plan,
            "input_level": classification["input_level"],
            "knowledge_health": knowledge_health,
            "references": self._job_reference(job),
        }
        if consultant_constraints:
            result["consultant_constraints"] = consultant_constraints
        if v2_ok:
            query_plan = query_builders.schedule_query_plan_v1(
                query_builders.compile_query_plan_v1(v2),
                self._query_plan_learning_metrics(int(job.get("id") or 0)),
            )
            golden_replay = strategy_v2.build_golden_candidate_replay(archetype, query_plan)
            result["strategy_v2"] = v2
            result["query_plan_v1"] = query_plan
            result["golden_candidate_replay_v1"] = golden_replay
            if sourcing_adjustment_input:
                # 这里只冻结本轮真实输入。状态转换由 WorkflowEngine 在策略产物落库事务内完成。
                result["sourcing_adjustment_consumption"] = {
                    "status": "awaiting_artifact_persistence",
                    "job_id": int(job.get("id") or 0),
                    "adjustment_ids": [item["id"] for item in sourcing_adjustment_input],
                    "input": sourcing_adjustment_input,
                }
            content = "# 多渠道寻访策略（strategy_v2）\n\n```json\n" + json.dumps(v2, ensure_ascii=False, indent=2) + "\n```"
            result["artifacts"] = [
                self._artifact(
                    "search_strategy", "多渠道寻访策略", content=content,
                    metadata={
                        "plan": plan,
                        "strategy_v2": v2,
                        "query_plan_v1": query_plan,
                        "golden_candidate_replay_v1": golden_replay,
                        "schema_version": strategy_v2.STRATEGY_V2_VERSION,
                        "coverage_report": coverage_report,
                        "sourcing_adjustment_input": sourcing_adjustment_input,
                    },
                )
            ]
        else:
            # 即使策略未通过 schema 校验，脚本能力也必须返回一个可审计的结果。
            # 这个产物是诊断记录，不是可执行策略；后置条件保持失败，避免无效策略
            # 被工作流误判为完成，同时把真实校验错误传给工作流和前端。
            error_payload = {
                "errors": v2_errors,
                "trace": classification["trace"][-12:],
                "input_level": classification["input_level"],
                "schema_version": strategy_v2.STRATEGY_V2_VERSION,
                "client": str(job.get("client") or ""),
                "job": str(job.get("title") or ""),
            }
            result["strategy_v2_error"] = error_payload
            readable_errors = v2_errors or ["未提供具体校验错误"]
            diagnostic_lines = [
                "# 多渠道寻访策略校验诊断",
                "",
                f"岗位：{job.get('client') or ''}｜{job.get('title') or ''}",
                f"输入分级：{classification['input_level']}",
                "",
                "## 校验错误",
                *[f"- {error}" for error in readable_errors],
                "",
                "## 执行追踪",
                *[f"- {item}" for item in error_payload["trace"]],
                "",
                "该记录仅用于审计和定位问题，未生成可执行寻访策略。",
            ]
            result["summary"] = "寻访策略未通过 strategy_v2 校验，已生成诊断记录。"
            result["artifacts"] = [
                self._artifact(
                    "search_strategy",
                    "寻访策略校验诊断",
                    content="\n".join(diagnostic_lines),
                    validation="failed",
                    metadata={
                        "diagnostic": True,
                        "schema_version": strategy_v2.STRATEGY_V2_VERSION,
                        "strategy_v2_error": error_payload,
                    },
                )
            ]
            result["postcondition"] = {
                "verified": False,
                "recoverable": True,
                "reason": "寻访策略未通过 strategy_v2 校验：" + "；".join(readable_errors[:4]),
            }
        return result

    def run_multi_channel_sourcing(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        job = self._job(context)
        strategy = self._workflow_strategy(inputs)
        approved_snapshot = self._approved_sourcing_snapshot(str(inputs.get("workflow_id") or ""))
        if approved_snapshot:
            strategy = {
                "strategy_summary": approved_snapshot.get("summary") or "",
                "channels": approved_snapshot.get("channels") or {},
                "target_companies": [
                    item.get("name") for item in (approved_snapshot.get("company_pool") or [])
                    if isinstance(item, dict) and item.get("name")
                ],
                "consultant_constraints": approved_snapshot.get("locked_constraints") or [],
            }
        try:
            preflight = self._run_json([self.python, str(MULTICHANNEL), "preflight", "--db", str(self.service.db_path), "--client", str(job["client"]), "--job", str(job["title"]), "--port", str(int(inputs.get("cdp_port") or 9223))], 90)
        except Exception as exc:
            preflight = {"ok": False, "status": "preflight_unavailable", "error": str(exc)[:1000]}
        channels = preflight.get("channels") or preflight.get("preflight") or {}
        ticket = {
            "client": job["client"], "job": job["title"], "preflight": preflight,
            "workflow_id": str(inputs.get("workflow_id") or ""),
            "target_count": int(approved_snapshot.get("target_count") or inputs.get("target_count") or self._target_count(inputs.get("objective")) or 10),
            "cdp_port": int(inputs.get("cdp_port") or 9223),
            "strategy": strategy,
            "strategy_snapshot": approved_snapshot,
            "strategy_hash": str(approved_snapshot.get("strategy_hash") or ""),
            "query_plan_v1": approved_snapshot.get("query_plan_v1") or {},
            "query_plan_hash": str(approved_snapshot.get("query_plan_hash") or ""),
            "required_result": {"verified": True, "channel_runs": [], "intake": {}, "audit": {}},
        }
        return {
            "summary": "渠道预检完成，已生成受约束寻访任务；渠道执行完成并读回前不会进入评估。",
            "references": self._job_reference(job), "external_action_executed": False,
            "external_request": ticket,
            "auto_execute_request": ticket if preflight.get("ok") is not False else None,
            "artifacts": [self._artifact("sourcing_ticket", "多渠道寻访执行任务", content=json.dumps(ticket, ensure_ascii=False, indent=2), validation="pending_execution", metadata={"channels": channels})],
        }

    def _workflow_strategy(self, inputs: dict[str, Any]) -> dict[str, Any]:
        workflow_id = str(inputs.get("workflow_id") or "")
        if not workflow_id:
            return {}
        conn = self.service._connect()
        try:
            row = conn.execute(
                "SELECT output_json FROM agent_workflow_steps WHERE workflow_id=? AND capability_id='search_strategy' AND status='completed' ORDER BY sequence DESC LIMIT 1",
                (workflow_id,),
            ).fetchone()
            output = _loads(row["output_json"], {}) if row else {}
            return output.get("strategy") if isinstance(output.get("strategy"), dict) else {}
        finally:
            conn.close()

    def _approved_sourcing_snapshot(self, workflow_id: str) -> dict[str, Any]:
        if not workflow_id:
            return {}
        conn = self.service._connect()
        try:
            row = conn.execute(
                """
                SELECT preflight_json FROM agent_approvals
                WHERE workflow_id=? AND action_type='multi_channel_sourcing' AND status='approved'
                ORDER BY id DESC LIMIT 1
                """,
                (workflow_id,),
            ).fetchone()
            preflight = _loads(row["preflight_json"], {}) if row else {}
            snapshot = preflight.get("strategy_snapshot") if isinstance(preflight.get("strategy_snapshot"), dict) else {}
            if snapshot and preflight.get("strategy_hash") == snapshot.get("strategy_hash"):
                return snapshot
            return {}
        finally:
            conn.close()

    @staticmethod
    def _target_count(objective: Any) -> int:
        match = re.search(r"(\d+)\s*(?:位|个|人)", str(objective or ""))
        return min(100, int(match.group(1))) if match else 0

    def run_job_publish_prepare(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        job = self._job(context)
        overrides = inputs.get("publish_fields") if isinstance(inputs.get("publish_fields"), dict) else {}
        draft = {
            "client_company": overrides.get("client_company") or job.get("client"),
            "job_title": overrides.get("job_title") or job.get("title"),
            "city_keyword": overrides.get("city_keyword") or job.get("position_location") or job.get("location"),
            "city_choice": overrides.get("city_choice") or job.get("position_location") or job.get("location"),
            "salary_low_k": overrides.get("salary_low_k"), "salary_high_k": overrides.get("salary_high_k"),
            "salary_months": overrides.get("salary_months") or 12,
            "job_category_keyword": overrides.get("job_category_keyword"), "job_category_choice": overrides.get("job_category_choice"),
            "industry_keyword": overrides.get("industry_keyword"), "industry_choice": overrides.get("industry_choice"),
            "work_year_keyword": overrides.get("work_year_keyword") or job.get("experience"),
            "work_year_choice": overrides.get("work_year_choice"), "work_year_low": overrides.get("work_year_low"), "work_year_high": overrides.get("work_year_high"),
            "education_choice": overrides.get("education_choice") or job.get("education"), "education_tongzhao": bool(overrides.get("education_tongzhao", False)),
            "private_job": bool(overrides.get("private_job", False)), "recruit_count": int(overrides.get("recruit_count") or job.get("headcount") or 1),
            "close_date": overrides.get("close_date") or job.get("deadline"),
            "description": overrides.get("description") or "\n".join(filter(None, [str(job.get("responsibilities") or ""), str(job.get("requirements") or job.get("summary") or "")])),
        }
        salary = str(job.get("salary") or "")
        numbers = [int(value) for value in re.findall(r"\d+", salary)]
        if not draft["salary_low_k"] and numbers:
            draft["salary_low_k"] = numbers[0]
        if not draft["salary_high_k"] and len(numbers) > 1:
            draft["salary_high_k"] = numbers[1]
        required = ["client_company", "job_title", "city_keyword", "salary_low_k", "salary_high_k", "job_category_choice", "industry_choice", "description", "close_date"]
        missing = [key for key in required if draft.get(key) in (None, "")]
        path = self._path("job_publish", f"{job['client']}-{job['title']}-draft", "json")
        path.write_text(json.dumps(draft, ensure_ascii=False, indent=2), encoding="utf-8")
        artifacts = [
            self._artifact(
                "job_publish_draft", "猎聘岗位发布草稿", file_path=path, mime_type="application/json",
                content=json.dumps(draft, ensure_ascii=False, indent=2),
                validation="needs_input" if missing else "passed",
                metadata={"client": job.get("client"), "job": job.get("title"), "job_id": job.get("id")},
            )
        ]
        result = {"summary": "岗位发布草稿已生成。" if not missing else "岗位发布草稿缺少关键字段，已阻塞正式发布。",
                  "references": self._job_reference(job), "draft": draft, "missing_inputs": missing,
                  "artifacts": artifacts}
        if missing:
            result["blocked"] = True
            return result
        prepare_log = self._path("job_publish", "liepin-publish-prepare-readback", "json")
        payload = self._run_json([
            self.python, str(LIEPIN_PUBLISH), "--mode", "prepare",
            "--port", str(int(inputs.get("cdp_port") or 9223)), "--draft", str(path), "--log", str(prepare_log),
        ], 180)
        failed = payload.get("ok") is False or str(payload.get("status") or "").lower() in {"blocked", "failed", "error"}
        artifacts.append(self._artifact(
            "job_publish_prepare_readback", "猎聘岗位发布预检读回", file_path=prepare_log,
            mime_type="application/json", content=json.dumps(payload, ensure_ascii=False, indent=2),
            validation="blocked" if failed else "passed",
            metadata={"draft_path": str(path), "client": job.get("client"), "job": job.get("title"), "job_id": job.get("id")},
        ))
        result["prepare_readback"] = payload
        result["summary"] = "岗位发布草稿已填入猎聘发布表单并完成读回预检。" if not failed else "猎聘岗位发布预检未通过，已阻塞正式发布。"
        if failed:
            result["blocked"] = True
            result["missing_inputs"] = ["修正猎聘发布预检问题"]
        return result

    def run_job_publish_execute(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        draft = self._dependency_file(inputs, "job_publish_draft")
        readback = self._dependency_file(inputs, "job_publish_prepare_readback")
        if not draft or not readback:
            return self._blocked("没有通过校验和读回的岗位发布草稿。", ["job_publish_draft", "job_publish_prepare_readback"])
        log = self._path("job_publish", "liepin-publish-readback", "json")
        payload = self._run_json([self.python, str(LIEPIN_PUBLISH), "--mode", "publish", "--confirm", "PUBLISH", "--port", str(int(inputs.get("cdp_port") or 9223)), "--draft", str(draft), "--log", str(log)], 180)
        verified = bool(payload.get("verified") or payload.get("status") in {"published", "submitted", "auditing"})
        if not verified:
            raise RuntimeError("猎聘发布动作未通过结果页或职位列表读回验证")
        job = self._job(context)
        conn = self.service._connect()
        try:
            if self._table(conn, "positions"):
                columns = {str(row[1]) for row in conn.execute("PRAGMA table_info(positions)").fetchall()}
                assignments = ["status='已发布'", "updated_at=datetime('now','localtime')"]
                if "liepin_status" in columns:
                    assignments.append("liepin_status='已发布/已验证'")
                if "liepin_published_at" in columns:
                    assignments.append("liepin_published_at=datetime('now','localtime')")
                if "liepin_verify_log" in columns:
                    assignments.append("liepin_verify_log=?")
                    conn.execute(f"UPDATE positions SET {','.join(assignments)} WHERE client=? AND title=?", (str(log), job["client"], job["title"]))
                else:
                    conn.execute(f"UPDATE positions SET {','.join(assignments)} WHERE client=? AND title=?", (job["client"], job["title"]))
                conn.commit()
        finally:
            conn.close()
        return {"summary": "猎聘岗位已发布并完成页面读回验证。", "external_action_executed": True,
                "external_result": payload, "artifacts": [self._artifact("external_action_receipt", "猎聘岗位发布回执", file_path=log, mime_type="application/json", content=json.dumps(payload, ensure_ascii=False, indent=2))]}

    def _dependency_file(self, inputs: dict[str, Any], artifact_type: str) -> Path | None:
        workflow_id = str(inputs.get("workflow_id") or "")
        if not workflow_id:
            return None
        conn = self.service._connect()
        try:
            row = conn.execute("SELECT file_path,validation_status FROM agent_artifacts WHERE workflow_id=? AND artifact_type=? ORDER BY id DESC LIMIT 1", (workflow_id, artifact_type)).fetchone()
            if row and row["file_path"] and row["validation_status"] == "passed" and Path(row["file_path"]).exists():
                return Path(row["file_path"])
            return None
        finally:
            conn.close()

    def run_resume_export(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        candidate = self._candidate(context)
        from docx import Document
        identity, position = candidate["identity"], candidate["position"]
        doc = Document()
        doc.add_heading(str(identity.get("name") or "候选人简历"), 0)
        for label, value in (("当前公司", identity.get("company")), ("当前职位", identity.get("title")), ("城市", identity.get("city")), ("学历", identity.get("education")), ("经验", identity.get("experience")), ("目标岗位", f"{position.get('client','')} / {position.get('job','')}")):
            doc.add_paragraph(f"{label}：{value or '不详'}")
        doc.add_heading("履历原始证据", level=1)
        profiles = candidate.get("source_profiles") or []
        if not profiles:
            doc.add_paragraph("暂无来源简历正文。")
        for profile in profiles:
            raw = _loads(profile.get("raw_json"), {})
            text = raw.get("profile_text") or raw.get("resume_text") or raw.get("text") or ""
            if text:
                doc.add_paragraph(str(text))
        path = self._path("resumes", f"{identity.get('name')}-{position.get('job')}-结构化简历", "docx")
        doc.save(path)
        return {"summary": "结构化简历 DOCX 已生成。", "references": self._candidate_reference(candidate),
                "artifacts": [self._artifact("resume_document", "结构化简历", file_path=path, mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")]}

    def run_matching_report(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        candidate = self._candidate(context)
        relation, identity, position = candidate["relation"], candidate["identity"], candidate["position"]
        assessment = self._latest_assessment(int(relation["job_candidate_id"]))
        criteria = assessment.get("criteria") or {}
        status_icon = {"met": "通过", "partial": "部分", "not_met": "不通过", "unknown": "待核验"}
        hard = [[item.get("criterion"), "；".join(item.get("evidence") or []) or item.get("reason") or "无证据", status_icon.get(item.get("status"), "待核验")] for item in criteria.get("hard_requirements") or []]
        star_map = {"met": 5, "partial": 3, "not_met": 1, "unknown": 2}
        matches = [{"duty": item.get("criterion"), "evidence": "；".join(item.get("evidence") or []) or "待核验", "stars": star_map.get(item.get("status"), 2)} for item in criteria.get("core_abilities") or []]
        risks = [{"level": "高" if "关键" in str(value) else "中", "title": str(value), "description": str(value), "verify": "在下一次沟通中核验并记录证据"} for value in assessment.get("risks") or assessment.get("gaps") or []]
        data = {
            "candidate": identity.get("name"), "company": position.get("client"), "position": position.get("job"),
            "hard_gates": hard, "responsibility_matches": matches,
            "bonus_items": [[value, "来自 ASA 当前评估", "通过"] for value in (assessment.get("strengths") or [])],
            "risks": risks, "scores": {"综合匹配": int(assessment.get("fit_score") or 0)},
            "total_score": int(assessment.get("fit_score") or 0),
            "interview_suggestions": [["证据核验", value] for value in (assessment.get("verification_questions") or [])],
            "verdict": str(assessment.get("recommendation") or "待复核"), "conclusion_summary": str(assessment.get("next_action") or "人工复核"),
        }
        data_path = self._path("reports", f"{identity.get('name')}-matching-data", "json")
        output = self._path("reports", f"人岗匹配-{position.get('client')}-{position.get('job')}-{identity.get('name')}", "docx")
        data_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        self._run([self.python, str(MATCHING_REPORT), "--data", f"@{data_path}", "--output", str(output)], 120)
        return {"summary": "人岗匹配分析报告已生成。", "references": self._candidate_reference(candidate),
                "artifacts": [self._artifact(
                    "matching_report", "人岗匹配分析报告", file_path=output,
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    metadata={
                        "assessment_id": assessment.get("id"), "job_candidate_id": relation.get("job_candidate_id"),
                        "person_id": relation.get("person_id"), "candidate_id": relation.get("source_candidate_id"),
                        "job_id": relation.get("job_id"), "client": position.get("client"), "job": position.get("job"),
                    },
                )]}

    def _s6_assessment_doc(self, job_candidate_id: int, job_id: Any) -> dict[str, Any] | None:
        """S6 判人评估 artifact（candidate_assessment，人×岗）；不存在/岗位号缺失 → None。"""
        try:
            jid = int(job_id or 0)
        except (TypeError, ValueError):
            return None
        if not jid:
            return None
        conn = self.service._connect()
        try:
            payload = candidate_assessment.get_assessment(conn, int(job_candidate_id), jid)
            return payload["assessment"] if payload else None
        finally:
            conn.close()

    def run_recommendation_report(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        attachment_candidate = inputs.get("attachment_candidate") if isinstance(inputs.get("attachment_candidate"), dict) else None
        if attachment_candidate:
            # attachment-only：候选人不在系统内（仅存在于上传简历附件），
            # 本质是“简历→嘉驰模板”转换，合法绕过 S6-3 判人评估前置。
            return self._run_attachment_recommendation_report(attachment_candidate)
        candidate = self._candidate(context)
        relation, identity, position = candidate["relation"], candidate["identity"], candidate["position"]
        # S6-3：推荐报告强制引用判人评估块（PRD §1 ②）——无评估不许退回纯简历罗列。
        s6 = self._s6_assessment_doc(int(relation["job_candidate_id"]), relation.get("job_id"))
        if s6 is None:
            return self._blocked(
                "推荐报告必须引用判人评估结论：该人选还没有判人评估，请先在人选卡「评估」区生成判人评估"
                "（职业轨迹/在同龄人里的位置/动机与时机/需要核实的问题），再生成推荐报告。",
                ["先完成判人评估（candidate_assessment）"],
                self._candidate_reference(candidate),
            )
        assessment_block = candidate_assessment.report_reference_block(s6)
        assessment = self._latest_assessment(int(relation["job_candidate_id"]))
        if float(assessment.get("evidence_coverage") or 0) < 0.75:
            return self._blocked("证据覆盖不足，不能生成可发送推荐报告。", ["核验问题完成", "证据覆盖率>=0.75"], self._candidate_reference(candidate))
        if not JIASHI_TEMPLATE.exists():
            return self._blocked("嘉驰标准模板不存在。", [str(JIASHI_TEMPLATE)], self._candidate_reference(candidate))
        profile_summary = str((candidate.get("candidate_profile") or {}).get("profile_summary") or "履历职责待进一步结构化核验")
        data = {
            "customer": position.get("client"), "position": position.get("job"), "name": identity.get("name"),
            "current_location": identity.get("city") or "不详", "expected_location": "不详",
            "consultant_comments": assessment_block["lines"] + (assessment.get("strengths") or [])[:6],
            "education": [str(identity.get("education") or "不详")],
            "work_experience": [f"时间不详 {identity.get('company') or '公司待核验'}\n担任职位：{identity.get('title') or '职位待核验'}\n工作职责：{profile_summary}"],
            "project_experience": ["暂无明确项目经历"], "motivation": "待核验", "leaving_reason": "待核验",
        }
        data_path = self._path("reports", f"{identity.get('name')}-jiashi-data", "json")
        output = self._path("reports", f"嘉驰国际-{position.get('client')}-{position.get('job')}-{identity.get('name')}", "docx")
        data_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        self._run([self.python, str(JIASHI_REPORT), "--template", str(JIASHI_TEMPLATE), "--data", str(data_path), "--output", str(output)], 180)
        self._sanitize_docx_privacy(output)
        audit = self._run([self.python, str(JIASHI_AUDIT), str(output)], 120)
        return {"summary": "嘉驰推荐报告草稿已生成并通过模板审计，已引用判人评估块（评估只辅助判断，发送前请顾问复核）。", "references": self._candidate_reference(candidate),
                "artifacts": [self._artifact(
                    "recommendation_report", "嘉驰推荐报告", file_path=output,
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    content=audit.stdout[-4000:],
                    metadata={
                        "assessment_id": assessment.get("id"), "job_candidate_id": relation.get("job_candidate_id"),
                        "person_id": relation.get("person_id"), "candidate_id": relation.get("source_candidate_id"),
                        "job_id": relation.get("job_id"), "client": position.get("client"), "job": position.get("job"),
                        "attached_to_candidate": True, "external_submitted": False,
                        "s6_assessment": {
                            "artifact_id": f"candidate_assessment_{int(relation['job_candidate_id'])}_{int(relation.get('job_id') or 0)}",
                            "as_of": assessment_block["as_of"],
                            "assessor_version": assessment_block["assessor_version"],
                            "trajectory_verdict": assessment_block["trajectory_verdict"],
                            "percentile_band": assessment_block["percentile_band"],
                            "percentile_band_label": assessment_block["percentile_band_label"],
                            "reference_n": assessment_block["reference_n"],
                            "top_risks": assessment_block["top_risks"],
                            "risks_pending": assessment_block["risks_pending"],
                        },
                    },
                )]}

    def _run_attachment_recommendation_report(self, attachment_candidate: dict[str, Any]) -> dict[str, Any]:
        """attachment-only 候选人：从上传简历文本解析字段 → 嘉驰模板 → 审计。

        候选人不挂任何 job_candidate 关系，artifact 标注 source=attachment_resume、
        不含 s6_assessment，并提示“基于上传简历生成，发送前请顾问复核”。
        """
        name = str(attachment_candidate.get("name") or "").strip()
        resume_text = str(attachment_candidate.get("resume_text") or "").strip()
        file_name = str(attachment_candidate.get("file_name") or "").strip()
        references = [
            {"type": "local_attachment", "id": "", "label": file_name or "上传简历附件", "subtitle": "用户上传的简历附件"}
        ]
        if not name or not resume_text:
            return self._blocked("上传简历缺少可识别的候选人姓名或正文，无法生成推荐报告。", ["包含姓名与正文的简历附件"], references)
        if not JIASHI_TEMPLATE.exists():
            return self._blocked("嘉驰标准模板不存在。", [str(JIASHI_TEMPLATE)], references)
        fields = _parse_jiashi_resume_fields(resume_text)
        customer = str(fields.get("customer") or attachment_candidate.get("customer") or "不详")
        position = str(fields.get("position") or attachment_candidate.get("position") or "不详")
        data = {
            "customer": customer,
            "position": position,
            "name": name,
            "gender": str(fields.get("gender") or "不详"),
            "marital_status": str(fields.get("marital_status") or "不详"),
            "current_location": str(fields.get("current_location") or "不详"),
            "expected_location": "不详",
            "current_salary": str(fields.get("current_salary") or "不详"),
            "expected_salary": str(fields.get("expected_salary") or "待核验"),
            "consultant_comments": list(fields.get("consultant_comments") or []) or ["基于上传简历生成，顾问评语待补充。"],
            "education": list(fields.get("education") or []) or ["不详"],
            "work_experience": list(fields.get("work_experience") or []) or ["详见简历原文"],
            "project_experience": list(fields.get("project_experience") or []) or ["暂无明确项目经历"],
            "motivation": "待核验",
            "leaving_reason": "待核验",
        }
        data_path = self._path("reports", f"{name}-jiashi-attachment-data", "json")
        output = self._path("reports", f"嘉驰国际-{customer}-{position}-{name}", "docx")
        data_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        self._run([self.python, str(JIASHI_REPORT), "--template", str(JIASHI_TEMPLATE), "--data", str(data_path), "--output", str(output)], 180)
        self._sanitize_docx_privacy(output)
        audit = self._run([self.python, str(JIASHI_AUDIT), str(output)], 120)
        return {
            "summary": (
                f"已基于上传简历生成 {name} 的嘉驰推荐报告草稿并通过模板审计。"
                "候选人尚未入库，报告基于上传简历生成，发送前请顾问复核。"
            ),
            "references": references,
            "artifacts": [self._artifact(
                "recommendation_report", "嘉驰推荐报告", file_path=output,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                content=audit.stdout[-4000:],
                metadata={
                    "source": "attachment_resume",
                    "attachment_file": file_name,
                    "candidate_name": name,
                    "client": customer,
                    "job": position,
                    "attached_to_candidate": False,
                    "external_submitted": False,
                    "review_notice": "基于上传简历生成，发送前请顾问复核",
                },
            )],
        }

    @staticmethod
    def _sanitize_docx_privacy(path: Path) -> None:
        from docx import Document
        doc = Document(path)
        private = re.compile(r"(?:手机号|手机号码|联系电话|电话|微信|WeChat|微信号|邮箱|Email|E-mail)|(?:\+?86[-\s]?)?1[3-9]\d{9}|[A-Za-z0-9_.+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", re.I)
        for paragraph in doc.paragraphs:
            if private.search(paragraph.text or ""):
                paragraph.text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if private.search(paragraph.text or ""):
                            paragraph.text = ""
        doc.save(path)

    def run_interview_followup(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        return self._lifecycle_note(context, inputs, "interview_followup", "面试跟进", "interview_followup", "interview_note")

    def run_salary_negotiation(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        return self._lifecycle_note(context, inputs, "salary_negotiation", "谈薪跟进", "salary_negotiation", "salary_negotiation_note")

    def run_decision_coaching(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        return self._lifecycle_note(context, inputs, "decision_coaching", "候选人决策辅导", "decision_coaching", "decision_coaching")

    def run_onboarding_followup(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        return self._lifecycle_note(context, inputs, "onboarding_followup", "入职跟进", "onboarding_followup", "onboarding_note", days=7)

    def _lifecycle_note(self, context: dict[str, Any], inputs: dict[str, Any], event_type: str, label: str, task_type: str, artifact_type: str, days: int = 2) -> dict[str, Any]:
        candidate = self._candidate(context)
        objective = str(inputs.get("objective") or "").strip()
        if not objective:
            return self._blocked(f"{label}缺少业务事实。", ["objective"], self._candidate_reference(candidate))
        event_id = self._candidate_event(candidate, event_type, "recorded", objective, inputs)
        task_id = self._followup(candidate, task_type, objective, inputs, days)
        content = f"# {label}\n\n- 记录：{objective}\n- 事件 ID：{event_id}\n- 跟进任务 ID：{task_id or '未创建'}\n"
        return {"summary": f"{label}事实已记录，并创建后续任务。", "references": self._candidate_reference(candidate),
                "artifacts": [self._artifact(artifact_type, label, content=content, metadata={"event_id": event_id, "task_id": task_id})]}

    def run_salary_verification(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        candidate = self._candidate(context)
        data = inputs.get("salary_data") if isinstance(inputs.get("salary_data"), dict) else None
        if not data or not data.get("records"):
            return self._blocked("薪资核验需要结构化流水证据，未找到时不会生成虚假报告。", ["salary_data.records"], self._candidate_reference(candidate))
        data = {**data, "candidate_name": data.get("candidate_name") or candidate["identity"].get("name"), "report_date": data.get("report_date") or datetime.now().strftime("%Y-%m-%d")}
        data_path = self._path("salary", f"{candidate['identity'].get('name')}-salary-data", "json")
        out_dir = self.output_dir / "salary"
        data_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        proc = self._run([self.python, str(SALARY_REPORT), "--input", str(data_path), "--output-dir", str(out_dir)], 180)
        files = [Path(line.strip()) for line in proc.stdout.splitlines() if line.strip().endswith(".docx")]
        return {"summary": f"薪资证据报告已生成，共 {len(files)} 个文件。", "references": self._candidate_reference(candidate),
                "artifacts": [self._artifact("salary_report", path.stem, file_path=path, mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document") for path in files]}

    @staticmethod
    def _message_hash(job_candidate_id: Any, message: str) -> str:
        raw = f"{job_candidate_id}|{message}".encode("utf-8")
        return hashlib.sha256(raw).hexdigest()[:24]

    def _outreach_message(self, candidate: dict[str, Any], inputs: dict[str, Any]) -> str:
        explicit = " ".join(str(inputs.get("message") or "").split())
        if explicit:
            return explicit
        identity, position = candidate["identity"], candidate["position"]
        return (
            f"{identity.get('name') or '你好'}，你好。我这边有一个{position.get('client') or ''}"
            f"{position.get('job') or '岗位'}机会，和你{identity.get('company') or ''}"
            f"{identity.get('title') or ''}经历比较相关，方便了解一下吗？"
        )

    def _outreach_targets(self, context: dict[str, Any], inputs: dict[str, Any]) -> list[dict[str, Any]]:
        context_type = str(context.get("type") or "")
        if context_type == "candidate":
            return [self._candidate(context)]
        filters = context.get("filters") if isinstance(context.get("filters"), dict) else {}
        queue = str(inputs.get("queue") or filters.get("queue") or "待联系")
        limit = max(1, min(int(inputs.get("limit") or filters.get("limit") or 20), 20))
        inbox = self.service.get_flow_inbox(queue=queue, limit=limit)
        targets: list[dict[str, Any]] = []
        for item in inbox.get("items") or []:
            job_candidate_id = item.get("job_candidate_id")
            if not job_candidate_id:
                continue
            try:
                targets.append(build_candidate_context(self.service.db_path, int(job_candidate_id)))
            except Exception:
                continue
        return targets

    def run_outreach_prepare(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        targets = self._outreach_targets(context, inputs)
        if not targets:
            return self._blocked("当前目标或队列没有可触达的人选。", ["选择当前目标、当前队列或具体人选"])
        items: list[dict[str, Any]] = []
        skipped: list[dict[str, Any]] = []
        for candidate in targets[:20]:
            relation, identity, position = candidate["relation"], candidate["identity"], candidate["position"]
            if is_stopped(candidate):
                skipped.append({"job_candidate_id": relation["job_candidate_id"], "candidate": identity.get("name"), "reason": "关系已停止"})
                continue
            message = self._outreach_message(candidate, inputs)
            message_hash = self._message_hash(relation["job_candidate_id"], message)
            items.append({
                "job_candidate_id": relation["job_candidate_id"], "person_id": relation.get("person_id"),
                "job_id": relation.get("job_id"), "candidate": identity.get("name"),
                "company": identity.get("company"), "title": identity.get("title"),
                "client": position.get("client"), "job": position.get("job"), "channel": "猎聘职聊",
                "message": message, "message_hash": message_hash,
                "before": relation.get("clean_stage") or relation.get("raw_status") or "未触达",
                "after": "猎聘消息发送并读回后进入已触达/待回复",
                "status": "pending",
            })
        if not items:
            return self._blocked("本批人选均不可触达。", ["选择未停止且有明确人岗关系的人选"], [])
        payload = {
            "version": "1.0", "prepared_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "batch_limit": 20, "item_count": len(items), "items": items, "skipped": skipped,
        }
        path = self._path("outreach", "猎聘触达锁定草稿", "json")
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        references = [
            {"type": "candidate", "id": item["job_candidate_id"], "label": item["candidate"], "subtitle": f"{item['client']} / {item['job']}"}
            for item in items[:8]
        ]
        return {
            "summary": f"已锁定 {len(items)} 条猎聘触达草稿，正式发送前需要批量确认。",
            "references": references,
            "outreach_draft_batch": payload,
            "artifacts": [self._artifact(
                "outreach_draft_batch", "猎聘触达锁定草稿", file_path=path, mime_type="application/json",
                content=json.dumps(payload, ensure_ascii=False, indent=2),
                metadata={"item_count": len(items), "batch_limit": 20},
            )],
        }

    def _sent_message_exists(self, job_candidate_id: int, message_hash: str) -> bool:
        conn = self.service._connect()
        try:
            row = conn.execute(
                """
                SELECT 1 FROM candidate_events
                WHERE job_candidate_id=? AND event_type='candidate_outreach'
                  AND event_status='sent_verified' AND raw_json LIKE ?
                LIMIT 1
                """,
                (int(job_candidate_id), f"%{message_hash}%"),
            ).fetchone()
            return row is not None
        finally:
            conn.close()

    def _load_outreach_batch(self, inputs: dict[str, Any]) -> dict[str, Any] | None:
        path = self._dependency_file(inputs, "outreach_draft_batch")
        if not path:
            return None
        return _loads(path.read_text(encoding="utf-8"), {})

    def run_outreach_execute(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        batch = self._load_outreach_batch(inputs)
        if not batch:
            return self._blocked("没有已锁定并通过审批的触达草稿。", ["outreach_draft_batch"])
        results: list[dict[str, Any]] = []
        references: list[dict[str, Any]] = []
        for item in (batch.get("items") or [])[:20]:
            job_candidate_id = int(item.get("job_candidate_id") or 0)
            message = str(item.get("message") or "").strip()
            message_hash = str(item.get("message_hash") or self._message_hash(job_candidate_id, message))
            if not job_candidate_id or not message:
                results.append({**item, "status": "failed", "error": "缺少人岗关系或文案"})
                continue
            if self._sent_message_exists(job_candidate_id, message_hash):
                results.append({**item, "status": "skipped", "reason": "同一锁定文案已发送并验证"})
                continue
            candidate = build_candidate_context(self.service.db_path, job_candidate_id)
            if is_stopped(candidate):
                results.append({**item, "status": "skipped", "reason": "关系已停止"})
                continue
            identity, position = candidate["identity"], candidate["position"]
            base = [
                self.python, str(LIEPIN_OUTREACH), "--port", str(int(inputs.get("cdp_port") or 9223)),
                "--candidate", str(identity.get("name")), "--message", message, "--check", str(position.get("job")),
            ]
            try:
                dry = self._run_json(base, 90)
                if dry.get("status") != "dry_run_ok":
                    raise RuntimeError("猎聘触达预检未通过")
                sent = self._run_json(base + ["--send"], 120)
                if sent.get("status") != "sent_verified":
                    raise RuntimeError("猎聘消息点击后未通过会话读回验证")
                event_id = self._candidate_event(
                    candidate, "candidate_outreach", "sent_verified", f"猎聘触达已验证：{message[:80]}",
                    {**inputs, "message_hash": message_hash, "locked_message": message, "channel_result": sent},
                )
                conn = self.service._connect()
                try:
                    conn.execute(
                        "UPDATE job_candidates SET raw_status='contacted',clean_stage='已触达',flow_bucket='已触达/待回复',updated_at=datetime('now','localtime') WHERE id=?",
                        (job_candidate_id,),
                    )
                    conn.commit()
                finally:
                    conn.close()
                result_item = {**item, "status": "sent_verified", "event_id": event_id, "dry_run": dry, "receipt": sent}
            except Exception as exc:
                result_item = {**item, "status": "failed", "error": str(exc)[:1000]}
            results.append(result_item)
            references.append({"type": "candidate", "id": job_candidate_id, "label": item.get("candidate"), "subtitle": f"{item.get('client','')} / {item.get('job','')}"})
        success = [item for item in results if item.get("status") in {"sent_verified", "skipped"}]
        failed = [item for item in results if item.get("status") == "failed"]
        if not success and failed:
            raise RuntimeError("本批猎聘触达全部失败：" + "；".join(str(item.get("error") or "") for item in failed[:3]))
        payload = {
            "verified": bool(success), "batch_status": "partial_failed" if failed else "completed",
            "sent_count": len([item for item in results if item.get("status") == "sent_verified"]),
            "skipped_count": len([item for item in results if item.get("status") == "skipped"]),
            "failed_count": len(failed), "items": results,
        }
        return {
            "summary": f"猎聘触达完成：发送 {payload['sent_count']} 人，跳过 {payload['skipped_count']} 人，失败 {payload['failed_count']} 人。",
            "references": references[:8], "external_action_executed": True, "external_result": payload,
            "artifacts": [self._artifact(
                "external_action_receipt", "猎聘触达批量回执", content=json.dumps(payload, ensure_ascii=False, indent=2),
                metadata={"sent_count": payload["sent_count"], "failed_count": payload["failed_count"]},
            )],
        }

    def run_client_recommendation(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        candidate = self._candidate(context)
        report = self._dependency_file(inputs, "recommendation_report")
        if not report:
            return self._blocked("客户推荐缺少通过审计的推荐报告。", ["recommendation_report"], self._candidate_reference(candidate))
        request = {"candidate": candidate["identity"].get("name"), "client": candidate["position"].get("client"), "job": candidate["position"].get("job"), "report": str(report), "channel": inputs.get("channel") or "manual_client_channel"}
        return {"summary": "客户推荐材料已锁定，等待指定客户渠道完成发送并读回。", "references": self._candidate_reference(candidate),
                "external_action_executed": False, "external_request": request,
                "artifacts": [self._artifact("external_action_ticket", "客户推荐执行任务", content=json.dumps(request, ensure_ascii=False, indent=2), validation="pending_execution")]}

    def run_offer_confirmation(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        candidate = self._candidate(context)
        terms = inputs.get("offer_terms") if isinstance(inputs.get("offer_terms"), dict) else None
        if not terms:
            return self._blocked("Offer 确认缺少明确条件，不改变候选人阶段。", ["offer_terms"], self._candidate_reference(candidate))
        event_id = self._candidate_event(candidate, "offer_confirmation", "confirmed", "Offer 条件已人工确认", {**inputs, "offer_terms": terms})
        return {"summary": "Offer 条件已记录；未代表候选人已经接受。", "references": self._candidate_reference(candidate), "external_action_executed": True,
                "artifacts": [self._artifact("offer_confirmation", "Offer 条件确认", content=json.dumps(terms, ensure_ascii=False, indent=2), metadata={"event_id": event_id})]}

    def run_project_retrospective(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        if context.get("type") == "candidate":
            candidate = self._candidate(context)
            refs = self._candidate_reference(candidate)
            relation_id = candidate["relation"]["job_candidate_id"]
            conn = self.service._connect()
            try:
                events = [_row(row) for row in conn.execute("SELECT event_type,event_status,event_time,summary FROM candidate_events WHERE job_candidate_id=? ORDER BY event_time,id", (relation_id,)).fetchall()]
            finally:
                conn.close()
            content = "# 项目复盘\n\n" + "\n".join(f"- {item.get('event_time') or ''} · {item.get('event_type')} · {item.get('summary') or ''}" for item in events)
        else:
            job = self._job(context)
            refs = self._job_reference(job)
            conn = self.service._connect()
            try:
                counts = {row["clean_stage"] or "未整理": int(row["total"]) for row in conn.execute("SELECT clean_stage,COUNT(*) total FROM job_candidates WHERE job_id=? GROUP BY clean_stage", (job["id"],)).fetchall()}
            finally:
                conn.close()
            content = "# 项目复盘\n\n```json\n" + json.dumps(counts, ensure_ascii=False, indent=2) + "\n```"
        return {"summary": "已基于 v3 事件与漏斗生成项目复盘。", "references": refs,
                "artifacts": [self._artifact("project_retrospective", "项目复盘", content=content)]}

    def run_memory_capture(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        content = str(inputs.get("confirmed_memory") or "").strip()
        if not content:
            return self._blocked("长期记忆只沉淀明确确认的信息。", ["confirmed_memory"])
        scope = context.get("type") if context.get("type") in {"job", "candidate"} else "global"
        result = self.service.store_memory(scope, str(context.get("id") or ""), "workflow_outcome", content, "agent_workflow", str(inputs.get("workflow_id") or ""), 1.0)
        return {"summary": "经确认的业务经验已写入 ASA 长期记忆。", "memory": result}

    def run_identity_merge_preflight(self, context: dict[str, Any], inputs: dict[str, Any]) -> dict[str, Any]:
        candidate = self._candidate(context)
        other_id = inputs.get("other_job_candidate_id")
        if not other_id:
            return self._blocked("身份合并预检需要另一条候选人关系。", ["other_job_candidate_id"], self._candidate_reference(candidate))
        other = build_candidate_context(self.service.db_path, int(other_id))
        comparison = {"left": candidate["identity"], "right": other["identity"], "same_name": candidate["identity"].get("name") == other["identity"].get("name"), "same_company": candidate["identity"].get("company") == other["identity"].get("company"), "same_title": candidate["identity"].get("title") == other["identity"].get("title")}
        comparison["allowed"] = comparison["same_name"] and (comparison["same_company"] or comparison["same_title"])
        return {"summary": "身份对比完成；该步骤不会执行合并。", "comparison": comparison,
                "references": self._candidate_reference(candidate) + self._candidate_reference(other),
                "artifacts": [self._artifact("identity_comparison", "候选人身份对比", content=json.dumps(comparison, ensure_ascii=False, indent=2), validation="passed" if comparison["allowed"] else "blocked")]}
