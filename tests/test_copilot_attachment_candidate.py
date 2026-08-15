"""上传简历附件的候选人识别（会话 copilot_da64e131281b 修复）。

场景：顾问上传候选人简历 docx（如「嘉驰国际-产能三佳-高级电气工程师-蔡敏明(1).docx」）
并发「给这个人选做份推荐报告」。候选人不在 DB（attachment-only）。

修复前：意图解释发生在附件证据收集之前，意图 LLM 看不到附件内容，
「这个人选」无法解析 → needs_clarification=true + missing_fields=["candidate_id"]
→ forced_answer「请问具体是哪位人选？」。

覆盖：
- 意图 LLM payload 携带附件摘要（文件名 + 截断正文 + untrusted 标注，过 sanitize 后姓名存活）；
- attachment-only 候选人：不追问人选、不走 create_goal，直接路由 recommendation_report
  并从附件简历生成嘉驰报告（source=attachment_resume、无 s6_assessment、提示顾问复核）；
- 附件姓名命中 DB：切到既有候选人流程（selected 变为该候选人，走 S6-3 既有门禁）；
- 简历正文字段解析与文件名姓名解析的纯函数行为。
"""

from __future__ import annotations

import hashlib
import json
import sqlite3
import subprocess
from pathlib import Path

import pytest
from docx import Document

from a_system_agent import AgentService, FakeLLM
from a_system_agent.capability_runtime import RecruitingCapabilityRuntime
from a_system_agent.copilot_impl import _attachment_candidate_identity
from a_system_agent.capability_runtime_delivery import _parse_jiashi_resume_fields
from test_a_system_agent_v1 import BASE_SCHEMA, fake_assessment

RESUME_TEXT = "\n".join(
    [
        "推荐岗位：高级电气工程师",
        "所属中心：产能三佳",
        "姓名：蔡敏明",
        "出生年月：1990.05",
        "性别：男",
        "婚育：已婚已育",
        "工作地址：江苏苏州",
        "综合年薪：35万",
        "期望薪酬：40万",
        "联系电话 13800138000",
        "推荐理由：8年电气设计经验，主导多条产线电气改造，稳定性好。",
        "教育背景：2008.09-2012.06 某某大学 电气工程及其自动化 本科",
        "工作经历：2012.07-至今 某某自动化有限公司 高级电气工程师 负责产线电气设计",
        "项目经历：2019 某某产线改造项目 电气负责人",
    ]
)
ATTACHMENT_FILE = "嘉驰国际-产能三佳-高级电气工程师-蔡敏明(1).docx"
REPORT_REQUEST = "给这个人选做份推荐报告，用嘉驰通用的模板"


@pytest.fixture()
def db_path(tmp_path: Path) -> Path:
    path = tmp_path / "agent.db"
    conn = sqlite3.connect(path)
    conn.executescript(BASE_SCHEMA)
    conn.execute("INSERT INTO clients VALUES (1,'长越科技')")
    conn.execute(
        "INSERT INTO jobs VALUES (10,1,'机械高级工程师','杭州','已发布','','','','','精密设备机械核心岗','2026-07-14')"
    )
    conn.execute(
        "INSERT INTO people VALUES (20,'张航','ASM中国集团公司','高级机械设计工程师','上海','本科','8年')"
    )
    conn.execute(
        "INSERT INTO job_candidates VALUES (30,10,20,'长越科技','机械高级工程师','new','','X1 待复核','正式流程','2026-07-14','40')"
    )
    conn.commit()
    conn.close()
    return path


class ClarifyingRecommendationLLM(FakeLLM):
    """模拟修复前的意图 LLM：动作识别为 recommendation，但缺 candidate_id 要追问。"""

    def __init__(self) -> None:
        super().__init__(fake_assessment(), chat_text="好的，已生成。", intent_understanding=self._intent)
        self.intent_payloads: list[dict] = []

    def _intent(self, payload: dict) -> dict | None:
        self.intent_payloads.append(json.loads(json.dumps(payload, ensure_ascii=False)))
        return {
            "speech_act": "propose",
            "action": "recommendation",
            "topic": "candidate_match",
            "objective": "给这个人选做份推荐报告",
            "target": {"type": "global", "id": None, "client": "", "label": ""},
            "constraints": [],
            "refers_to_previous": True,
            "confidence": 0.9,
            "needs_clarification": True,
            "missing_fields": ["candidate_id"],
            "clarification_question": "请问具体是哪位人选？",
        }


def _insert_attachment(db_path: Path, *, file_name: str, text: str, attachment_id: str = "att_resume") -> str:
    token = f"token-{attachment_id}-with-sufficient-length"
    conn = sqlite3.connect(db_path)
    conn.execute(
        """
        INSERT INTO agent_copilot_attachments
        (attachment_id,access_token_hash,file_name,file_type,size_bytes,content_sha256,extracted_text,status)
        VALUES (?,?,?,?,?,?,?,?)
        """,
        (
            attachment_id,
            hashlib.sha256(token.encode("utf-8")).hexdigest(),
            file_name,
            "docx",
            len(text),
            "test-sha256",
            text,
            "已读取附件正文。",
        ),
    )
    conn.commit()
    conn.close()
    return token


def _attachment_context(token: str, *, file_name: str, attachment_id: str = "att_resume") -> dict:
    return {
        "type": "global",
        "source": "asa_floating",
        "uploaded_attachments": [
            {
                "attachment_id": attachment_id,
                "access_token": token,
                "file_name": file_name,
                "file_type": "docx",
                "status": "已读取附件正文。",
            }
        ],
    }


@pytest.fixture()
def fake_jiashi_runtime(monkeypatch: pytest.MonkeyPatch, tmp_path: Path):
    """不依赖本机嘉驰模板/填充脚本：模板换成空 docx，子进程换成假填充+假审计。"""
    template = tmp_path / "嘉驰模板.docx"
    Document().save(template)
    monkeypatch.setattr("a_system_agent.capability_runtime_delivery.JIASHI_TEMPLATE", template)
    runs: list[list[str]] = []

    def fake_run(self, command, timeout: int = 300, *, cancel_check=None):
        runs.append([str(part) for part in command])
        if "--output" in command:
            output = Path(command[command.index("--output") + 1])
            data = json.loads(Path(command[command.index("--data") + 1]).read_text(encoding="utf-8"))
            doc = Document()
            doc.add_paragraph(f"{data.get('name')} 推荐报告")
            doc.add_paragraph(str(data.get("position") or ""))
            doc.save(output)
            return subprocess.CompletedProcess(command, 0, stdout="filled\n", stderr="")
        return subprocess.CompletedProcess(command, 0, stdout="AUDIT OK\n", stderr="")

    monkeypatch.setattr(RecruitingCapabilityRuntime, "_run", fake_run)
    return runs


def test_intent_payload_includes_uploaded_attachment_excerpt(db_path: Path, fake_jiashi_runtime) -> None:
    llm = ClarifyingRecommendationLLM()
    service = AgentService(db_path, llm)
    try:
        long_text = RESUME_TEXT + "\n" + ("补充经历" * 3000)
        token = _insert_attachment(db_path, file_name=ATTACHMENT_FILE, text=long_text)
        service.copilot(REPORT_REQUEST, context=_attachment_context(token, file_name=ATTACHMENT_FILE))
        assert llm.intent_payloads, "意图 LLM 未被调用"
        attachments = llm.intent_payloads[0].get("uploaded_attachments")
        assert attachments, "意图 payload 缺少上传附件摘要"
        item = attachments[0]
        assert "蔡敏明" in item["file_name"]
        assert item["untrusted_document_content"] is True
        assert len(item["text_excerpt"]) <= 4000
        assert "推荐理由" in item["text_excerpt"]
        # sanitize_payload 后姓名存活、手机号被脱敏
        assert "蔡敏明" in item["text_excerpt"]
        assert "13800138000" not in item["text_excerpt"]
    finally:
        service.close()


def test_attachment_only_candidate_requires_confirmation_without_clarification(
    db_path: Path, fake_jiashi_runtime
) -> None:
    llm = ClarifyingRecommendationLLM()
    service = AgentService(db_path, llm)
    try:
        token = _insert_attachment(db_path, file_name=ATTACHMENT_FILE, text=RESUME_TEXT)
        result = service.copilot(REPORT_REQUEST, context=_attachment_context(token, file_name=ATTACHMENT_FILE))
        assert result["ok"] is True
        assert "哪位人选" not in result["answer"]
        assert result["intent_understanding"]["needs_clarification"] is False
        # attachment-only：不走 create_goal，不追问姓名；首轮只创建持久化命令。
        assert result["goal_id"] is None
        assert result["workflow_id"] is None
        assert result["context"] == {"type": "candidate", "id": None}
        assert result["skill_runs"] == []
        assert fake_jiashi_runtime == []
        command = result.get("pending_command") or {}
        assert command.get("command_type") == "recommendation_report"
        assert (command.get("target") or {}).get("label") == "蔡敏明"
        assert (command.get("snapshot") or {}).get("attachment", {}).get("attachment_id") == "att_resume"
        interaction_card = result.get("interaction_card") or {}
        assert interaction_card.get("state") == "pending"
        assert (interaction_card.get("target") or {}).get("candidate") == "蔡敏明"
        assert (interaction_card.get("target") or {}).get("client") == "产能三佳"
        assert (interaction_card.get("target") or {}).get("job") == "高级电气工程师"

        preflight = service.preflight_copilot_command(str(command["command_id"]))
        decision = service.decide_copilot_command(
            str(command["command_id"]),
            decision="approve",
            confirmation_token=str(preflight["confirmation_token"]),
        )
        receipt = decision.get("receipt") or {}
        assert receipt.get("verified") is True
        assert receipt.get("state") == "已完成"
        artifacts = receipt.get("artifacts") or []
        assert artifacts, "推荐报告未产出 artifact"
        metadata = artifacts[0].get("metadata") or {}
        assert metadata.get("source") == "attachment_resume"
        assert "s6_assessment" not in metadata
        assert metadata.get("attached_to_candidate") is False
        assert "复核" in str(metadata.get("review_notice") or "")
        assert artifacts[0].get("file_path"), "artifact 缺少文件路径"
    finally:
        service.close()


def test_attachment_report_refresh_preserves_key_condition(
    db_path: Path, fake_jiashi_runtime,
) -> None:
    service = AgentService(db_path, ClarifyingRecommendationLLM())
    try:
        token = _insert_attachment(db_path, file_name=ATTACHMENT_FILE, text=RESUME_TEXT)
        result = service.copilot(
            REPORT_REQUEST,
            session_id="attachment-refresh-card",
            context=_attachment_context(token, file_name=ATTACHMENT_FILE),
        )
        command = dict(result.get("pending_command") or {})
        conn = sqlite3.connect(db_path)
        row = conn.execute(
            "SELECT id,structured_json FROM agent_copilot_messages WHERE session_id=? AND role='assistant' ORDER BY id DESC LIMIT 1",
            ("attachment-refresh-card",),
        ).fetchone()
        structured = json.loads(row[1])
        structured["interaction_card"]["key_conditions"] = ["用嘉驰通用的模板"]
        conn.execute(
            "UPDATE agent_copilot_messages SET structured_json=? WHERE id=?",
            (json.dumps(structured, ensure_ascii=False), row[0]),
        )
        conn.execute(
            "UPDATE agent_copilot_commands SET expires_at='2000-01-01 00:00:00' WHERE command_id=?",
            (command["command_id"],),
        )
        conn.commit()
        conn.close()

        refreshed = service.refresh_copilot_command(
            str(command["command_id"]),
            request_id="attachment-refresh-card-1",
            expected_command_hash=str(command["command_hash"]),
        )
        restored = service.get_copilot_session("attachment-refresh-card")
        assistant = next(message for message in reversed(restored["messages"]) if message["role"] == "assistant")
        assert (assistant.get("pending_command") or {}).get("command_id") == (
            refreshed.get("command") or {}
        ).get("command_id")
        assert (assistant.get("interaction_card") or {}).get("key_conditions") == ["用嘉驰通用的模板"]
    finally:
        service.close()


def test_attachment_candidate_matching_db_uses_existing_candidate_flow(db_path: Path) -> None:
    llm = ClarifyingRecommendationLLM()
    service = AgentService(db_path, llm)
    try:
        file_name = "嘉驰国际-长越科技-机械高级工程师-张航.docx"
        token = _insert_attachment(db_path, file_name=file_name, text=RESUME_TEXT.replace("蔡敏明", "张航"))
        result = service.copilot(REPORT_REQUEST, context=_attachment_context(token, file_name=file_name))
        assert "哪位人选" not in result["answer"]
        # DB 命中：等价于顾问手动选中该候选人，先进入统一报告命令确认。
        assert result["context"] == {"type": "candidate", "id": 30}
        command = result.get("pending_command") or {}
        assert command.get("command_type") == "recommendation_report"
        assert (command.get("target") or {}).get("id") == 30
        preflight = service.preflight_copilot_command(str(command["command_id"]))
        decision = service.decide_copilot_command(
            str(command["command_id"]),
            decision="approve",
            confirmation_token=str(preflight["confirmation_token"]),
        )
        # 既有 S6-3 门禁不变：无判人评估时不生成报告，并给出真实阻塞回执。
        receipt = decision.get("receipt") or {}
        assert receipt.get("state") == "流程阻塞"
        assert receipt.get("verified") is True
        assert "判人评估" in str(receipt.get("summary") or "")
    finally:
        service.close()


def test_attachment_candidate_identity_from_filename_and_db(db_path: Path) -> None:
    service = AgentService(db_path, FakeLLM(fake_assessment()))
    try:
        evidence = {
            "content_available": True,
            "items": [
                {
                    "file_name": ATTACHMENT_FILE,
                    "extracted_text": "一份没有姓名标签的简历正文",
                    "content_available": True,
                }
            ],
        }
        identity = _attachment_candidate_identity(service, evidence)
        assert identity["name"] == "蔡敏明"
        assert identity.get("customer") == "产能三佳"
        assert identity.get("position") == "高级电气工程师"
        assert identity.get("job_candidate_id") is None

        hit = _attachment_candidate_identity(
            service,
            {
                "content_available": True,
                "items": [{"file_name": "嘉驰国际-长越科技-机械高级工程师-张航.docx", "extracted_text": "姓名：张航"}],
            },
        )
        assert hit["name"] == "张航"
        assert hit["job_candidate_id"] == 30
    finally:
        service.close()


def test_parse_jiashi_resume_fields() -> None:
    fields = _parse_jiashi_resume_fields(RESUME_TEXT)
    assert fields["name"] == "蔡敏明"
    assert fields["position"] == "高级电气工程师"
    assert fields["customer"] == "产能三佳"
    assert fields["current_location"] == "江苏苏州"
    assert any("8年电气设计经验" in line for line in fields["consultant_comments"])
    assert any("电气工程及其自动化" in line for line in fields["education"])
    assert any("高级电气工程师" in line for line in fields["work_experience"])
    assert any("电气负责人" in line for line in fields["project_experience"])
