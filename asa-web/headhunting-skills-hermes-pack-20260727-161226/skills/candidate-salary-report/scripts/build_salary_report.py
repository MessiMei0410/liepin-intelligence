#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import sys
from collections import defaultdict
from dataclasses import dataclass
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BLACK = RGBColor(31, 35, 40)
MUTED = RGBColor(91, 101, 114)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F6F8FA"
NOTE_FILL = "FFF7E6"
TOTAL_FILL = "DDEBFF"
GRID = "D9DEE7"
LATIN_FONT = "Arial"
CJK_FONT = "Hiragino Sans GB"
FONT_FILE = "/System/Library/Fonts/Hiragino Sans GB.ttc"


def add_documents_helpers_to_path() -> None:
    candidates = [
        Path.home() / ".codex/plugins/cache/openai-primary-runtime/documents",
        Path.home() / ".codex/plugins/cache/openai-bundled/documents",
    ]
    for root in candidates:
        if not root.exists():
            continue
        for path in sorted(root.glob("*/skills/documents/scripts"), reverse=True):
            sys.path.append(str(path))
            return


add_documents_helpers_to_path()
try:
    from table_geometry import apply_table_geometry
except Exception:
    apply_table_geometry = None


@dataclass
class Record:
    month: str
    category: str
    subtype: str
    payer: str
    income: Decimal
    tax: Decimal
    note: str
    source: str


@dataclass
class BankEntry:
    date: str
    summary: str
    payer: str
    account: str
    amount: Decimal
    balance: Decimal | None
    note: str
    source: str
    include_in_main_totals: bool


def dec(value: Any) -> Decimal:
    if value in (None, ""):
        return Decimal("0")
    return Decimal(str(value))


def money(value: Decimal) -> str:
    value = value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return f"{value:,.2f}"


def pct(value: Decimal) -> str:
    return f"{value.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)}%"


def month_label(month: str) -> str:
    return month.replace("-", ".")


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = LATIN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    rfonts.set(qn("w:cs"), LATIN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def format_paragraph(paragraph, *, before=0, after=6, line_spacing=1.1, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    if align is not None:
        paragraph.alignment = align


def add_run(paragraph, text, *, size=11, bold=False, color=BLACK, italic=False):
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return run


def add_paragraph(doc, text="", *, size=11, bold=False, color=BLACK, before=0, after=6, align=None, italic=False):
    p = doc.add_paragraph()
    format_paragraph(p, before=before, after=after, align=align)
    if text:
        add_run(p, text, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        format_paragraph(p, before=14, after=7, line_spacing=1.15)
        add_run(p, text, size=16, bold=True, color=BLUE)
    elif level == 2:
        format_paragraph(p, before=10, after=5, line_spacing=1.12)
        add_run(p, text, size=13, bold=True, color=BLUE)
    else:
        format_paragraph(p, before=8, after=4, line_spacing=1.1)
        add_run(p, text, size=12, bold=True, color=DARK_BLUE)
    return p


def set_cell_text(cell, text, *, size=9.2, bold=False, color=BLACK, align=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.text = ""
    format_paragraph(p, before=0, after=0, line_spacing=1.08, align=align)
    add_run(p, text, size=size, bold=bold, color=color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_table(table, header_rows=1):
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            if r_idx < header_rows:
                shade_cell(cell, HEADER_FILL)
            elif r_idx % 2 == 0:
                shade_cell(cell, LIGHT_FILL)


def add_table(doc, headers, rows, widths, font_size=9.0, total_row=False):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_text(hdr[i], text, size=font_size, bold=True, color=BLACK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        is_total = total_row and row_idx == len(rows) - 1
        for i, text in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER
            if len(row_data) > 2 and i in (1, len(row_data) - 1):
                align = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(text), size=font_size, bold=is_total, color=BLACK, align=align)
            if is_total:
                shade_cell(cells[i], TOTAL_FILL)
    if apply_table_geometry:
        apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=120)
    style_table(table)
    add_paragraph(doc, "", after=3)
    return table


def add_note_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, NOTE_FILL)
    set_cell_border(cell, color="E6D8A8")
    set_cell_text(cell, text, size=9.5, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT)
    if apply_table_geometry:
        apply_table_geometry(table, [9360], table_width_dxa=9360, indent_dxa=120)
    add_paragraph(doc, "", after=4)


def parse_records(data: dict[str, Any]) -> list[Record]:
    records = []
    for item in data.get("records", []):
        records.append(
            Record(
                month=item["month"],
                category=item.get("category", "其他"),
                subtype=item.get("subtype", ""),
                payer=item.get("payer", ""),
                income=dec(item.get("income")),
                tax=dec(item.get("tax")),
                note=item.get("note", ""),
                source=item.get("source", ""),
            )
        )
    return records


def parse_bank_entries(data: dict[str, Any]) -> list[BankEntry]:
    entries = []
    for item in data.get("bank_entries", []):
        balance = item.get("balance")
        entries.append(
            BankEntry(
                date=item.get("date", ""),
                summary=item.get("summary", ""),
                payer=item.get("payer", ""),
                account=item.get("account", ""),
                amount=dec(item.get("amount")),
                balance=dec(balance) if balance not in (None, "") else None,
                note=item.get("note", ""),
                source=item.get("source", ""),
                include_in_main_totals=bool(item.get("include_in_main_totals", False)),
            )
        )
    return entries


def totals(records: list[Record]):
    total_income = sum((r.income for r in records), Decimal("0"))
    total_tax = sum((r.tax for r in records), Decimal("0"))
    by_cat = defaultdict(lambda: {"income": Decimal("0"), "tax": Decimal("0")})
    by_month = defaultdict(lambda: {"income": Decimal("0"), "tax": Decimal("0"), "items": []})
    for record in records:
        by_cat[record.category]["income"] += record.income
        by_cat[record.category]["tax"] += record.tax
        by_month[record.month]["income"] += record.income
        by_month[record.month]["tax"] += record.tax
        by_month[record.month]["items"].append(record)
    return total_income, total_tax, by_cat, by_month


def records_with_included_bank_entries(records: list[Record], bank_entries: list[BankEntry]) -> list[Record]:
    combined = list(records)
    for entry in bank_entries:
        if not entry.include_in_main_totals:
            continue
        month = entry.date[:7] if len(entry.date) >= 7 else "待补充"
        combined.append(
            Record(
                month=month,
                category="银行入账",
                subtype=entry.summary,
                payer=entry.payer,
                income=entry.amount,
                tax=Decimal("0"),
                note=entry.note or "用户指定并入主统计，税额按0处理",
                source=entry.source,
            )
        )
    return combined


def infer_findings(data, records, bank_entries, total_income, total_tax, by_cat, by_month):
    findings = []
    salary_income = by_cat["工资薪金"]["income"]
    salary_months = sorted(m for m, entry in by_month.items() if any(i.category == "工资薪金" for i in entry["items"]))
    if salary_months:
        salary_avg = salary_income / Decimal(len(salary_months))
        findings.append(
            {
                "title": "固定薪资稳定性",
                "text": f"工资薪金覆盖{len(salary_months)}个月，合计{money(salary_income)}元，月均约{money(salary_avg)}元；需结合工资单确认固定/浮动构成。",
            }
        )
    for category, values in sorted(by_cat.items()):
        if category != "工资薪金" and values["income"] > 0:
            findings.append(
                {
                    "title": f"{category}口径",
                    "text": f"{category}合计{money(values['income'])}元，税额{money(values['tax'])}元；建议确认是否可持续、是否属于谈薪总包口径。",
                }
            )
    if by_month:
        salary_month_values = []
        for month, entry in by_month.items():
            salary = sum((r.income for r in entry["items"] if r.category == "工资薪金"), Decimal("0"))
            if salary:
                salary_month_values.append((month, salary))
        if salary_month_values:
            avg = sum(v for _, v in salary_month_values) / Decimal(len(salary_month_values))
            for month, value in salary_month_values:
                if value < avg * Decimal("0.85") or value > avg * Decimal("1.15"):
                    findings.append(
                        {
                            "title": f"{month_label(month)}工资波动",
                            "text": f"当月工资薪金{money(value)}元，相对平均值{money(avg)}元存在明显差异，需确认出勤、扣款、发放周期或薪资结构变化。",
                        }
                    )
    if bank_entries:
        amount = sum((entry.amount for entry in bank_entries), Decimal("0"))
        findings.append(
            {
                "title": "银行流水补充",
                "text": f"银行入账证据合计{money(amount)}元，默认作为补充材料单独列示；需确认与税务记录对应关系、纳税处理和服务期/返还义务。",
            }
        )
    findings.extend(data.get("custom_findings", []))
    if not findings:
        findings.append({"title": "薪资口径", "text": "当前证据不足以形成明确判断，需补充工资单、银行流水或个税导出记录。"})
    return findings


def default_questions(data, records, bank_entries):
    questions = [
        {"topic": "固定薪资", "question": "请确认当前税前月固定工资金额，以及其中基本工资、绩效、津贴、补贴、预发奖金的构成。"},
        {"topic": "税后实发", "question": "请补充近12个月工资卡流水或工资单，以核对社保、公积金、专项扣除和其他扣款后的实发金额。"},
        {"topic": "奖金规则", "question": "请确认奖金对应年度、发放月份、计算规则、绩效系数、是否仍有未发放或递延部分。"},
        {"topic": "长期激励", "question": "请确认是否有股票、RSU、期权、虚拟股、分红、项目奖、专利奖、留任奖或递延奖金。"},
        {"topic": "离职限制", "question": "请确认是否存在竞业限制、服务期、培训协议、补贴返还、奖金返还或未解锁权益损失。"},
        {"topic": "谈薪预期", "question": "请确认目标年包、最低可接受现金部分、结构偏好、签字费/保底奖金诉求和最早到岗时间。"},
    ]
    if any(r.category != "工资薪金" for r in records):
        questions.insert(
            2,
            {"topic": "非常规收入", "question": "请解释奖金、偶然所得、补贴或其他一次性收入的来源、持续性、纳税情况和是否可计入谈薪口径。"},
        )
    if bank_entries:
        questions.insert(
            3,
            {"topic": "银行入账", "question": "请解释银行入账的具体性质，是否与个税记录对应，是否已纳税，是否存在服务期或返还条件。"},
        )
    questions.extend(data.get("custom_questions", []))
    return questions


def make_chart(path: Path, by_month):
    months = sorted(by_month)
    if not months:
        return
    width, height = 1400, 720
    margin_l, margin_r, margin_t, margin_b = 105, 60, 70, 120
    plot_w = width - margin_l - margin_r
    plot_h = height - margin_t - margin_b
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        font_title = ImageFont.truetype(FONT_FILE, 36)
        font_small = ImageFont.truetype(FONT_FILE, 20)
    except Exception:
        font_title = ImageFont.load_default()
        font_small = ImageFont.load_default()

    draw.text((margin_l, 22), "收入趋势（申报收入口径）", fill=(31, 35, 40), font=font_title)
    max_income = max(by_month[m]["income"] for m in months)
    y_max = Decimal("1")
    while y_max < max_income:
        y_max *= Decimal("2")
        if y_max < max_income:
            y_max = (y_max / Decimal("2")) * Decimal("5")
    if y_max < Decimal("100000"):
        y_max = Decimal("100000")
    for i in range(0, 7):
        val = y_max * Decimal(i) / Decimal(6)
        y = margin_t + plot_h - int((float(val) / float(y_max)) * plot_h)
        draw.line((margin_l, y, width - margin_r, y), fill=(225, 230, 236), width=2)
        draw.text((18, y - 12), f"{int(val / Decimal('10000'))}万", fill=(91, 101, 114), font=font_small)

    bar_gap = 18
    bar_w = max(18, int((plot_w - bar_gap * (len(months) + 1)) / len(months)))
    base_color = (46, 116, 181)
    extra_color = (230, 126, 34)
    for idx, month in enumerate(months):
        x0 = margin_l + bar_gap + idx * (bar_w + bar_gap)
        items = by_month[month]["items"]
        salary = sum((item.income for item in items if item.category == "工资薪金"), Decimal("0"))
        extra = by_month[month]["income"] - salary
        y_base = margin_t + plot_h
        salary_h = int((float(salary) / float(y_max)) * plot_h)
        extra_h = int((float(extra) / float(y_max)) * plot_h)
        draw.rectangle((x0, y_base - salary_h, x0 + bar_w, y_base), fill=base_color)
        if extra_h > 0:
            draw.rectangle((x0, y_base - salary_h - extra_h, x0 + bar_w, y_base - salary_h), fill=extra_color)
        draw.text((x0 - 6, y_base + 18), month[2:], fill=(31, 35, 40), font=font_small)
    draw.rectangle((margin_l, margin_t, width - margin_r, margin_t + plot_h), outline=(160, 170, 185), width=2)
    draw.rectangle((margin_l, height - 58, margin_l + 28, height - 34), fill=base_color)
    draw.text((margin_l + 38, height - 62), "工资薪金", fill=(31, 35, 40), font=font_small)
    draw.rectangle((margin_l + 185, height - 58, margin_l + 213, height - 34), fill=extra_color)
    draw.text((margin_l + 223, height - 62), "其他收入", fill=(31, 35, 40), font=font_small)
    image.save(path)


def setup_document(title: str):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    header_p = section.header.paragraphs[0]
    header_p.text = ""
    format_paragraph(header_p, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(header_p, title, size=9, color=MUTED)
    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    format_paragraph(footer_p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(footer_p, "金额单位：人民币元；数据来源：用户提供材料", size=8.5, color=MUTED)
    return doc


def build_report(data, output_path: Path, work_dir: Path):
    records = parse_records(data)
    bank_entries = parse_bank_entries(data)
    records_for_totals = records_with_included_bank_entries(records, bank_entries)
    total_income, total_tax, by_cat, by_month = totals(records_for_totals)
    net_total = total_income - total_tax
    effective_tax = (total_tax / total_income * Decimal("100")) if total_income else Decimal("0")
    findings = infer_findings(data, records_for_totals, bank_entries, total_income, total_tax, by_cat, by_month)
    questions = default_questions(data, records, bank_entries)

    chart_path = work_dir / "salary_trend.png"
    make_chart(chart_path, by_month)

    title = data.get("title", "候选人薪资流水报告")
    period = data.get("period", {})
    period_text = f"{period.get('start', '')} 至 {period.get('end', '')}".strip(" 至")
    period_label = period.get("label", "统计周期")

    doc = setup_document(title)
    doc.core_properties.title = title
    doc.core_properties.subject = "候选人薪资统计与核实报告"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""

    p = doc.add_paragraph()
    format_paragraph(p, before=0, after=4)
    add_run(p, title, size=24, bold=True, color=BLACK)
    subtitle = f"{period_label}统计（{period_text}） | 数据来源：{data.get('source_summary', '用户提供材料')}"
    add_paragraph(doc, subtitle, size=11.5, color=MUTED, after=14)

    meta_rows = [
        ["统计周期", period_text or "待补充", "主要扣缴义务人", data.get("primary_payer", "待补充")],
        ["核心口径", "以结构化记录中的收入/税额为主", "报告日期", data.get("report_date", "待补充")],
        ["特别说明", "银行入账默认单独列示，未确认前不并入主统计", "候选人姓名", data.get("candidate_name", "待补充")],
    ]
    add_table(doc, ["项目", "内容", "项目", "内容"], meta_rows, [1350, 3590, 1450, 2970], font_size=8.8)

    add_heading(doc, "一、核心结论", 1)
    salary_income = by_cat["工资薪金"]["income"]
    salary_month_count = len({r.month for r in records if r.category == "工资薪金"}) or 1
    metric_rows = [
        ["申报收入合计", money(total_income), "已申报个税合计", money(total_tax)],
        ["税后估算（仅扣个税）", money(net_total), "综合个税率", pct(effective_tax)],
        ["工资薪金合计", money(salary_income), "工资薪金月均", money(salary_income / Decimal(salary_month_count))],
    ]
    add_table(doc, ["指标", "金额/比例", "指标", "金额/比例"], metric_rows, [2200, 2480, 2200, 2480], font_size=9.2)
    add_note_box(doc, "口径提示：税后估算=收入-已申报税额，仅反映个人所得税后的粗略现金水平，未核对社保、公积金、专项扣除、工资卡实发金额及任何未在证据中体现的长期激励。")

    add_table(doc, ["观察点", "判断"], [[f["title"], f["text"]] for f in findings], [2100, 7260], font_size=9.2)

    add_heading(doc, "二、收入统计", 1)
    cat_rows = []
    for category in sorted(by_cat.keys(), key=lambda c: (c != "工资薪金", c)):
        inc = by_cat[category]["income"]
        tax = by_cat[category]["tax"]
        share = pct(inc / total_income * Decimal("100")) if total_income else "0.00%"
        cat_rows.append([category, money(inc), money(tax), money(inc - tax), share])
    cat_rows.append(["合计", money(total_income), money(total_tax), money(net_total), "100.00%" if total_income else "0.00%"])
    add_table(doc, ["收入类别", "收入合计", "已申报税额", "税后估算", "收入占比"], cat_rows, [1900, 1900, 1900, 1900, 1760], font_size=9.0, total_row=True)

    if chart_path.exists():
        p = add_paragraph(doc, "图示：蓝色为工资薪金，橙色为其他收入。", size=9.5, color=MUTED, before=0, after=4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(chart_path), width=Inches(6.3))
        add_paragraph(doc, "", after=2)

    monthly_rows = []
    for month in sorted(by_month):
        entry = by_month[month]
        parts = [f"{item.category} {money(item.income)}" for item in entry["items"]]
        notes = [item.note for item in entry["items"] if item.note and item.category != "工资薪金"]
        monthly_rows.append([
            month,
            "；".join(parts),
            money(entry["income"]),
            money(entry["tax"]),
            money(entry["income"] - entry["tax"]),
            "；".join(notes) if notes else "-",
        ])
    add_table(doc, ["月份", "收入构成", "收入合计", "已申报税额", "税后估算", "备注"], monthly_rows, [1050, 2950, 1350, 1350, 1350, 1310], font_size=7.9)

    if bank_entries:
        add_heading(doc, "三、银行入账补充", 1)
        bank_rows = []
        for entry in bank_entries:
            treatment = "并入主统计" if entry.include_in_main_totals else "单独列示，待核实"
            bank_rows.append([
                entry.date,
                entry.payer,
                entry.summary,
                money(entry.amount),
                treatment,
                entry.note or entry.source or "-",
            ])
        add_table(doc, ["日期", "付款方", "摘要", "金额", "口径处理", "备注"], bank_rows, [1100, 2200, 1600, 1350, 1500, 1610], font_size=8.2)

    add_heading(doc, "四、需重点核实事项", 1)
    risk_rows = [[q["topic"], q["question"], "待候选人确认"] for q in questions]
    add_table(doc, ["事项", "建议核实问题", "当前状态"], risk_rows, [1700, 5960, 1700], font_size=8.4)

    if data.get("known_context"):
        add_heading(doc, "五、补充背景", 1)
        add_table(doc, ["序号", "背景信息"], [[str(i + 1), text] for i, text in enumerate(data["known_context"])], [700, 8660], font_size=9.0)

    add_heading(doc, "报告限制", 1)
    add_note_box(doc, "本报告基于用户提供材料整理，不构成审计、税务或法律意见。若用于Offer定薪、背调或薪资证明审核，建议以候选人提供的工资单、银行流水、税务App导出记录和公司奖金/股权文件进行交叉验证。")

    images = data.get("source_images", [])
    if images:
        doc.add_page_break()
        add_heading(doc, "附录：原始材料", 1)
        for idx, item in enumerate(images):
            image_path = Path(item["path"]).expanduser()
            if not image_path.exists():
                continue
            if idx > 0:
                doc.add_page_break()
            caption = item.get("caption") or item.get("label") or image_path.name
            add_paragraph(doc, caption, size=10.5, bold=True, color=BLACK, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                with Image.open(image_path) as im:
                    ratio = im.width / im.height
                width = min(6.1, 7.4 * ratio)
            except Exception:
                width = 5.2
            p.add_run().add_picture(str(image_path), width=Inches(width))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def build_checklist(data, output_path: Path):
    records = parse_records(data)
    bank_entries = parse_bank_entries(data)
    questions = default_questions(data, records, bank_entries)
    period = data.get("period", {})
    period_text = f"{period.get('start', '')} 至 {period.get('end', '')}".strip(" 至")

    doc = setup_document("薪资沟通核实清单")
    p = doc.add_paragraph()
    format_paragraph(p, before=0, after=4)
    add_run(p, "人选薪资沟通核实清单", size=22, bold=True, color=BLACK)
    add_paragraph(doc, f"基于候选人薪资材料整理｜统计口径：{period_text or '待补充'}", size=10.5, color=MUTED, after=10)
    add_note_box(doc, "使用方式：逐项向人选确认，右侧“记录/结论”用于补充人选原话、证明材料、风险判断和下一步动作。")

    priority = questions[: min(5, len(questions))]
    add_heading(doc, "优先核实事项", 1)
    add_table(
        doc,
        ["优先级", "核实事项", "为什么要问"],
        [[str(i + 1), item["topic"], item["question"]] for i, item in enumerate(priority)],
        [760, 2300, 6300],
        font_size=9.0,
    )

    sections = [
        (
            "一、统计口径与固定薪资",
            [
                ("最近周期", f"是否认可当前统计周期（{period_text or '待补充'}）及收入/税额口径？"),
                ("月薪构成", "当前月固定工资由哪些部分组成？基本工资、岗位工资、绩效、津贴、补贴分别是多少？"),
                ("固定/浮动", "月薪是否为固定税前月薪，还是包含浮动项、补贴或其他发放？"),
                ("社保公积金", "社保、公积金缴纳城市、基数、个人比例、公司比例分别是多少？"),
                ("专项扣除", "是否有专项附加扣除或其他扣除项，对个税口径影响是否较大？"),
            ],
        ),
        (
            "二、奖金、补贴与非常规收入",
            [
                ("奖金规则", "全年一次性奖金对应哪一年度？发放规则、绩效系数、未发部分如何？"),
                ("补贴性质", "补贴、偶然所得、政府奖励或其他一次性收入的性质、持续性和纳税情况如何？"),
                ("银行入账", "银行入账是否与个税记录对应？是否有通知、纳税证明、服务期或返还条件？"),
                ("离职损失", "如果现在离职，会损失多少奖金、绩效、递延收入或其他现金收入？"),
            ],
        ),
        (
            "三、真实年包、长期激励与期望薪资",
            [
                ("完整年包", "过去一个完整自然年真实总包是多少？今年如果不跳槽预计全年总包是多少？"),
                ("长期激励", "是否有股票、RSU、虚拟股、分红或其他长期激励？未归属价值和归属时间如何？"),
                ("薪资底线", "当前期望新机会税前年包目标是多少？最低能接受的税前年包或现金部分是多少？"),
                ("结构偏好", "更看重base、奖金、签字费、保底奖金，还是长期激励？"),
            ],
        ),
        (
            "四、决策因素与到岗风险",
            [
                ("其他机会", "手上是否还有其他机会或offer？分别到了什么阶段？"),
                ("限制条款", "是否存在竞业限制、脱敏期、服务期或培训协议？"),
                ("离职周期", "正常离职周期多久？最早可入职时间是哪一天？"),
                ("挽留风险", "直属领导是否可能挽留、加薪或卡离职？"),
            ],
        ),
        (
            "五、需补充证明材料",
            [
                ("工资单", "是否愿意补充最近12个月工资单或薪资条？"),
                ("个税材料", "是否愿意补充完整个税App截图、导出记录或纳税记录？"),
                ("银行流水", "是否能提供关键银行流水原件或对应通知？"),
                ("长期激励", "是否能提供股票、RSU、长期激励授予或归属截图？"),
                ("合同文件", "是否能提供劳动合同薪资页、调薪通知或补贴协议？"),
            ],
        ),
    ]

    for title, rows in sections:
        add_heading(doc, title, 1)
        table_rows = [["□", item, question, ""] for item, question in rows]
        add_table(doc, ["勾选", "核实项", "沟通问题", "记录/结论"], table_rows, [620, 1900, 4920, 1920], font_size=8.8)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Build candidate salary report DOCX files from structured JSON.")
    parser.add_argument("--input", required=True, help="Path to salary report JSON input.")
    parser.add_argument("--output-dir", required=True, help="Directory for generated DOCX files.")
    parser.add_argument("--report-name", default="候选人薪资流水报告.docx")
    parser.add_argument("--checklist-name", default="薪资沟通核实清单.docx")
    parser.add_argument("--report-only", action="store_true")
    parser.add_argument("--checklist-only", action="store_true")
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()
    work_dir = output_dir / "_salary_report_work"
    work_dir.mkdir(parents=True, exist_ok=True)
    data = json.loads(input_path.read_text(encoding="utf-8"))

    outputs = []
    if not args.checklist_only:
        outputs.append(build_report(data, output_dir / args.report_name, work_dir))
    if not args.report_only:
        outputs.append(build_checklist(data, output_dir / args.checklist_name))

    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
