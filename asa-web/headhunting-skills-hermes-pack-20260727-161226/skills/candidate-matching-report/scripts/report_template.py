#!/usr/bin/env python3
"""
候选人-岗位匹配报告 .docx 生成模板 (v1.4.0)

用法：
  python3 report_template.py --data '<json_string>' --output '/path/to/report.docx'

JSON 结构示例见底部 EXAMPLE_DATA。

设计目标：
  - 把所有 docx 生成逻辑集中在此文件，彻底消除 Hermes cron 任务里
    内联 heredoc Python 脚本时中文引号引发的 SyntaxError 问题。
  - 调用方只传结构化数据，不再生成 Python 代码字符串。
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("缺少 python-docx，请先执行: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── 颜色常量 ──────────────────────────────────────────────────
C_TITLE      = RGBColor(0x1A, 0x52, 0x76)   # #1A5276 深蓝大标题
C_H2         = RGBColor(0x2C, 0x3E, 0x50)   # #2C3E50 二级标题
C_H3         = RGBColor(0x34, 0x49, 0x5E)   # #34495E 三级标题
C_BODY       = RGBColor(0x33, 0x33, 0x33)   # #333333 正文
C_MUTED      = RGBColor(0x78, 0x78, 0x78)   # #787878 辅助文字
C_RISK_HIGH  = RGBColor(0xE7, 0x4C, 0x3C)   # #E74C3C 🔴 高风险
C_RISK_MID   = RGBColor(0xF3, 0x9C, 0x12)   # #F39C12 🟡 中风险
C_RISK_LOW   = RGBColor(0x27, 0xAE, 0x60)   # #27AE60 🟢 低风险
C_TH_BG      = RGBColor(0x1A, 0x52, 0x76)   # 表头背景
C_ROW_ALT    = RGBColor(0xEB, 0xF5, 0xFB)   # 交替行背景


def set_cn_font(run, size_pt=10, bold=False, color: RGBColor = C_BODY):
    """统一设置中英文字体、字号、颜色。"""
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color
    # 中文字体需通过 XML 设置 eastAsia
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "微软雅黑")
    rPr.insert(0, rFonts)


def add_heading(doc: Document, text: str, level: int = 1):
    """添加标题段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    size_map = {1: (16, C_TITLE), 2: (12, C_H2), 3: (10.5, C_H3)}
    size, color = size_map.get(level, (10, C_BODY))
    set_cn_font(run, size_pt=size, bold=True, color=color)
    return p


def add_body(doc: Document, text: str, color: RGBColor = C_BODY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, size_pt=10, color=color)
    return p


def shade_cell(cell, rgb: RGBColor):
    """给表格单元格着色。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    # python-docx exposes RGBColor as a tuple-like value across versions.
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc: Document, headers: list, rows: list):
    """添加格式化表格（表头深蓝，数据行交替浅蓝）。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], C_TH_BG)
        for run in hdr_cells[i].paragraphs[0].runs:
            set_cn_font(run, size_pt=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # 数据行
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            if r_idx % 2 == 1:
                shade_cell(cells[c_idx], C_ROW_ALT)
            for run in cells[c_idx].paragraphs[0].runs:
                set_cn_font(run, size_pt=10)

    doc.add_paragraph()


def score_bar(pct: int) -> str:
    """生成评分条字符串，pct 范围 0-100。"""
    n = min(20, max(0, pct // 5))
    return f"{'█' * n}{'░' * (20 - n)}  {pct}%"


def risk_color(level: str) -> RGBColor:
    level = level.strip().upper()
    if "高" in level or "RED" in level:
        return C_RISK_HIGH
    if "中" in level or "YELLOW" in level or "MID" in level:
        return C_RISK_MID
    return C_RISK_LOW


def build_report(data: dict, output_path: str):
    doc = Document()

    # ── 页边距 ─────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    candidate = data.get("candidate", "候选人")
    company   = data.get("company", "企业")
    position  = data.get("position", "岗位")
    version   = data.get("plugin_version", "v1.4.0")
    date_str  = datetime.now().strftime("%Y-%m-%d")

    # ── 封面 ───────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run(f"人选匹配分析报告")
    set_cn_font(r, size_pt=26, bold=True, color=C_TITLE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(f"{candidate}  ×  {company} · {position}")
    set_cn_font(r2, size_pt=14, color=C_MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = meta.add_run(f"生成日期：{date_str}    插件版本：{version}")
    set_cn_font(r3, size_pt=9, color=C_MUTED)

    doc.add_paragraph()

    jd_warn = data.get("jd_incomplete_warning", "")
    if jd_warn:
        add_body(doc, f"⚠️ {jd_warn}", color=C_RISK_MID)
        doc.add_paragraph()

    # ── Section 1：硬性门槛对照 ────────────────────────────────
    add_heading(doc, "一、硬性门槛对照", level=1)
    gate_rows = data.get("hard_gates", [])
    if gate_rows:
        add_table(doc, ["要求项", "候选人情况", "判定"], gate_rows)

    # ── Section 2：岗位职责逐条匹配 ───────────────────────────
    add_heading(doc, "二、岗位职责逐条匹配", level=1)
    target_biz = data.get("target_company_background", "")
    if target_biz:
        add_heading(doc, "目标公司背景", level=3)
        add_body(doc, target_biz)
    for item in data.get("responsibility_matches", []):
        stars = "★" * item.get("stars", 3) + "☆" * (5 - item.get("stars", 3))
        add_heading(doc, f"{stars}  {item.get('duty', '')}", level=3)
        add_body(doc, item.get("evidence", ""))

    # ── Section 3：加分项核查 ──────────────────────────────────
    add_heading(doc, "三、加分项核查", level=1)
    bonus_rows = data.get("bonus_items", [])
    if bonus_rows:
        add_table(doc, ["加分项", "候选人情况", "判定"], bonus_rows)

    # ── Section 4：潜在风险与关注点 ───────────────────────────
    add_heading(doc, "四、潜在风险与关注点", level=1)
    for risk in data.get("risks", []):
        level  = risk.get("level", "低")
        title  = risk.get("title", "")
        desc   = risk.get("description", "")
        verify = risk.get("verify", "")
        color  = risk_color(level)
        add_heading(doc, f"[{level}风险] {title}", level=3)
        p = doc.add_paragraph()
        r_desc = p.add_run(desc)
        set_cn_font(r_desc, size_pt=10, color=color)
        if verify:
            add_body(doc, f"验证建议：{verify}", color=C_MUTED)

    # ── Section 5：综合匹配度评分 ─────────────────────────────
    add_heading(doc, "五、综合匹配度评分", level=1)
    scores = data.get("scores", {})
    total  = data.get("total_score", 0)
    for dim, pct in scores.items():
        p = doc.add_paragraph()
        r = p.add_run(f"{dim:<12}  {score_bar(pct)}")
        set_cn_font(r, size_pt=10)

    p_total = doc.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_total = p_total.add_run(f"综合匹配度  {total}%")
    set_cn_font(r_total, size_pt=22, bold=True, color=C_TITLE)
    doc.add_paragraph()

    # ── Section 6：面试考察建议 ────────────────────────────────
    add_heading(doc, "六、面试考察建议", level=1)
    interview_rows = data.get("interview_suggestions", [])
    if interview_rows:
        add_table(doc, ["重点考察维度", "建议问题方向"], interview_rows)

    # ── Section 7：初次电话沟通指南 ───────────────────────────
    add_heading(doc, "七、初次电话沟通指南", level=1)
    phone = data.get("phone_guide", {})
    for stage, script in phone.get("stages", {}).items():
        add_heading(doc, stage, level=3)
        add_body(doc, script)

    checklist_rows = phone.get("checklist", [])
    if checklist_rows:
        add_heading(doc, "电话后评估清单", level=3)
        add_table(doc, ["考察维度", "通过标准", "红灯信号"], checklist_rows)

    taboos = phone.get("taboos", [])
    if taboos:
        add_heading(doc, "禁忌提醒", level=3)
        for t in taboos:
            add_body(doc, f"❌  {t}")

    # ── Section 8：结论 ────────────────────────────────────────
    add_heading(doc, "八、结论", level=1)
    verdict = data.get("verdict", "")
    p_verdict = doc.add_paragraph()
    p_verdict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_v = p_verdict.add_run(verdict)
    set_cn_font(r_v, size_pt=18, bold=True, color=C_TITLE)
    add_body(doc, data.get("conclusion_summary", ""))

    # ── 写入文件 ───────────────────────────────────────────────
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[report_template] 报告已保存：{output_path}")


# ── CLI 入口 ──────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="生成匹配分析 .docx 报告")
    parser.add_argument("--data",   required=True, help="JSON 数据字符串（或 @file.json）")
    parser.add_argument("--output", required=True, help="输出 .docx 路径")
    args = parser.parse_args()

    raw = args.data
    if raw.startswith("@"):
        with open(raw[1:], "r", encoding="utf-8") as f:
            raw = f.read()
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"JSON 解析失败：{e}", file=sys.stderr)
        sys.exit(1)

    build_report(data, args.output)


# ── 示例数据（供开发调试用）─────────────────────────────────
EXAMPLE_DATA = {
    "candidate": "张三",
    "company": "示例科技",
    "position": "高级产品经理",
    "plugin_version": "v1.4.0",
    "hard_gates": [
        ["学历", "本科 计算机科学", "✅"],
        ["工作年限", "7年", "✅"],
        ["目前公司主营业务", "B2B SaaS 企业软件（与目标岗位高度吻合）", "✅"],
        ["薪资区间", "当前约 40-45w，目标岗预算 45-55w，在范围内", "✅"],
        ["跳槽频率", "平均在职 24 个月，稳定性良好", "✅"],
    ],
    "target_company_background": "目标公司为国内头部新能源汽车智能座舱软件供应商，B端客户为主机厂，当前处于快速扩张阶段（C轮，年营收约3亿）。",
    "responsibility_matches": [
        {"duty": "负责产品路线图规划与迭代", "stars": 5, "evidence": "在前公司主导过3个核心产品从0到1，年活用户超100万。"},
        {"duty": "跨部门协作推动需求落地", "stars": 4, "evidence": "曾牵头研发+设计+商务三方项目，按时交付率92%。"},
    ],
    "bonus_items": [
        ["有汽车行业背景", "无直接汽车行业经验，但有IoT平台经验", "⚠️"],
        ["CPA或PMP认证", "持有PMP证书", "✅"],
    ],
    "risks": [
        {
            "level": "高",
            "title": "跨行业风险（无汽车行业经验）",
            "description": "候选人来自SaaS软件行业，汽车行业供应链流程、主机厂对接经验为零，适应期预计3-6个月。",
            "verify": "电话中询问：是否有意愿系统学习汽车行业知识？对主机厂甲方文化有何预期？"
        },
        {
            "level": "中",
            "title": "薪资期望待确认",
            "description": "简历未披露期望薪资，当前包估算40-45w，目标岗预算上限55w，有空间但需核实。",
            "verify": "直接询问当前税前总包及期望，确认是否接受股权部分。"
        },
    ],
    "scores": {
        "教育背景":     85,
        "技术深度":     80,
        "行业经验":     55,
        "产品思维":     90,
        "主营业务匹配": 60,
    },
    "total_score": 74,
    "interview_suggestions": [
        ["产品规划能力", "请描述你主导过的最复杂产品决策，当时如何在资源约束下做取舍？"],
        ["行业适应意愿", "你对汽车行业的了解程度如何？有什么具体的学习计划？"],
    ],
    "phone_guide": {
        "stages": {
            "阶段一：开场破冰（2-3分钟）": "您好，我是XX猎头，看到您在B2B产品领域有很深的积累，想和您聊一个智能座舱软件方向的机会...",
            "阶段二：现状摸底（5-7分钟）": "方便透露一下目前的薪资结构吗？另外现在对新机会的开放程度如何？",
        },
        "checklist": [
            ["稳定性意愿", "明确有换工作意向", "含糊推脱或近期刚入职"],
            ["薪资匹配", "期望在预算范围内", "期望超预算30%以上"],
        ],
        "taboos": ["不询问候选人婚育状况", "不透露客户公司全称（初次通话）"],
    },
    "verdict": "🏆 可推面试",
    "conclusion_summary": "候选人产品能力突出，稳定性良好，薪资在预算范围内。主要风险为无汽车行业背景，建议面试中重点考察学习意愿和行业认知深度。",
}

if __name__ == "__main__":
    main()
