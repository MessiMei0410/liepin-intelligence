"""
Generate a recruitment strategy .docx from a markdown file or inline content.

Usage:
    python3 scripts/gen_strategy_docx.py <input.md> [output.docx]

Styling:
    - Title: 微软雅黑 18pt, dark blue (#003366), centered
    - Body: 微软雅黑 10.5pt
    - Tables: Light Grid Accent 1 style, 9pt
    - Margins: 2.5cm all sides

The script reads markdown with ## sections and | tables,
converts them to python-docx paragraphs and tables with proper styling.
"""
import sys
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def build_docx(md_path: str, output_path: str, title: str = None, date: str = None):
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    if title:
        h = doc.add_heading(title, level=0)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            run.font.size = Pt(18)

    if date:
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dr = dp.add_run(date)
        dr.font.size = Pt(10)
        dr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # Parse markdown
    with open(md_path, 'r') as f:
        content = f.read()

    sections = re.split(r'\n(?=## )', content)
    for section in sections:
        lines = section.strip().split('\n')
        if not lines:
            continue

        # Heading
        hm = re.match(r'^(#{1,3})\s+(.+)', lines[0])
        if hm:
            level = min(len(hm.group(1)), 3)
            h = doc.add_heading(hm.group(2), level=level)
            for run in h.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            lines = lines[1:]

        # Process lines — tables and text
        table_rows = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                if table_rows:
                    _flush_table(doc, table_rows)
                    table_rows = []
                continue

            # Table row
            if stripped.startswith('|') and '|' in stripped[1:]:
                if '---' not in stripped:
                    cells = [c.strip() for c in stripped.split('|')[1:-1]]
                    if cells:
                        table_rows.append(cells)
                continue

            # Flush pending table
            if table_rows:
                _flush_table(doc, table_rows)
                table_rows = []

            # Text
            cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', stripped)
            cleaned = re.sub(r'`([^`]+)`', r'\1', cleaned)
            cleaned = re.sub(r'^[-*]\s+', '• ', cleaned)
            if cleaned:
                p = doc.add_paragraph(cleaned)
                for run in p.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10.5)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # Flush trailing table
        if table_rows:
            _flush_table(doc, table_rows)

    doc.save(output_path)
    print(f'Saved: {output_path}')

def _flush_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = text
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = '微软雅黑'
                        r.font.size = Pt(9)
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

if __name__ == '__main__':
    md_path = sys.argv[1] if len(sys.argv) > 1 else None
    out_path = sys.argv[2] if len(sys.argv) > 2 else md_path.replace('.md', '.docx')
    if not md_path:
        print("Usage: python gen_strategy_docx.py <input.md> [output.docx]")
        sys.exit(1)
    build_docx(md_path, out_path)
