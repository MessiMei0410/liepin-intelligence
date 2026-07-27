#!/usr/bin/env python3
"""Generate Word analysis report from recruitment data + analysis dict."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime, sys, json

# Read JSON from stdin: {"positions": [...], "analysis": {...}}
inp = json.load(sys.stdin)
positions = inp["positions"]
analysis = inp.get("analysis", {})

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10.5)
    run.bold = bold

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = '微软雅黑'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

# Cover
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('半导体设备招聘需求分析报告')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run(f'分析日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
doc.add_page_break()

# Sections (template — customize per analysis)
add_heading_styled('一、需求全景', 1)
add_para(f"本次共梳理 {len(positions)} 个招聘岗位，总缺口 {sum(p['gap'] for p in positions)} 人。")

# Department table
dept_data = analysis.get("departments", [])
if dept_data:
    add_heading_styled('1.1 部门维度', 2)
    add_table(['部门', '岗位数', '缺口', '核心方向'], dept_data, [3, 2, 2, 6])

# Position list
pos_rows = [[p.get('title',''), p.get('dept',''), p.get('team',''), p.get('level',''),
             p.get('edu',''), p.get('exp',''), str(p.get('gap','')), p.get('deadline','')] for p in positions]
add_heading_styled('1.2 岗位清单', 2)
add_table(['岗位名称', '部门', '小组', '层级', '学历', '经验', '缺口', '到岗'], pos_rows,
          [3.5, 2, 2.5, 1.5, 1.5, 1.5, 1, 2])

# Strategy signals
signals = analysis.get("signals", [])
if signals:
    doc.add_page_break()
    add_heading_styled('二、战略信号解读', 1)
    for s in signals:
        add_para(f'➤ {s["title"]}', bold=True)
        add_para(s["desc"], indent=True)

# Difficulty
diff_data = analysis.get("difficulty", [])
if diff_data:
    add_heading_styled('三、招聘难度评估', 1)
    add_table(['难度', '岗位', '原因', '建议策略'], diff_data, [2, 3, 5.5, 5.5])

# Summary
summary = analysis.get("summary", "")
if summary:
    add_heading_styled('四、总结', 1)
    add_para(summary)

out = sys.argv[1] if len(sys.argv) > 1 else "/Users/messi/Desktop/招聘需求分析报告.docx"
doc.save(out)
print(f"✅ {out}")
